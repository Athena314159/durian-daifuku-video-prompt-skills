#!/usr/bin/env python3
"""Export the sole user-facing Jimeng deliverable from one fresh compile snapshot."""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import subprocess
import sys
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from correction_memory import normalize_memory
from pipeline import canonical_input_hashes, normalized_prompt_length_contract, normalized_skill_release_lock


BLUE = RGBColor(43, 102, 158)
ORANGE = RGBColor(237, 123, 40)
PALE_BLUE = "EEF4FA"
PERFORMANCE_LAYERS = (
    ("emotion_trigger", "情绪触发"),
    ("gaze", "视线"),
    ("facial_microreaction", "面部微反应"),
    ("body_hand_preparation", "身体与手部准备"),
    ("breath_pause", "呼吸与停顿"),
    ("voice_speech", "声音与口播"),
)

# ``pipeline.py compile`` stores each canonical JSON input twice: once as a
# byte-level hash in generation_pack.json and once as editable JSON inside the
# immutable input snapshot.  Export must prove both halves.  Merely trusting a
# copied hash map would allow an empty or mixed snapshot to masquerade as the
# compile that produced the Prompt files.
SNAPSHOT_PAYLOAD_FIELDS = {
    "project.json": "project",
    "library/product_bible.json": "product_bible",
    "library/product_library.json": "product_library",
    "library/style_bible.json": "style_bible",
    "library/correction_memory.json": "correction_memory",
    "library/knowledge_index.json": "knowledge_index",
    "library/avatar_library.json": "avatar_library",
    "planning/story_plan.json": "story_plan",
    "planning/asset_reuse_plan.json": "asset_reuse_plan",
    "source/source_manifest.json": "source_manifest",
    "shots/shot_manifest.json": "shot_manifest",
}


def read_json(path: Path) -> dict[str, Any]:
    with path.open(encoding="utf-8") as handle:
        value = json.load(handle)
    if not isinstance(value, dict):
        raise ValueError(f"Expected JSON object: {path}")
    return value


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def digest_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def resolve(project: Path, value: str | None) -> Path | None:
    if not value:
        return None
    path = Path(value).expanduser()
    return path if path.is_absolute() else project / path


def prompt_from_markdown(path: Path) -> str:
    value = path.read_text(encoding="utf-8")
    match = re.search(r"```text\s*\n(.*?)\n```", value, re.S)
    if not match:
        raise ValueError(f"No copyable text code block in {path}")
    return match.group(1).strip()


def as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    return value if isinstance(value, list) else [value]


def has_complete_performance_layers(unit: dict[str, Any]) -> bool:
    layers = unit.get("source_performance_layers")
    return isinstance(layers, dict) and all(isinstance(layers.get(key), dict) for key, _ in PERFORMANCE_LAYERS)


def expected_pack_shot_payload(shot: dict[str, Any]) -> dict[str, Any]:
    """Project the canonical shot fields that compile must freeze verbatim."""
    asset_links = shot.get("asset_links") or {}
    source_units = [item for item in (shot.get("source_units") or []) if isinstance(item, dict)]
    inserted_units = [item for item in (shot.get("inserted_units") or []) if isinstance(item, dict)]
    return {
        "shot_id": shot.get("id"),
        "title": shot.get("title"),
        "timecode": shot.get("timecode") or {},
        "visual_type": shot.get("visual_type"),
        "narrative_role": shot.get("narrative_role"),
        "delivery_mode": (shot.get("audio") or {}).get("delivery_mode"),
        "script_segment_ids": as_list(shot.get("script_segment_ids")),
        "risk": shot.get("risk"),
        "source_shot_ids": [item.get("source_shot_id") for item in source_units],
        "inserted_shot_ids": [item.get("inserted_shot_id") for item in inserted_units],
        "source_first_frame": asset_links.get("source_first_frame"),
        "selected_beauty_keyframe": asset_links.get("selected_beauty_keyframe"),
        "approved_generation_first_frame": asset_links.get("approved_generation_first_frame"),
        "product_references": as_list(asset_links.get("product_references")),
        "prompt_file": f"prompts/{shot.get('id')}.md",
        "source_units": source_units,
        "inserted_units": inserted_units,
    }


def validate_unit_delivery_asset_binding(
    project_dir: Path,
    shot_id: str,
    item_id: str,
    unit: dict[str, Any],
    inventory: dict[str, dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[str]]:
    """Resolve every ordered asset owned by one SRC/ADD and verify its bytes."""
    errors: list[str] = []
    asset_ids = [str(value) for value in unit.get("delivery_asset_ids", [])]
    if not asset_ids:
        return [], [f"{shot_id}/{item_id}: at least one approved target-frame asset_id is required"]
    if len(asset_ids) != len(set(asset_ids)):
        errors.append(f"{shot_id}/{item_id}: delivery_asset_ids contains a duplicate")
    assets: list[dict[str, Any]] = []
    seen_paths: set[str] = set()
    seen_hashes: set[str] = set()
    for asset_id in asset_ids:
        asset = inventory.get(asset_id)
        if not isinstance(asset, dict):
            errors.append(f"{shot_id}/{item_id}: delivery asset {asset_id} is absent from inventory")
            continue
        assets.append(asset)
        if asset.get("approval_status") != "user_approved":
            errors.append(f"{shot_id}/{item_id}: delivery asset {asset_id} was not approved by the user after the full gallery was shown")
        user_approval = asset.get("user_approval") if isinstance(asset.get("user_approval"), dict) else {}
        if user_approval.get("status") != "user_approved" or user_approval.get("asset_sha256") != asset.get("sha256"):
            errors.append(f"{shot_id}/{item_id}: delivery asset {asset_id} lacks a byte-bound user approval receipt")
        asset_path = resolve(project_dir, asset.get("path"))
        if asset_path is None or not asset_path.is_file():
            errors.append(f"{shot_id}/{item_id}: approved image file missing for {asset_id}")
            continue
        resolved_path = str(asset_path.resolve())
        recorded_hash = asset.get("sha256")
        actual_hash = digest(asset_path)
        if resolved_path in seen_paths or actual_hash in seen_hashes:
            errors.append(f"{shot_id}/{item_id}: target-frame images must be distinct action states; duplicate path/hash found for {asset_id}")
        seen_paths.add(resolved_path)
        seen_hashes.add(actual_hash)
        if not isinstance(recorded_hash, str) or not recorded_hash:
            errors.append(f"{shot_id}/{item_id}: inventory sha256 missing for {asset_id}")
        elif recorded_hash != actual_hash:
            errors.append(f"{shot_id}/{item_id}: approved image hash differs from inventory for {asset_id}")
    return assets, errors


def delivery_asset_responsibility(unit: dict[str, Any], asset: dict[str, Any]) -> str:
    asset_id = str(asset.get("asset_id") or "")
    overrides = unit.get("delivery_asset_roles") or {}
    if isinstance(overrides, dict) and isinstance(overrides.get(asset_id), str) and overrides[asset_id].strip():
        return overrides[asset_id].strip()
    for field in ("responsibility", "frame_role", "caption"):
        value = asset.get(field)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return ""


def target_frame_caption(item_id: str, asset: dict[str, Any], responsibility: str, *, continuity: bool = False) -> str:
    continuity_label = "｜连续边界参考" if continuity else ""
    return f"{item_id}｜{asset.get('asset_id')}｜已批准{continuity_label}｜职责：{responsibility}"


def storyboard_units(shot: dict[str, Any]) -> list[tuple[str, dict[str, Any]]]:
    values: list[tuple[str, dict[str, Any]]] = []
    values.extend(("source", unit) for unit in (shot.get("source_units") or []) if isinstance(unit, dict))
    values.extend(("inserted", unit) for unit in (shot.get("inserted_units") or []) if isinstance(unit, dict))
    return sorted(values, key=lambda item: float((item[1].get("generation_timecode") or {}).get("start", 0)))


def unit_id(unit: dict[str, Any]) -> str:
    return str(unit.get("source_shot_id") or unit.get("inserted_shot_id") or "UNIT")


def exact_time_label(kind: str, unit: dict[str, Any]) -> str:
    generated = unit.get("generation_timecode") or {}
    if kind == "source":
        source = unit.get("source_timecode") or {}
        return (
            f"原片 {float(source.get('start', 0)):.3f}–{float(source.get('end', 0)):.3f} 秒｜"
            f"生成镜内 {float(generated.get('start', 0)):.3f}–{float(generated.get('end', 0)):.3f} 秒"
        )
    return (
        "新增镜头（无原片秒数）｜"
        f"生成镜内 {float(generated.get('start', 0)):.3f}–{float(generated.get('end', 0)):.3f} 秒"
    )


def format_performance_layers(unit: dict[str, Any]) -> str | None:
    layers = unit.get("source_performance_layers")
    if not isinstance(layers, dict):
        return None
    parts: list[str] = []
    for key, label in PERFORMANCE_LAYERS:
        evidence = layers.get(key)
        if not isinstance(evidence, dict):
            continue
        status = str(evidence.get("status") or "未记录")
        observable = str(evidence.get("observable_evidence") or "未记录")
        source_timecode = evidence.get("source_timecode")
        if isinstance(source_timecode, dict) and "start" in source_timecode and "end" in source_timecode:
            time_text = f"{float(source_timecode['start']):.3f}–{float(source_timecode['end']):.3f}秒"
        else:
            time_text = "无源片秒数"
        # The machine manifest retains the exact reference path and hash.  The
        # user-facing Word only needs to know whether a reference was locked;
        # local/project-relative paths are not actionable delivery content.
        frame = "有参考帧（内部已锁定）" if evidence.get("source_reference_frame") else "无参考帧"
        confidence = evidence.get("confidence")
        confidence_text = f"；置信度 {confidence}" if confidence is not None else ""
        gap_reason = evidence.get("gap_reason")
        gap_text = f"；模板补缺原因 {gap_reason}" if gap_reason not in (None, "") else ""
        parts.append(f"{label}[{status}] {time_text}｜{frame}｜可观察证据 {observable}{gap_text}{confidence_text}")
    return "；".join(parts) if parts else None


def matching_break_occurrences(story: dict[str, Any], shot_id: str, item_id: str) -> list[dict[str, Any]]:
    matches: list[dict[str, Any]] = []
    for occurrence in ((story.get("break_plan") or {}).get("occurrences") or []):
        if not isinstance(occurrence, dict) or str(occurrence.get("shot_id")) != shot_id:
            continue
        source_ids = occurrence.get("source_shot_ids") or []
        if isinstance(source_ids, str):
            source_ids = [source_ids]
        bound_ids = {
            str(value)
            for value in (
                [occurrence.get("source_shot_id"), occurrence.get("inserted_shot_id")]
                + list(source_ids)
            )
            if value
        }
        if item_id in bound_ids:
            matches.append(occurrence)
    return matches


def matching_eating_occurrences(story: dict[str, Any], shot_id: str, item_id: str) -> list[dict[str, Any]]:
    matches: list[dict[str, Any]] = []
    for occurrence in ((story.get("eating_plan") or {}).get("occurrences") or []):
        if not isinstance(occurrence, dict):
            continue
        if str(occurrence.get("shot_id")) == shot_id and str(occurrence.get("unit_id")) == item_id:
            matches.append(occurrence)
    return matches


def format_eating_occurrence(occurrence: dict[str, Any]) -> str:
    timecode = occurrence.get("generation_timecode") or {}
    phases = " → ".join(map(str, occurrence.get("required_phases") or []))
    transition = occurrence.get("speech_transition_rule") or occurrence.get("speech_transition")
    if not transition:
        transition = "产品离嘴且闭口咀嚼结束/嘴部可说时可按原片节奏马上接口播；不强制吞咽或吃后反应。"
    return (
        f"{occurrence.get('id')}｜{occurrence.get('origin')}｜"
        f"生成镜内 {float(timecode.get('start', 0)):.3f}–{float(timecode.get('end', 0)):.3f} 秒｜"
        f"节奏锚点 {occurrence.get('rhythm_anchor')}｜口播锚点 {occurrence.get('script_anchor')}｜"
        f"动作阶段 {phases}｜说话衔接 {transition}"
    )


def format_break_occurrence(occurrence: dict[str, Any]) -> str:
    proof = occurrence.get("crisp_proof") or {}
    crumbs = proof.get("crumbs") or {}
    timecode = occurrence.get("generation_timecode") or {}
    source_evidence_values = occurrence.get("source_evidence") or []
    if isinstance(source_evidence_values, str):
        source_evidence_values = [source_evidence_values]
    source_evidence = "、".join(map(str, source_evidence_values)) or "无（新增镜头）"
    insertion_rationale = str(occurrence.get("insertion_rationale") or "无（原片镜头）")
    return (
        f"{occurrence.get('id')}｜{occurrence.get('mode')}｜{occurrence.get('origin')}｜"
        f"生成镜内 {float(timecode.get('start', 0)):.3f}–{float(timecode.get('end', 0)):.3f} 秒｜"
        f"节奏依据 {occurrence.get('rhythm_rationale')}｜源片证据 {source_evidence}｜新增依据 {insertion_rationale}｜"
        f"同一根一次咔嚓={proof.get('single_snap')}；断点清楚={proof.get('fracture_visible')}；"
        f"两段质量守恒={proof.get('material_conservation_locked')}；"
        f"碎屑 {crumbs.get('minimum')}–{crumbs.get('maximum')}；"
        f"互补橙金断面={proof.get('complementary_orange_gold_fracture')}；"
        f"同一根两段守恒={proof.get('same_stick_two_piece_conservation')}；"
        f"音画同步={proof.get('sound_sync')}；拟音={proof.get('foley')}"
    )


def compact_json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def format_package_face(face: dict[str, Any]) -> str:
    evidence = face.get("qa_evidence") or {}
    checkpoints = evidence.get("visible_region_checkpoints") or []
    checkpoint_text = "、".join(
        f"{item.get('id')}={item.get('status')}" for item in checkpoints if isinstance(item, dict)
    ) or "无"
    expected = "、".join(map(str, face.get("expected_visible_regions") or [])) or "无（隐藏面）"
    return (
        f"{face.get('box_id')}/{face.get('face')}｜visibility_state={face.get('visibility_state')}｜"
        f"visible_extent={face.get('visible_extent')}｜polygon={compact_json(face.get('expected_visible_polygon'))}｜"
        f"area_ratio={face.get('visible_area_ratio')}｜legibility_required={compact_json(face.get('legibility_required'))}｜"
        f"master={face.get('master_reference')}｜expected_checkpoints={expected}｜"
        f"crop={evidence.get('candidate_face_crop')}｜crop_sha256={evidence.get('candidate_face_crop_sha256')}｜"
        f"parent_asset={evidence.get('delivery_asset_id')}｜parent_sha256={evidence.get('parent_image_sha256')}｜"
        f"crop_rect={compact_json(evidence.get('crop_rect_xywh'))}｜master_sha256={evidence.get('master_sha256')}｜"
        f"checkpoint_results={checkpoint_text}｜text_legibility={evidence.get('text_legibility')}｜"
        f"orientation={evidence.get('orientation')}｜cross_edge={evidence.get('cross_edge_registration')}｜"
        f"occlusion_scope={evidence.get('occlusion_scope')}｜qa_status={face.get('qa_status')}｜"
        f"hidden_reason={face.get('not_applicable_reason')}"
    )


def format_package_face_for_word(face: dict[str, Any]) -> str:
    """Editable user-facing packaging proof without paths, hashes or raw JSON."""
    evidence = face.get("qa_evidence") or {}
    checkpoints = evidence.get("visible_region_checkpoints") or []
    checkpoint_text = "、".join(
        f"{item.get('id')}={item.get('status')}" for item in checkpoints if isinstance(item, dict)
    ) or "无（隐藏面）"
    parent_asset = evidence.get("delivery_asset_id") or "未绑定"
    return (
        f"{face.get('box_id')}/{face.get('face')}｜可见状态 {face.get('visibility_state')}｜"
        f"可见范围 {face.get('visible_extent')}｜自然遮挡/出框 {face.get('natural_crop_or_occlusion')}｜"
        f"母版投射方式 {face.get('projection_method')}｜正文目标帧 {parent_asset}｜"
        f"应见区域核对 {checkpoint_text}｜文字 {evidence.get('text_legibility')}｜"
        f"方向 {evidence.get('orientation')}｜跨棱登记 {evidence.get('cross_edge_registration')}｜"
        f"遮挡范围 {evidence.get('occlusion_scope')}｜批准状态 {face.get('qa_status')}"
    )


def clock_label(seconds: Any) -> str:
    value = max(0.0, float(seconds or 0.0))
    minutes = int(value // 60)
    remainder = value - minutes * 60
    return f"{minutes:02d}:{remainder:06.3f}"


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    props.append(shade)


def set_cell_margins(cell, top: int = 120, start: int = 150, bottom: int = 120, end: int = 150) -> None:
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_run(paragraph, text: str, size: float = 10.5, color: RGBColor | None = None, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Arial Unicode MS"
    for family in ("ascii", "hAnsi", "eastAsia", "cs"):
        run._element.rPr.rFonts.set(qn(f"w:{family}"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_label(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    add_run(paragraph, label + "：", 9.5, BLUE, True)
    add_run(paragraph, value, 9.5)


def add_header_footer(document: Document, project_id: str) -> None:
    section = document.sections[0]
    add_run(section.header.paragraphs[0], project_id, 8, RGBColor(128, 136, 146))
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_image(cell, path: Path, width: float, caption: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = cell.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(caption_paragraph, caption, 8, RGBColor(100, 100, 100))


def rerun_lint(project_dir: Path) -> dict[str, Any]:
    result = subprocess.run(
        [sys.executable, str(Path(__file__).resolve().parent / "pipeline.py"), "lint", "--project-dir", str(project_dir)],
        capture_output=True,
        text=True,
        check=False,
    )
    report_path = project_dir / "review" / "lint_report.json"
    if not report_path.is_file():
        raise SystemExit("DOCX export blocked: fresh lint did not produce review/lint_report.json")
    report = read_json(report_path)
    if result.returncode == 2 or report.get("counts", {}).get("ERROR", 0):
        raise SystemExit("DOCX export blocked by fresh project lint:\n" + (result.stdout or result.stderr))
    return report


def validate_compile_snapshot(
    project_dir: Path,
    project: dict[str, Any],
    shots: list[dict[str, Any]],
    pack: dict[str, Any],
) -> tuple[dict[str, Any], Path, list[str]]:
    errors: list[str] = []
    compile_id = pack.get("compile_id")
    if not isinstance(compile_id, str) or not compile_id:
        errors.append("generation_pack.compile_id missing")
    current_hashes = canonical_input_hashes(project_dir)
    if pack.get("canonical_input_hashes") != current_hashes:
        errors.append("canonical inputs changed after compile; rerun pipeline.py compile")
    contract = normalized_prompt_length_contract(project)
    if pack.get("prompt_length_contract") != contract:
        errors.append("project prompt_length_contract differs from the compile snapshot")
    release_lock = normalized_skill_release_lock(project)
    legacy_release = release_lock.get("bundle_release_id") == "unmanaged-legacy"
    if pack.get("skill_release_lock") != release_lock and not (legacy_release and pack.get("skill_release_lock") is None):
        errors.append("project skill_release_lock differs from the compile snapshot")

    history_value = pack.get("history_dir")
    history_dir = resolve(project_dir, history_value if isinstance(history_value, str) else None)
    expected_history_root = (project_dir / "prompts" / "history").resolve()
    if history_dir is None:
        history_dir = expected_history_root / "__missing__"
        errors.append("generation_pack.history_dir missing")
    else:
        history_dir = history_dir.resolve()
        if not history_dir.is_relative_to(expected_history_root) or history_dir.name != compile_id:
            errors.append("generation_pack.history_dir is not the immutable directory for compile_id")
    history_pack_path = history_dir / "generation_pack.json"
    snapshot_path = history_dir / "input_snapshot.json"
    if not history_pack_path.is_file():
        errors.append("historical generation_pack.json missing")
    else:
        history_pack = read_json(history_pack_path)
        if history_pack != pack:
            errors.append("active generation_pack differs from its immutable historical copy")
    snapshot: dict[str, Any] = {}
    if not snapshot_path.is_file():
        errors.append("input_snapshot.json missing for compile_id")
    else:
        snapshot = read_json(snapshot_path)
        if snapshot.get("compile_id") != compile_id:
            errors.append("input_snapshot.compile_id mismatch")
        if snapshot.get("canonical_input_hashes") != current_hashes:
            errors.append("input_snapshot canonical hashes do not match current inputs")
        if snapshot.get("prompt_length_contract") != contract:
            errors.append("input_snapshot prompt length contract mismatch")
        if snapshot.get("skill_release_lock") != release_lock and not (
            legacy_release and snapshot.get("skill_release_lock") is None
        ):
            errors.append("input_snapshot Skill release lock mismatch")
        current_keys = set(current_hashes)
        mapped_keys = set(SNAPSHOT_PAYLOAD_FIELDS)
        if current_keys != mapped_keys:
            errors.append(
                "canonical snapshot payload mapping is incomplete: "
                f"missing={sorted(current_keys - mapped_keys)} extra={sorted(mapped_keys - current_keys)}"
            )
        for relative_path, snapshot_field in SNAPSHOT_PAYLOAD_FIELDS.items():
            canonical_path = project_dir / relative_path
            if snapshot_field not in snapshot:
                errors.append(f"input_snapshot.{snapshot_field} missing")
                continue
            current_value = read_json(canonical_path) if canonical_path.is_file() else None
            if relative_path == "library/correction_memory.json" and isinstance(current_value, dict):
                current_value, _ = normalize_memory(
                    current_value,
                    project_id=str(project.get("project_id") or ""),
                    product_profile=str(project.get("product_profile") or ""),
                    style_profile=str(project.get("style_profile") or ""),
                )
            if canonical_path.is_file() and snapshot.get(snapshot_field) != current_value:
                errors.append(f"input_snapshot.{snapshot_field} differs from current {relative_path}")

    current_source = read_json(project_dir / "source" / "source_manifest.json")
    if pack.get("project_id") != project.get("project_id"):
        errors.append("generation_pack.project_id differs from current project")
    if pack.get("source_sha256") != current_source.get("sha256"):
        errors.append("generation_pack.source_sha256 differs from current source manifest")

    expected_shot_ids = [str(shot.get("id")) for shot in shots]
    pack_shots = pack.get("shots") or []
    if [str(item.get("shot_id")) for item in pack_shots if isinstance(item, dict)] != expected_shot_ids:
        errors.append("generation_pack shot order/set differs from canonical shot_manifest")
    by_id = {str(item.get("shot_id")): item for item in pack_shots if isinstance(item, dict)}
    for shot in shots:
        shot_id = str(shot.get("id"))
        item = by_id.get(shot_id)
        if item is None:
            continue
        for field, expected in expected_pack_shot_payload(shot).items():
            if field not in item:
                errors.append(f"{shot_id}: generation_pack.{field} missing")
            elif item.get(field) != expected:
                errors.append(f"{shot_id}: generation_pack.{field} is stale or mixed")
        product_bible = read_json(project_dir / "library" / "product_bible.json")
        style_bible = read_json(project_dir / "library" / "style_bible.json")
        for field, expected in (
            ("product_profile", product_bible.get("profile_id")),
            ("product_version", product_bible.get("version")),
            ("style_profile", style_bible.get("profile_id")),
            ("style_version", style_bible.get("version")),
        ):
            if field not in item or item.get(field) != expected:
                errors.append(f"{shot_id}: generation_pack.{field} differs from the canonical bible")
        prompt_path = resolve(project_dir, item.get("prompt_file"))
        history_prompt = history_dir / f"{shot_id}.md"
        expected_prompt_path = (project_dir / "prompts" / f"{shot_id}.md").resolve()
        if prompt_path is None or prompt_path.resolve() != expected_prompt_path:
            errors.append(f"{shot_id}: prompt_file must be the canonical prompts/{shot_id}.md")
            prompt_path = expected_prompt_path
        if not prompt_path.is_file():
            errors.append(f"{shot_id}: canonical prompt file missing")
            continue
        prompt = prompt_from_markdown(prompt_path)
        if item.get("prompt_file_sha256") != digest(prompt_path):
            errors.append(f"{shot_id}: canonical prompt file changed after compile")
        if item.get("prompt_sha256") != digest_text(prompt):
            errors.append(f"{shot_id}: canonical Prompt text hash differs from compile snapshot")
        if not history_prompt.is_file() or digest(history_prompt) != digest(prompt_path):
            errors.append(f"{shot_id}: current prompt differs from immutable compile history")
        count = len(re.sub(r"\s+", "", prompt))
        if item.get("prompt_non_whitespace_characters") != count:
            errors.append(f"{shot_id}: stored Prompt character count is stale")
        if contract["enabled"] and not (
            contract["minimum_non_whitespace_characters"] <= count <= contract["maximum_non_whitespace_characters"]
        ):
            errors.append(
                f"{shot_id}: Prompt has {count} non-whitespace characters; enabled project contract requires "
                f"{contract['minimum_non_whitespace_characters']}–{contract['maximum_non_whitespace_characters']}"
            )
    return snapshot, snapshot_path, errors


def main() -> int:
    parser = argparse.ArgumentParser(description="Export a fresh, editable Jimeng storyboard DOCX.")
    parser.add_argument("--project-dir", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument(
        "--manifest-out",
        type=Path,
        help="Internal manifest path; defaults to <project>/review/<docx-stem>.manifest.json.",
    )
    # Accepted only to produce an explicit migration error.  Project JSON is
    # the one and only length-contract owner.
    parser.add_argument("--min-prompt-chars", type=int, default=None, help=argparse.SUPPRESS)
    parser.add_argument("--max-prompt-chars", type=int, default=None, help=argparse.SUPPRESS)
    args = parser.parse_args()
    if args.min_prompt_chars is not None or args.max_prompt_chars is not None:
        raise SystemExit(
            "DOCX export blocked: CLI Prompt-length flags were retired. Configure project.json.prompt_length_contract once."
        )

    project_dir = args.project_dir.expanduser().resolve()
    out = args.out.expanduser().resolve()
    workflow_path = project_dir / "planning" / "workflow_state.json"
    if workflow_path.is_file():
        workflow = read_json(workflow_path)
        if workflow.get("status") == "images_revoked" or workflow.get("docx_export_authorized") is False:
            raise SystemExit("DOCX export blocked: an image revocation is active; replace the revoked assets, show the complete updated gallery, record new user approval, and recompile first.")
    active_revocations: list[str] = []
    for revocation_path in sorted((project_dir / "review").glob("*revocation*.json")):
        try:
            if read_json(revocation_path).get("status") == "active":
                active_revocations.append(str(revocation_path))
        except (OSError, ValueError, json.JSONDecodeError):
            active_revocations.append(str(revocation_path))
    if active_revocations:
        raise SystemExit(
            "DOCX export blocked: active/unreadable revocation record(s) exist. Run invalidate_revoked_delivery.py and repair the affected assets:\n- "
            + "\n- ".join(active_revocations)
        )
    rerun_lint(project_dir)

    reuse_plan_path = project_dir / "planning" / "asset_reuse_plan.json"
    reuse_audit = subprocess.run(
        [sys.executable, str(Path(__file__).resolve().parent / "audit_asset_reuse.py"), "--plan", str(reuse_plan_path), "--stage", "pre-word"],
        capture_output=True,
        text=True,
        check=False,
    )
    if reuse_audit.returncode != 0:
        raise SystemExit("DOCX export blocked by asset reuse audit:\n" + (reuse_audit.stdout or reuse_audit.stderr))

    project = read_json(project_dir / "project.json")
    shots = read_json(project_dir / "shots" / "shot_manifest.json").get("shots", [])
    if not isinstance(shots, list):
        raise ValueError("shots/shot_manifest.json.shots must be a list")
    story = read_json(project_dir / "planning" / "story_plan.json")
    pack_path = project_dir / "prompts" / "generation_pack.json"
    pack = read_json(pack_path)
    snapshot, snapshot_path, snapshot_errors = validate_compile_snapshot(project_dir, project, shots, pack)
    if snapshot_errors:
        raise SystemExit("DOCX export blocked by stale or mixed compile snapshot:\n- " + "\n- ".join(snapshot_errors))

    pack_by_id = {str(item.get("shot_id")): item for item in (pack.get("shots") or []) if isinstance(item, dict)}
    prompt_contract = normalized_prompt_length_contract(project)
    reuse_plan = read_json(reuse_plan_path)
    reuse_assets = {str(item.get("asset_id")): item for item in (reuse_plan.get("inventory") or []) if isinstance(item, dict)}
    reuse_decisions = {str(item.get("shot_id")): item for item in (reuse_plan.get("shot_decisions") or []) if isinstance(item, dict)}

    def shot_delivery_assets(shot_id: str) -> list[dict[str, Any]]:
        decision = reuse_decisions.get(shot_id) or {}
        return [reuse_assets[item] for item in decision.get("selected_asset_ids", []) if item in reuse_assets]

    def unit_delivery_assets(shot: dict[str, Any], unit: dict[str, Any]) -> list[dict[str, Any]]:
        item_id = unit_id(unit)
        values, _ = validate_unit_delivery_asset_binding(
            project_dir, str(shot.get("id", "<unknown>")), item_id, unit, reuse_assets
        )
        return values

    export_errors: list[str] = []
    expected_unit_order: list[str] = []
    expected_asset_order: list[str] = []
    asset_owner_by_id: dict[str, str] = {}
    asset_owner_by_path: dict[str, str] = {}
    asset_owner_by_hash: dict[str, str] = {}
    canonical_units_by_id: dict[str, tuple[str, dict[str, Any], dict[str, Any]]] = {}
    for shot in shots:
        shot_id = str(shot.get("id", "<unknown>"))
        meta = pack_by_id.get(shot_id)
        if not meta:
            export_errors.append(f"{shot_id}: generation_pack entry missing")
            continue
        units = storyboard_units(shot)
        if not units:
            export_errors.append(f"{shot_id}: no source_units or inserted_units")
        expected_selected: list[str] = []
        for kind, unit in units:
            item_id = unit_id(unit)
            expected_unit_order.append(item_id)
            canonical_units_by_id[item_id] = (shot_id, unit, shot)
            if not unit.get("storyboard_description"):
                export_errors.append(f"{shot_id}/{item_id}: storyboard_description missing")
            if not unit.get("script_text"):
                export_errors.append(f"{shot_id}/{item_id}: script_text missing; use '无' only for a genuinely silent beat")
            if not has_complete_performance_layers(unit):
                export_errors.append(f"{shot_id}/{item_id}: all six structured performance layers are required")
            if kind == "inserted":
                if not unit.get("insertion_rationale") or not unit.get("rhythm_anchor"):
                    export_errors.append(f"{shot_id}/{item_id}: insertion_rationale or rhythm_anchor missing")
                if not unit.get("source_reference_shot_ids") or not unit.get("source_reference_frame"):
                    export_errors.append(f"{shot_id}/{item_id}: source references missing")
            assets, asset_errors = validate_unit_delivery_asset_binding(
                project_dir, shot_id, item_id, unit, reuse_assets
            )
            export_errors.extend(asset_errors)
            responsibilities: list[str] = []
            for asset in assets:
                asset_id = str(asset.get("asset_id"))
                responsibility = delivery_asset_responsibility(unit, asset)
                if not responsibility:
                    export_errors.append(f"{shot_id}/{item_id}/{asset_id}: target-frame responsibility missing")
                responsibilities.append(responsibility)
                asset_path = resolve(project_dir, asset.get("path"))
                if asset_path is None or not asset_path.is_file():
                    continue
                actual_path = str(asset_path.resolve())
                actual_hash = digest(asset_path)
                for key, owner_map, value in (
                    ("asset_id", asset_owner_by_id, asset_id),
                    ("path", asset_owner_by_path, actual_path),
                    ("sha256", asset_owner_by_hash, actual_hash),
                ):
                    prior = owner_map.get(value)
                    if prior and prior != item_id:
                        export_errors.append(f"{shot_id}/{item_id}: {key} is already owned by {prior}; cross-unit image reuse is forbidden")
                    owner_map[value] = item_id
                expected_selected.append(asset_id)
                expected_asset_order.append(asset_id)
            if len(assets) > 1 and len(responsibilities) == len(assets) and len(set(responsibilities)) != len(responsibilities):
                export_errors.append(f"{shot_id}/{item_id}: multi-frame responsibilities must be distinct")
        selected_ids = [str(value) for value in (reuse_decisions.get(shot_id) or {}).get("selected_asset_ids", [])]
        if selected_ids != expected_selected:
            export_errors.append(f"{shot_id}: selected_asset_ids must exactly follow its SRC/ADD card order")
        if project.get("product_mode") == "replace_product" and not meta.get("product_references"):
            export_errors.append(f"{shot_id}: product_references missing")
    if len(set(expected_unit_order)) != len(expected_unit_order):
        export_errors.append("SRC/ADD identifiers must be globally unique in DOCX order")
    if len(set(expected_asset_order)) != len(expected_asset_order):
        export_errors.append("Each selected target-frame asset_id must appear once in its owner SRC/ADD delivery list")

    shot_position = {str(shot.get("id")): index for index, shot in enumerate(shots)}
    continuity_display_count = 0
    for index, shot in enumerate(shots):
        shot_id = str(shot.get("id"))
        references = shot.get("continuity_boundary_references") or []
        if not isinstance(references, list):
            export_errors.append(f"{shot_id}: continuity_boundary_references must be a list")
            continue
        for reference in references:
            if not isinstance(reference, dict):
                export_errors.append(f"{shot_id}: invalid continuity boundary reference")
                continue
            owner_unit_id = str(reference.get("owner_unit_id") or "")
            asset_id = str(reference.get("asset_id") or "")
            owner = canonical_units_by_id.get(owner_unit_id)
            if reference.get("continuity_boundary_reference") is not True:
                export_errors.append(f"{shot_id}/{owner_unit_id}/{asset_id}: continuity_boundary_reference=true is required")
            if owner is None or asset_id not in [str(value) for value in owner[1].get("delivery_asset_ids", [])]:
                export_errors.append(f"{shot_id}/{owner_unit_id}/{asset_id}: boundary asset is not owned by the declared unit")
                continue
            if abs(shot_position.get(owner[0], -999) - index) != 1:
                export_errors.append(f"{shot_id}/{owner_unit_id}/{asset_id}: boundary display is not adjacent to the owner segment")
            if not isinstance(reference.get("responsibility"), str) or not reference["responsibility"].strip():
                export_errors.append(f"{shot_id}/{owner_unit_id}/{asset_id}: continuity responsibility missing")
            continuity_display_count += 1
    if export_errors:
        raise SystemExit("DOCX export blocked:\n- " + "\n- ".join(export_errors))

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.60)
    section.bottom_margin = Inches(0.60)
    section.left_margin = Inches(0.80)
    section.right_margin = Inches(0.80)
    add_header_footer(document, str(project.get("project_id", "")))

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(28)
    add_run(paragraph, "逐分镜执行稿", 28, BLUE, True)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(paragraph, f"{project.get('project_id')}｜{project.get('product_profile')}", 15, ORANGE, True)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if prompt_contract["enabled"]:
        prompt_contract_text = (
            f"Prompt {prompt_contract['minimum_non_whitespace_characters']}–"
            f"{prompt_contract['maximum_non_whitespace_characters']} 字硬门"
        )
    else:
        prompt_contract_text = "Prompt 长度硬门未启用，以完整可执行为准"
    source_count = sum(len(shot.get("source_units") or []) for shot in shots)
    inserted_count = sum(len(shot.get("inserted_units") or []) for shot in shots)
    add_run(
        paragraph,
        f"{source_count} 个原片分镜完整保留｜{inserted_count} 个有依据新增镜头｜"
        f"合并为 {len(shots)} 个连续生成片段｜{prompt_contract_text}",
        10,
        RGBColor(120, 130, 140),
    )

    document.add_page_break()
    heading = document.add_heading("当前结构结论", level=1)
    heading.runs[0].font.color.rgb = BLUE
    conclusions = [
        f"完整保留 {source_count} 个 SRC 原片原子分镜，并按原顺序合并为 {len(shots)} 个不少于4秒的连续生成段；短镜只合并，不删除。",
        f"按节奏加入 {inserted_count} 个 ADD 新增镜头；每个新增镜头均保留新增原因、节奏锚点与源片表演依据。",
        f"正文共绑定 {len(expected_asset_order)} 张批准目标帧；每个 SRC/ADD 至少1张，同一 unit 可保留多个不同动作关键状态，任何图不得冒充另一 unit。",
        "口播稿、准确时码、可编辑分镜描述、六层源证据、吃食/掰开/包装证据均写入正文；文字不会做成图片。",
        "吃食事件按全片节奏分散；一次事件的多个动作状态图仍属于同一事件，不会被误算成多次吃食。",
        "用户最终只接收这一份 Word；内部路径、JSON、handoff、alignment table 和审计文件不写入交付正文。",
    ]
    for value in conclusions:
        paragraph = document.add_paragraph(style="List Bullet")
        add_run(paragraph, value, 10)

    heading = document.add_heading("生成段总览", level=1)
    heading.runs[0].font.color.rgb = BLUE
    overview = document.add_table(rows=1, cols=4)
    overview.autofit = False
    headers = ("生成段", "动作镜头", "总时段", "主要职责")
    for cell, label in zip(overview.rows[0].cells, headers):
        set_cell_shading(cell, PALE_BLUE)
        set_cell_margins(cell)
        add_run(cell.paragraphs[0], label, 9, BLUE, True)
    for shot in shots:
        units = storyboard_units(shot)
        unit_ids = [unit_id(unit) for _, unit in units]
        tc = shot.get("timecode") or {}
        row = overview.add_row().cells
        values = (
            str(shot.get("id")),
            " + ".join(unit_ids),
            f"{clock_label(tc.get('start'))}–{clock_label(tc.get('end'))}",
            "；".join(str(unit.get("storyboard_description") or "") for _, unit in units),
        )
        for cell, value in zip(row, values):
            set_cell_margins(cell)
            add_run(cell.paragraphs[0], value, 8.5)

    manifest_shots: list[dict[str, Any]] = []
    for shot in shots:
        shot_id = str(shot.get("id"))
        meta = pack_by_id[shot_id]
        prompt_path = resolve(project_dir, meta.get("prompt_file"))
        prompt = prompt_from_markdown(prompt_path)
        units = storyboard_units(shot)
        shot_unit_ids = [unit_id(unit) for _, unit in units]
        shot_tc = shot.get("timecode") or {}
        segment_title = (
            f"{shot_id}｜{' + '.join(shot_unit_ids)}｜"
            f"{clock_label(shot_tc.get('start'))}–{clock_label(shot_tc.get('end'))}"
        )
        document.add_page_break()
        heading = document.add_heading(segment_title, level=1)
        heading.runs[0].font.color.rgb = BLUE
        add_label(
            document,
            "连续生成片段",
            f"{shot.get('title')}｜{float(shot_tc.get('duration') or 0):.3f} 秒｜"
            f"{shot.get('visual_type')}｜{shot.get('narrative_role')}",
        )
        shot_asset_count = sum(len(unit.get("delivery_asset_ids") or []) for _, unit in units)
        add_label(
            document,
            "对齐状态",
            f"动作镜头 {len(units)} 个｜批准目标帧 {shot_asset_count} 张｜"
            "本段全部动作状态帧必须一并使用，不能只上传段首图。",
        )
        if shot.get("merge_reason"):
            add_label(document, "合并原因", str(shot.get("merge_reason")))

        action_heading = document.add_heading("动作镜头对应", level=2)
        action_heading.runs[0].font.color.rgb = BLUE
        source_manifest: list[dict[str, Any]] = []
        inserted_manifest: list[dict[str, Any]] = []
        display_units: list[tuple[str, dict[str, Any], list[dict[str, Any]], dict[str, Any]]] = []
        for kind, unit in units:
            item_id = unit_id(unit)
            time_label = exact_time_label(kind, unit)
            subheading = document.add_heading(f"{item_id}｜{time_label}", level=3)
            subheading.runs[0].font.color.rgb = ORANGE
            add_label(document, "准确秒数", time_label)
            if kind == "inserted":
                add_label(document, "新增原因", str(unit.get("insertion_rationale")))
                add_label(document, "节奏锚点", str(unit.get("rhythm_anchor")))
                add_label(document, "源片表演依据", "、".join(map(str, unit.get("source_reference_shot_ids") or [])))
            add_label(document, "分镜描述", str(unit.get("storyboard_description")))
            add_label(document, "口播稿", str(unit.get("script_text") or "无"))
            layer_text = format_performance_layers(unit)
            eating_occurrences = matching_eating_occurrences(story, shot_id, item_id)
            eating_texts = [format_eating_occurrence(item) for item in eating_occurrences]
            for eating_text in eating_texts:
                add_label(document, "吃食节奏证据", eating_text)
            break_occurrences = matching_break_occurrences(story, shot_id, item_id)
            break_texts = [format_break_occurrence(item) for item in break_occurrences]
            for break_text in break_texts:
                add_label(document, "掰开酥脆证据", break_text)

            delivery_assets = unit_delivery_assets(shot, unit)
            asset_records: list[dict[str, Any]] = []
            for asset in delivery_assets:
                asset_path = resolve(project_dir, asset.get("path"))
                responsibility = delivery_asset_responsibility(unit, asset)
                asset_records.append(
                    {
                        "asset_id": asset.get("asset_id"),
                        "path": asset.get("path"),
                        "sha256": digest(asset_path),
                        "responsibility": responsibility,
                        "display_caption": target_frame_caption(item_id, asset, responsibility),
                    }
                )
            record: dict[str, Any] = {
                "generation_timecode": unit.get("generation_timecode") or {},
                "exact_time_label": time_label,
                "storyboard_description": unit.get("storyboard_description"),
                "script_text": unit.get("script_text"),
                "source_performance_layers": unit.get("source_performance_layers"),
                "performance_layers_text": layer_text,
                "eating_occurrences": eating_occurrences,
                "eating_occurrence_texts": eating_texts,
                "break_occurrences": break_occurrences,
                "break_occurrence_texts": break_texts,
                "delivery_assets": asset_records,
            }
            if kind == "source":
                record.update({"source_shot_id": item_id, "source_timecode": unit.get("source_timecode") or {}})
                source_manifest.append(record)
            else:
                record.update(
                    {
                        "inserted_shot_id": item_id,
                        "insertion_rationale": unit.get("insertion_rationale"),
                        "rhythm_anchor": unit.get("rhythm_anchor"),
                        "source_reference_shot_ids": unit.get("source_reference_shot_ids") or [],
                        "source_reference_frame": unit.get("source_reference_frame"),
                    }
                )
                inserted_manifest.append(record)
            display_units.append((item_id, unit, delivery_assets, record))

        frame_heading = document.add_heading("目标帧与职责", level=2)
        frame_heading.runs[0].font.color.rgb = BLUE
        for item_id, unit, delivery_assets, _ in display_units:
            add_label(document, "目标帧组", f"{item_id}｜{len(delivery_assets)} 张批准动作关键状态")
            frame_table = document.add_table(rows=0, cols=2)
            frame_table.autofit = False
            for asset_index in range(0, len(delivery_assets), 2):
                cells = frame_table.add_row().cells
                for offset, cell in enumerate(cells):
                    index = asset_index + offset
                    set_cell_margins(cell, 100, 100, 100, 100)
                    if index >= len(delivery_assets):
                        continue
                    asset = delivery_assets[index]
                    responsibility = delivery_asset_responsibility(unit, asset)
                    add_image(
                        cell,
                        resolve(project_dir, asset.get("path")),
                        2.20,
                        target_frame_caption(item_id, asset, responsibility),
                    )

        continuity_records: list[dict[str, Any]] = []
        for reference in shot.get("continuity_boundary_references") or []:
            owner_unit_id = str(reference.get("owner_unit_id"))
            asset_id = str(reference.get("asset_id"))
            asset = reuse_assets[asset_id]
            responsibility = str(reference.get("responsibility"))
            add_label(document, "连续边界参考", f"{owner_unit_id} 的批准目标帧只用于相邻段衔接，不计入本段其他 unit 的最低图片覆盖。")
            frame_table = document.add_table(rows=1, cols=1)
            frame_table.autofit = False
            caption = target_frame_caption(owner_unit_id, asset, responsibility, continuity=True)
            add_image(frame_table.cell(0, 0), resolve(project_dir, asset.get("path")), 2.40, caption)
            continuity_records.append(
                {
                    "owner_unit_id": owner_unit_id,
                    "asset_id": asset_id,
                    "sha256": digest(resolve(project_dir, asset.get("path"))),
                    "responsibility": responsibility,
                    "continuity_boundary_reference": True,
                    "display_caption": caption,
                }
            )

        assets = shot.get("asset_links") or {}
        package_faces = ((shot.get("product_state") or {}).get("package_artwork") or {}).get("visible_faces") or []
        package_face_texts = [format_package_face(face) for face in package_faces if isinstance(face, dict)]
        package_face_word_texts = [format_package_face_for_word(face) for face in package_faces if isinstance(face, dict)]
        if package_face_word_texts:
            package_heading = document.add_heading("产品与包装证据", level=2)
            package_heading.runs[0].font.color.rgb = BLUE
            add_label(document, "产品状态", str((shot.get("product_state") or {}).get("state", "")))
            for package_face_text in package_face_word_texts:
                add_label(document, "包装盒面", package_face_text)
        count = len(re.sub(r"\s+", "", prompt))
        prompt_heading = document.add_heading("可复制Prompt原文", level=2)
        prompt_heading.runs[0].font.color.rgb = BLUE
        paragraph = document.add_paragraph()
        add_run(paragraph, "即梦可复制 Prompt", 13, BLUE, True)
        add_run(paragraph, f"  {count} 字", 10, ORANGE, True)
        box = document.add_table(rows=1, cols=1).cell(0, 0)
        set_cell_shading(box, PALE_BLUE)
        set_cell_margins(box, 180, 180, 180, 180)
        paragraph = box.paragraphs[0]
        paragraph.paragraph_format.line_spacing = 1.08
        add_run(paragraph, prompt, 8.5)

        selected_assets = shot_delivery_assets(shot_id)
        manifest_shots.append(
            {
                "shot_id": shot_id,
                "source_first_frame": meta.get("source_first_frame"),
                "approved_generation_first_frame": meta.get("approved_generation_first_frame"),
                "source_units": source_manifest,
                "inserted_units": inserted_manifest,
                "selected_asset_ids": [str(item.get("asset_id")) for item in selected_assets],
                "continuity_boundary_references": continuity_records,
                "avatar_reference": assets.get("avatar_reference"),
                "product_references": meta.get("product_references") or [],
                "package_faces": package_faces,
                "package_face_texts": package_face_texts,
                "package_face_word_texts": package_face_word_texts,
                "prompt_sha256": digest_text(prompt),
                "prompt_file_sha256": digest(prompt_path),
                "prompt_non_whitespace_characters": count,
            }
        )

    document.add_page_break()
    heading = document.add_heading("最终确认清单", level=1)
    heading.runs[0].font.color.rgb = BLUE
    final_checks = [
        f"全部 {source_count} 个 SRC 与 {inserted_count} 个 ADD 都有准确时码、可编辑分镜描述和可编辑口播；内部六层证据已在导出前审计，不作为凑字栏目写进 Word。",
        f"全部 {len(expected_asset_order)} 张经用户总览确认的目标帧均在对应 owner unit 下展示；同一 unit 可多图，跨 unit 不复用 asset/path/hash。",
        "小于4秒的原片分镜只与相邻镜头合并生成，未删除、未省略任何原片动作节点。",
        "吃食事件按新版口播和原片节奏分散；不把同一事件的多张关键状态图误算为多次吃食，也不强制源片没有的吞咽或吃后反应。",
        "黄油脆丝棒掰开镜按人物出镜与纯手无人出镜两类核验；同一根、一次咔嚓、互补橙金断面、3–8片克制掉渣和音画同步均保留。",
        "可见包装盒面只使用批准母版的确定性投射证据；自然遮挡/出框允许，模型重绘、镜像、乱码或图案缺块均不得批准。",
        "每个生成段末尾均保留完整可复制 Prompt 原文；用户交付仅为本 Word，不附内部路径、JSON 或对齐表。",
    ]
    for value in final_checks:
        paragraph = document.add_paragraph(style="List Bullet")
        add_run(paragraph, value, 10)

    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(out)
    expected_word_images = reuse_plan.get("summary", {}).get("expected_word_image_count")
    with zipfile.ZipFile(out) as archive:
        embedded_media_count = len([name for name in archive.namelist() if name.startswith("word/media/")])
        document_xml = archive.read("word/document.xml")
        body_image_relationship_count = len(re.findall(rb"<a:blip\b", document_xml))
        visible_internal_path_tokens = [
            token.decode("utf-8", errors="replace")
            for token in (
                b"/Users/",
                b"/private/",
                b"/Volumes/",
                b"file://",
                b"source/",
                b"assets/",
                b"library/",
                b"planning/",
                b"review/",
                b"prompts/",
            )
            if token in document_xml
        ]
    if expected_word_images != len(expected_asset_order):
        out.unlink(missing_ok=True)
        raise SystemExit(
            f"DOCX export blocked: asset_reuse_plan expected_word_image_count {expected_word_images} "
            f"does not equal the {len(expected_asset_order)} ordered owner-unit target frames."
        )
    if embedded_media_count != expected_word_images:
        out.unlink(missing_ok=True)
        raise SystemExit(
            f"DOCX export blocked: embedded image count {embedded_media_count} does not equal "
            f"asset_reuse_plan expected_word_image_count {expected_word_images}."
        )
    expected_body_occurrences = len(expected_asset_order) + continuity_display_count
    if body_image_relationship_count != expected_body_occurrences:
        out.unlink(missing_ok=True)
        raise SystemExit(
            f"DOCX export blocked: body image occurrence count {body_image_relationship_count} does not equal "
            f"owner frames plus continuity references ({expected_body_occurrences})."
        )
    if visible_internal_path_tokens:
        out.unlink(missing_ok=True)
        raise SystemExit(
            "DOCX export blocked: editable Word text exposes internal local paths: "
            + ", ".join(visible_internal_path_tokens)
        )

    manifest_path = (
        args.manifest_out.expanduser().resolve()
        if args.manifest_out
        else project_dir / "review" / f"{out.stem}.manifest.json"
    )
    manifest_value = {
        "schema_version": "2.0",
        "project_id": project.get("project_id"),
        "compile_id": pack.get("compile_id"),
        "canonical_input_hashes": pack.get("canonical_input_hashes"),
        "prompt_length_contract": prompt_contract,
        "skill_release_lock": normalized_skill_release_lock(project),
        "generation_pack_sha256": digest(pack_path),
        "input_snapshot_path": str(snapshot_path.relative_to(project_dir)),
        "input_snapshot_sha256": digest(snapshot_path),
        "source_sha256": pack.get("source_sha256"),
        "docx_sha256": digest(out),
        "exported_at": datetime.now().astimezone().replace(microsecond=0).isoformat(),
        "unit_order": expected_unit_order,
        "reused_frame_count": reuse_plan.get("summary", {}).get("reused_frame_count"),
        "new_generation_count": reuse_plan.get("summary", {}).get("new_generation_count"),
        "expected_word_image_count": expected_word_images,
        "embedded_media_count": embedded_media_count,
        "expected_body_image_occurrence_count": expected_body_occurrences,
        "body_image_relationship_count": body_image_relationship_count,
        "continuity_boundary_reference_count": continuity_display_count,
        "prompt_non_whitespace_characters": {
            item["shot_id"]: item["prompt_non_whitespace_characters"] for item in manifest_shots
        },
        "compile_snapshot_validated": True,
        "shots": manifest_shots,
    }
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(manifest_value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(
        json.dumps(
            {"docx": str(out), "manifest": str(manifest_path), "compile_id": pack.get("compile_id"), "shot_count": len(manifest_shots)},
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except (FileNotFoundError, KeyError, ValueError, json.JSONDecodeError, zipfile.BadZipFile) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(2)
