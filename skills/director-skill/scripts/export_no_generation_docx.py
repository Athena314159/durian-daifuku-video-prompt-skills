#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import annotations
import argparse, hashlib, json, re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from pipeline import canonical_input_hashes, sha256_file

MODE = "no_generation_prompt_docx_alignment"

def load(path): return json.loads(path.read_text(encoding="utf-8"))
def unit_id(unit): return unit.get("source_shot_id") or unit.get("inserted_shot_id")
def prompt_text(path):
    m=re.search(r"```text\s*\n(.*?)\n```",path.read_text(encoding="utf-8"),re.S)
    if not m: raise ValueError(f"Canonical Prompt missing: {path}")
    return m.group(1).strip()
def clear_body(doc):
    body=doc._element.body
    for child in list(body):
        if not child.tag.endswith('}sectPr'): body.remove(child)
def add_label(doc,label,text):
    p=doc.add_paragraph(); r=p.add_run(label); r.bold=True; p.add_run(text)
def main():
    ap=argparse.ArgumentParser(); ap.add_argument('--project-dir',required=True,type=Path); ap.add_argument('--template-docx',required=True,type=Path); ap.add_argument('--output',required=True,type=Path); a=ap.parse_args()
    root=a.project_dir.resolve(); project=load(root/'project.json')
    if project.get('document_delivery_mode')!=MODE or project.get('execution_tier')!='prompt_only': raise ValueError('NO_GENERATION_MODE_NOT_AUTHORIZED')
    contract=load(root/'planning/no_generation_prompt_docx_alignment_contract.json')
    if contract.get('schema_version')!='no-generation-prompt-docx-alignment-v1.0' or contract.get('status')!='ready_for_prompt_compile_and_docx_alignment' or contract.get('approved_target_frame_count')!=0: raise ValueError('NO_GENERATION_CONTRACT_INVALID')
    pack=load(root/'prompts/generation_pack.json'); receipt=load(root/'review/prompt_delivery_receipt.json')
    if receipt.get('status') not in {'authorized','delivery_authorized'} or receipt.get('compile_id')!=pack.get('compile_id'): raise ValueError('PROMPT_DELIVERY_NOT_AUTHORIZED')
    manifest=load(root/'shots/shot_manifest.json'); story=load(root/'planning/story_plan.json'); lock=load(root/'planning/revised_script_lock.json')
    refs={x['owner_unit_id']:x for x in contract['references']}
    doc=Document(str(a.template_docx)); clear_body(doc)
    sec=doc.sections[0]; sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.5); sec.left_margin=Cm(1.6); sec.right_margin=Cm(1.6)
    normal=doc.styles['Normal']; normal.font.name='PingFang SC'; normal.font.size=Pt(9)
    title=doc.add_heading('榴莲大福｜即梦逐分镜执行稿',0); title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('停止生图版｜Prompt + 旧版图片职责对齐'); r.bold=True; r.font.color.rgb=RGBColor(156,87,0)
    add_label(doc,'交付声明：','本文档不声称图片生成门通过。旧图仅按构图参考或错误示例展示；SRC006为精确源片复用；批准目标帧数量为0。')
    doc.add_heading('完整可编辑口播稿（249字）',1); doc.add_paragraph(lock['editable_text'])
    prompts=[]; unit_order=[]
    for shot in manifest['shots']:
        sid=shot['id']; doc.add_page_break() if sid!='S001' else None
        doc.add_heading(f"{sid}｜{shot.get('title','')}",1)
        add_label(doc,'镜头职责：',shot.get('purpose','')); add_label(doc,'镜头时间：',f"{shot['timecode']['start']}–{shot['timecode']['end']}秒（独立生成{shot['timecode']['duration']}秒）")
        doc.add_heading('动作镜头对应',2)
        for kind in ('source_units','inserted_units'):
            for unit in shot.get(kind,[]) or []:
                uid=unit_id(unit); unit_order.append(uid); tc=unit.get('generation_timecode') or {}
                doc.add_heading(f"{uid}｜生成镜内{tc.get('start',0):.2f}–{tc.get('end',0):.2f}秒",2)
                add_label(doc,'准确秒数：',f"{tc.get('start',0):.2f}–{tc.get('end',0):.2f}秒")
                add_label(doc,'分镜描述：',unit.get('storyboard_description',''))
                add_label(doc,'口播稿：',unit.get('script_text') or '无')
                ref=refs[uid]; img=root/ref['path'];
                if sha256_file(img)!=ref['sha256']: raise ValueError(f'REFERENCE_HASH_MISMATCH:{uid}')
                doc.add_picture(str(img),width=Cm(5.3)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
                cp=doc.add_paragraph(ref['caption']); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; cp.runs[0].italic=True
        doc.add_heading('逐时动作',2)
        for beat in shot.get('action_beats',[]) or []:
            doc.add_paragraph(f"{beat.get('start',0):.2f}–{beat.get('end',0):.2f}秒｜{beat.get('action')}｜声音：{beat.get('voice_change')}｜产品：{beat.get('product_change')}",style='List Bullet')
        doc.add_heading('可复制Prompt原文',2)
        prompt=prompt_text(root/f'prompts/{sid}.md'); prompts.append({'shot_id':sid,'sha256':hashlib.sha256(prompt.encode()).hexdigest(),'count':len(re.sub(r'\s+','',prompt))})
        if prompts[-1]['count']>=4000: raise ValueError(f'PROMPT_TOO_LONG:{sid}')
        doc.add_paragraph(prompt)
    a.output.parent.mkdir(parents=True,exist_ok=True); doc.save(str(a.output))
    out={
      'schema_version':'no-generation-docx-export-v1.0','status':'exported_for_render_qa','document_delivery_mode':MODE,
      'project_id':project.get('project_id'),'compile_id':pack.get('compile_id'),'generation_pack_sha256':sha256_file(root/'prompts/generation_pack.json'),
      'canonical_input_hashes':canonical_input_hashes(root),'prompt_length_contract':project.get('prompt_length_contract'),'skill_release_lock':project.get('skill_release_lock'),
      'docx_path':str(a.output),'docx_sha256':sha256_file(a.output),'unit_order':unit_order,'reference_count':len(contract['references']),
      'approved_target_frame_count':0,'prompts':prompts,'script_effective_character_count':249,'compile_snapshot_validated':True,
    }
    mp=root/'review'/f'{a.output.stem}.manifest.json'; mp.write_text(json.dumps(out,ensure_ascii=False,indent=2)+'\n',encoding='utf-8')
    print(json.dumps({'status':'exported','docx':str(a.output),'manifest':str(mp),'sha256':out['docx_sha256']},ensure_ascii=False))
if __name__=='__main__': raise SystemExit(main())
