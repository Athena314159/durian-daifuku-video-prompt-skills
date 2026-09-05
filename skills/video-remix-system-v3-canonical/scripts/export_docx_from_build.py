#!/usr/bin/env python3
"""Export a V3 Canonical build into an auditable DOCX.

The exporter is deliberately a consumer of compiled manifests.  It does not
make shot, script, prompt, or asset decisions.  Validation happens completely
before the DOCX is created so a blocked build can never leave a misleading
partial document behind.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import sys
from pathlib import Path
from typing import Any


REQUIRED = (
    "final_generation_manifest.json",
    "script_shot_map.json",
    "prompt_task_manifest.json",
    "image_task_manifest.json",
    "rule_receipt.json",
)


class ExportBlocked(RuntimeError):
    """Raised when a build cannot produce a trustworthy DOCX."""


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_json(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ExportBlocked(f"无法读取 {path.name}: {exc}") from exc
    if not isinstance(value, dict):
        raise ExportBlocked(f"{path.name}必须是JSON对象")
    return value


def _tasks_by_shot(manifest: dict[str, Any], name: str) -> dict[str, dict[str, Any]]:
    tasks = manifest.get("tasks")
    if not isinstance(tasks, list):
        raise ExportBlocked(f"{name}缺tasks数组")
    result: dict[str, dict[str, Any]] = {}
    for task in tasks:
        if not isinstance(task, dict) or not task.get("shot_id"):
            raise ExportBlocked(f"{name}存在无效任务")
        shot = str(task["shot_id"])
        if shot in result:
            raise ExportBlocked(f"{name}重复镜头 {shot}")
        result[shot] = task
    return result


def _image_path(task: dict[str, Any], build_dir: Path) -> Path | None:
    """Resolve the approved/generated QA image field used by provider runners."""
    for key in (
        "approved_image_path",
        "generated_image_path",
        "generated_image",
        "qa_image_path",
        "qa_image",
        "image_path",
        "output_path",
        "result_path",
        "result_image_path",
    ):
        value = task.get(key)
        if isinstance(value, str) and value.strip():
            path = Path(value).expanduser()
            return path if path.is_absolute() else (build_dir / path).resolve()
    for key in ("generated_images", "qa_images", "images"):
        values = task.get(key)
        if isinstance(values, list):
            for value in values:
                if isinstance(value, str) and value.strip():
                    path = Path(value).expanduser()
                    return path if path.is_absolute() else (build_dir / path).resolve()
                if isinstance(value, dict):
                    for path_key in ("path", "image_path", "file"):
                        path_value = value.get(path_key)
                        if isinstance(path_value, str) and path_value.strip():
                            path = Path(path_value).expanduser()
                            return path if path.is_absolute() else (build_dir / path).resolve()
    for key in ("generated_image", "qa_image", "result_image"):
        value = task.get(key)
        if isinstance(value, dict):
            for path_key in ("path", "image_path", "file"):
                path_value = value.get(path_key)
                if isinstance(path_value, str) and path_value.strip():
                    path = Path(path_value).expanduser()
                    return path if path.is_absolute() else (build_dir / path).resolve()
    return None


def _image_hash(task: dict[str, Any]) -> str | None:
    for key in (
        "approved_image_sha256",
        "generated_image_sha256",
        "generated_image_hash",
        "qa_image_sha256",
        "image_sha256",
        "output_sha256",
        "sha256",
    ):
        value = task.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip().lower()
    return None


def validate_build(build_dir: Path) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    """Validate all five canonical inputs and return ordered shot records."""
    missing = [name for name in REQUIRED if not (build_dir / name).is_file()]
    if missing:
        raise ExportBlocked("缺少Canonical编译产物：" + ", ".join(missing))

    final = read_json(build_dir / REQUIRED[0])
    script = read_json(build_dir / REQUIRED[1])
    prompts = read_json(build_dir / REQUIRED[2])
    images = read_json(build_dir / REQUIRED[3])
    receipt = read_json(build_dir / REQUIRED[4])

    if final.get("canonical_sha256") and receipt.get("canonical_sha256"):
        if final["canonical_sha256"] != receipt["canonical_sha256"]:
            raise ExportBlocked("final_generation_manifest与rule_receipt不是同一Canonical版本")
    declared_script_hash = receipt.get("script_shot_map_sha256") or final.get("script_shot_map_sha256")
    if declared_script_hash:
        actual_script_hash = sha256(build_dir / "script_shot_map.json")
        if actual_script_hash != str(declared_script_hash).lower():
            raise ExportBlocked("script_shot_map哈希不一致，build stale")

    selected = final.get("selected_shots")
    if not isinstance(selected, list) or not selected or any(not isinstance(x, str) for x in selected):
        raise ExportBlocked("final_generation_manifest.selected_shots为空或无效")
    if len(selected) != len(set(selected)):
        raise ExportBlocked("selected_shots包含重复镜头")

    script_lines = script.get("lines")
    if not isinstance(script_lines, list):
        raise ExportBlocked("script_shot_map缺lines数组")
    # A stale map is one that no longer covers the selected shots, has duplicate
    # line IDs, or is not in its declared source order.
    line_ids: list[str] = []
    lines_by_shot: dict[str, list[dict[str, Any]]] = {shot: [] for shot in selected}
    for line in script_lines:
        if not isinstance(line, dict) or not line.get("line_id") or not line.get("shot_id"):
            raise ExportBlocked("script_shot_map存在无效口播行")
        line_id, shot_id = str(line["line_id"]), str(line["shot_id"])
        if line_id in line_ids:
            raise ExportBlocked(f"script_shot_map重复口播 {line_id}")
        if shot_id not in lines_by_shot:
            raise ExportBlocked(f"script_shot_map引用未选镜头 {shot_id}")
        line_ids.append(line_id)
        lines_by_shot[shot_id].append(line)
    orders = [line.get("order") for line in script_lines]
    if any(not isinstance(order, (int, float)) for order in orders) or orders != list(range(1, len(orders) + 1)):
        raise ExportBlocked("script_shot_map口播顺序过期或无效")

    prompt_by_shot = _tasks_by_shot(prompts, "prompt_task_manifest")
    image_by_shot = _tasks_by_shot(images, "image_task_manifest")
    if set(prompt_by_shot) != set(selected):
        raise ExportBlocked("prompt_task_manifest与selected_shots不一致，可能是stale")
    if set(image_by_shot) != set(selected):
        raise ExportBlocked("image_task_manifest与selected_shots不一致，可能是stale")

    receipt_prompt_hashes = receipt.get("prompt_file_sha256")
    if not isinstance(receipt_prompt_hashes, dict):
        raise ExportBlocked("rule_receipt缺prompt_file_sha256，无法证明Prompt为当前版本")

    records: list[dict[str, Any]] = []
    for shot in selected:
        prompt_task = prompt_by_shot[shot]
        prompt_value = prompt_task.get("prompt_file")
        if not isinstance(prompt_value, str) or not prompt_value:
            raise ExportBlocked(f"{shot}缺Prompt文件")
        prompt_path = Path(prompt_value).expanduser()
        if not prompt_path.is_absolute():
            prompt_path = (build_dir / prompt_path).resolve()
        if not prompt_path.is_file():
            raise ExportBlocked(f"{shot} Prompt文件不存在：{prompt_path}")
        actual_prompt_hash = sha256(prompt_path)
        declared_prompt_hash = prompt_task.get("prompt_file_sha256")
        receipt_hash = receipt_prompt_hashes.get(shot)
        if not isinstance(declared_prompt_hash, str) or actual_prompt_hash != declared_prompt_hash.lower():
            raise ExportBlocked(f"{shot} Prompt哈希不一致，build stale")
        if not isinstance(receipt_hash, str) or actual_prompt_hash != receipt_hash.lower():
            raise ExportBlocked(f"{shot} Prompt与rule_receipt哈希不一致，build stale")

        mapped_lines = lines_by_shot[shot]
        declared_line_ids = prompt_task.get("line_ids")
        if declared_line_ids is not None and list(declared_line_ids) != [line["line_id"] for line in mapped_lines]:
            raise ExportBlocked(f"{shot} script与Prompt任务口播映射过期")

        image_task = image_by_shot[shot]
        status = str(image_task.get("image_status", "")).lower()
        if status == "awaiting_generation" or "awaiting" in status:
            raise ExportBlocked(f"{shot}图像仍awaiting_generation")
        image_path = _image_path(image_task, build_dir)
        if image_path is None or not image_path.is_file():
            raise ExportBlocked(f"{shot}缺批准/生成QA图")
        declared_image_hash = _image_hash(image_task)
        if not declared_image_hash:
            raise ExportBlocked(f"{shot}缺QA图哈希")
        actual_image_hash = sha256(image_path)
        if actual_image_hash != declared_image_hash:
            raise ExportBlocked(f"{shot} QA图哈希不一致")

        records.append({
            "shot_id": shot,
            "lines": mapped_lines,
            "prompt_path": str(prompt_path),
            "prompt_sha256": actual_prompt_hash,
            "image_path": str(image_path),
            "image_sha256": actual_image_hash,
            "image_status": image_task.get("image_status"),
        })
    return {"final": final, "receipt": receipt}, records


def export_docx(build_dir: Path, output_docx: Path, alignment_path: Path | None = None) -> Path:
    metadata, records = validate_build(build_dir)
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:  # pragma: no cover - environment-specific
        raise ExportBlocked("当前Python环境缺bundled python-docx") from exc

    output_docx = output_docx.expanduser().resolve()
    alignment_path = (alignment_path or output_docx.with_name(output_docx.stem + "_alignment_manifest.json")).resolve()
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(f"V3 Canonical交付 · {metadata['final'].get('project_id', '')}", level=0)
    document.add_paragraph("本文件由Canonical编译产物排版生成；每个镜头的口播、Prompt和QA图均来自同一组已校验manifest。")
    alignment: dict[str, Any] = {
        "schema": "v3-canonical-docx-alignment-v1",
        "project_id": metadata["final"].get("project_id"),
        "selected_shots": [record["shot_id"] for record in records],
        "shots": [],
    }
    for record in records:
        shot = record["shot_id"]
        heading_index = len(document.paragraphs)
        heading = document.add_heading(shot, level=1)
        script_heading_index = len(document.paragraphs)
        script_heading = document.add_paragraph()
        script_heading.add_run("当前口播").bold = True
        line_indices: list[int] = []
        for line in record["lines"]:
            line_indices.append(len(document.paragraphs))
            paragraph = document.add_paragraph(style=None)
            paragraph.add_run(f"{line['line_id']} · {line.get('speaker_id', '')}：").bold = True
            paragraph.add_run(str(line.get("text", "")))
        prompt_heading_index = len(document.paragraphs)
        prompt_heading = document.add_paragraph()
        prompt_heading.add_run("当前 Prompt").bold = True
        prompt_index = len(document.paragraphs)
        prompt_paragraph = document.add_paragraph(Path(record["prompt_path"]).read_text(encoding="utf-8"))
        image_heading_index = len(document.paragraphs)
        image_heading = document.add_paragraph()
        image_heading.add_run("批准/生成 QA 图").bold = True
        image_index = len(document.paragraphs)
        image_paragraph = document.add_paragraph()
        image_run = image_paragraph.add_run()
        image_run.add_picture(record["image_path"], width=Inches(5.8))
        document.add_paragraph(f"图像状态：{record['image_status']} · SHA-256：{record['image_sha256']}")
        alignment["shots"].append({
            "shot_id": shot,
            "line_ids": [line["line_id"] for line in record["lines"]],
            "prompt_file": record["prompt_path"],
            "prompt_sha256": record["prompt_sha256"],
            "image_file": record["image_path"],
            "image_sha256": record["image_sha256"],
            "docx_paragraph_indices": {
                "heading": heading_index,
                "script_heading": script_heading_index,
                "lines": line_indices,
                "prompt_heading": prompt_heading_index,
                "prompt": prompt_index,
                "image_heading": image_heading_index,
                "image": image_index,
            },
        })
    # Save only after all images have been embedded successfully.
    document.save(output_docx)
    alignment["docx_sha256"] = sha256(output_docx)
    alignment_path.write_text(json.dumps(alignment, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return output_docx


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Export V3 Canonical build to DOCX")
    parser.add_argument("build_dir", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--alignment-manifest", type=Path, default=None)
    args = parser.parse_args(argv)
    try:
        output = export_docx(args.build_dir.resolve(), args.output_docx, args.alignment_manifest)
    except ExportBlocked as exc:
        print(f"DOCX_EXPORT_BLOCKED: {exc}", file=sys.stderr)
        return 2
    print(json.dumps({"status": "PASS", "docx": str(output)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
