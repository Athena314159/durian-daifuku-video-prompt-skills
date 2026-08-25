#!/usr/bin/env python3
"""Regression tests for lint_prompt_txt.py."""

from __future__ import annotations

import base64
import json
import tempfile
import unittest
from pathlib import Path

from lint_prompt_txt import lint, lint_delivery_directory


BASE_GOOD_PROMPT = """【生成目标与叙事职责】
4:3真实手机生活拍摄。本镜叙事作用是用男生藏着礼物回家、等待镜头后女生反应的短互动建立人物关系，并让礼物成为情绪转折的可见证据。

【口播原文与声源】
女生在镜头后说“今天七夕几点了才回来？”。声音靠近手机，起音直接，语速中快；男生只用画面反应承接这句镜外现场对白。

【原片叙事复原】
男生刚回到门口时带着一点心虚和想卖关子的克制，他手里的盒子说明此刻想把礼物藏到最后一拍再揭示。听到镜头后女生带着追问感的问句后，声音成为可听触发；他的意图是先用眼神试探对方，再让盒子回答“为什么回来晚”。情绪从短暂紧张转为压不住的得意，最终落到想把礼物转向镜头的轻快确认，让观众感受到两个人熟悉、带打趣感的关系，并承接下一拍的礼物展示。推断依据来自他先停手、再抬眼、眉峰短促抬起以及手腕随后转盒的连续动作。

【原片逐时动作】
0.00–1.00秒：女生在镜头后说“今天七夕几点了才回来？”。男生听见后，夹住盒底的手指先停半拍，视线从盒面抬起并看向镜头后女生的眼睛位置；眉峰短促抬起，眼睑随后放松，嘴角压住笑意，肩颈从轻微绷紧转为放松，手腕开始把盒面转向镜头。开口前的短促吸气变成无声停顿，让对方的问句先落地。

【产品与动作物理】
男生双手稳定承托同一只礼盒，手指受力先建立，盒体随后转向；包装纸随腕部运动产生轻微折光和摩擦声，盒子的位置、数量和持有者连续一致。

【摄影、灯光与声音】
场景保留普通住宅暖顶灯、门口脚步、包装摩擦和手机自动对焦。镜头保持生活化中近景，焦点先落在盒子，再随男生抬眼回到眉眼；采用轻微川渝年轻情侣城市口音，声音中低，平翘舌咬字偏松，起音轻，语速稍快，重音落在追问，尾音带一点笑气。

【最小纠错附录】
女生始终位于镜头后；禁止生成自动字幕和水印。
"""


def build_valid_prompt() -> str:
    """Create a 3000–4000-character fixture from positive narrative/action detail."""
    phases = ("听见问句", "停住手指", "抬眼试探", "眉峰抬起", "肩颈放松", "转动礼盒", "焦点回脸", "准备揭示")
    visible_results = ("盒面折光移动", "指腹压力加深", "衣料产生轻微牵动", "嘴角笑意逐渐露出", "目光落点变得明确", "包装摩擦声靠近", "手机焦点平滑回正", "下一拍展示方向建立")
    details = []
    for index in range(22):
        phase = phases[index % len(phases)]
        result = visible_results[(index * 3) % len(visible_results)]
        details.append(
            f"叙事动作补充{index + 1:02d}：在“{phase}”这个节拍里，人物先让上一动作的余势自然收住，"
            f"再由视线、手指受力和呼吸共同推动下一动作；画面可见“{result}”，情绪因此由试探向分享推进，"
            "摄影只跟随动作目的调整焦点，现场脚步、衣料与包装声音共同维持生活化节奏。"
        )
    return BASE_GOOD_PROMPT + "\n" + "\n".join(details)


GOOD_PROMPT = build_valid_prompt()


def make_txt(
    *,
    sid: str = "S001",
    speaker: str = "女生",
    extra_prompt: str = "",
    product: str = "完整未破、手持展示",
    scale_mode: str | None = None,
    projection_mode: str | None = None,
    projection_source: str | None = None,
    source_time: str = "00:00.000–00:03.000",
    source_ids: str = "SRC001",
    duration: float = 3.0,
    description: str = "按原片人物停步抬眼的节奏，把新版口播落在抬眼看向镜头后人物时",
    image_field: str = "SRC001=work/approved_frames/SRC001.png",
) -> str:
    prompt = GOOD_PROMPT + extra_prompt
    char_count = len("".join(prompt.split()))
    return f"""==================================================
{sid}｜“今天七夕几点了才回来？”
==================================================
原片时间：{source_time}
源分镜ID：{source_ids}
独立生成时长：{duration:.3f}秒
分镜描述：{description}
本镜句段数：1
实际可说时段：3.000秒
计划语速：3.33字/秒
人物位置：男生始终在画面中；女生始终在镜头后且不出镜
声音方式：女生镜外现场对白
产品形态：{product}
{f"尺度模式：{scale_mode}" if scale_mode else ""}
{f"投影模式：{projection_mode}" if projection_mode else ""}
{f"投影事实源：{projection_source}" if projection_source else ""}
生成首帧：待制作
分镜图：{image_field}
核心主体：男生
核心动作：抬眼、看向
核心产品：盒子
适用表演层：情绪触发、视线、五官、身体手部、呼吸停顿、声音口语

【口播稿】
{speaker}：“今天七夕几点了才回来？”

【完整Prompt｜主体非空白字符数：{char_count}】
{prompt}
【原片动作对应】
- 原片 00:00.000–00:01.000 的男生停步和抬眼，对应生成镜内 0.00–1.00 秒。

【内容审核记录】
- 人物事实：男生在画面中；女生在镜头后。
- 台词与声源事实：该句由女生在镜外说。
- 生成后像素复核：女生不得出镜。
"""


def make_role_lock(*, required_actions: list[dict[str, str]] | None = None) -> dict:
    return {
        "characters": {
            "A": {"label": "女生", "visibility": "offscreen_all", "camera_holder": True},
            "B": {"label": "男生", "visibility": "onscreen_all", "camera_holder": False},
        },
        "spatial_invariants": ["女生全程镜头后", "男生全程画面中"],
        "speech_plan": {
            "source": "creative_proposal",
            "disclosed_to_user": True,
            "summary": "轻微川渝年轻情侣城市口音；平翘舌不刻意咬正，语速中快，尾音带笑，不夸张模仿方言"
        },
        "dialogue": [
            {
                "speaker_id": "A",
                "label": "女生",
                "text": "今天七夕几点了才回来？",
                "delivery": "offscreen_live",
                "visible_lip_sync": False,
                "action_subject": "男生",
                "visual_evidence": "男生抬眼看向镜头后的女生",
            }
        ],
        "required_actions": required_actions or [],
    }


def make_short_merge_story(*, include_eating: bool = False) -> tuple[str, dict]:
    description = "按原片 SRC001 的停步抬眼与 SRC002 的拿稳盒子连续节奏，把新版口播落在抬眼之后"
    eating_prompt = ""
    if include_eating:
        eating_prompt = (
            "\n1.00–1.80秒：男生把食物接近口部，牙齿接触后完成咬合，产品离嘴，随后闭口咀嚼。"
            "\n1.80–2.40秒：闭口咀嚼结束后口型恢复，下一拍可以马上说新版口播，不补吞咽或吃后反应。"
        )
    txt = make_txt(
        duration=4.2,
        source_time="00:00.000–00:04.200",
        source_ids="SRC001、SRC002",
        description=description,
        image_field="SRC001=work/approved/SRC001.png；SRC002=work/approved/SRC002.png",
        extra_prompt=eating_prompt,
    )
    story: dict = {
        "source_duration_seconds": 4.2,
        "generation_time_policy": {
            "min_duration_seconds": 4.0,
            "max_duration_seconds": 12.0,
            "onscreen_speech_max_chars_per_second": 5.0,
            "voiceover_max_chars_per_second": 5.5,
        },
        "source_shot_inventory": [
            {
                "source_shot_id": "SRC001",
                "source_start": "00:00.000",
                "source_end": "00:01.800",
                "duration_seconds": 1.8,
                "action": "男生停步抬眼",
                "source_first_frame": "work/source/SRC001.png",
                "approved_delivery_image": "work/approved/SRC001.png",
            },
            {
                "source_shot_id": "SRC002",
                "source_start": "00:01.800",
                "source_end": "00:04.200",
                "duration_seconds": 2.4,
                "action": "男生拿稳盒子",
                "source_first_frame": "work/source/SRC002.png",
                "approved_delivery_image": "work/approved/SRC002.png",
            },
        ],
        "generation_shot_map": [
            {
                "shot_id": "S001",
                "origin": "source_merge",
                "source_shot_ids": ["SRC001", "SRC002"],
                "generation_duration_seconds": 4.2,
                "shot_description": description,
                "revised_script_anchor": "新版口播第1句",
                "merge_reason": "相邻源分镜均不足4秒且动作连续",
            }
        ],
    }
    if include_eating:
        story["eating_plan"] = {
            "minimum_nonconsecutive_events_for_30s": 3,
            "events": [
                {
                    "event_id": "EAT-S01",
                    "origin": "source",
                    "generation_shot_id": "S001",
                    "source_shot_ids": ["SRC001", "SRC002"],
                    "narrative_section": "原片吃食",
                    "rhythm_anchor": "原片1.00秒动作点",
                    "revised_script_anchor": "新版口播第1句之前",
                    "required_phases": ["approach", "bite_contact", "withdraw", "closed_chew"],
                    "speech_resume_after": "closed_chew_end",
                }
            ],
        }
    return txt, story


def make_30s_three_eating_fixture() -> tuple[str, dict]:
    blocks: list[str] = []
    inventory: list[dict] = []
    mappings: list[dict] = []
    events: list[dict] = []
    eating_indexes = {1: "开场口感", 3: "中段材质", 5: "收束复证"}
    for index in range(1, 7):
        sid = f"S{index:03d}"
        source_id = f"SRC{index:03d}"
        start = (index - 1) * 5
        end = index * 5
        source_time = f"00:{start:02d}.000–00:{end:02d}.000"
        description = f"按原片 {source_id} 第{index}段动作节奏，把新版口播第{index}句落在动作稳定点"
        eating_prompt = ""
        if index in eating_indexes:
            eating_prompt = (
                "\n1.00–2.20秒：男生把食物接近口部，牙齿接触并咬合，产品离嘴后闭口咀嚼。"
            )
        blocks.append(
            make_txt(
                sid=sid,
                source_time=source_time,
                source_ids=source_id,
                duration=5.0,
                description=description,
                image_field=f"{source_id}=work/approved/{source_id}.png",
                extra_prompt=eating_prompt,
            )
        )
        inventory.append(
            {
                "source_shot_id": source_id,
                "source_start": f"00:{start:02d}.000",
                "source_end": f"00:{end:02d}.000",
                "duration_seconds": 5.0,
                "action": f"原片第{index}段动作",
                "source_first_frame": f"work/source/{source_id}.png",
                "approved_delivery_image": f"work/approved/{source_id}.png",
            }
        )
        mappings.append(
            {
                "shot_id": sid,
                "origin": "source",
                "source_shot_ids": [source_id],
                "generation_duration_seconds": 5.0,
                "shot_description": description,
                "revised_script_anchor": f"新版口播第{index}句",
            }
        )
        if index in eating_indexes:
            events.append(
                {
                    "event_id": f"EAT-S{index:02d}",
                    "origin": "source",
                    "generation_shot_id": sid,
                    "source_shot_ids": [source_id],
                    "narrative_section": eating_indexes[index],
                    "rhythm_anchor": f"原片{start + 1:.3f}秒吃食动作点",
                    "revised_script_anchor": f"新版口播第{index}句之前",
                    "required_phases": ["approach", "bite_contact", "withdraw", "closed_chew"],
                    "speech_resume_after": "closed_chew_end",
                }
            )
    story = {
        "source_duration_seconds": 30.0,
        "generation_time_policy": {
            "min_duration_seconds": 4.0,
            "max_duration_seconds": 12.0,
            "onscreen_speech_max_chars_per_second": 5.0,
            "voiceover_max_chars_per_second": 5.5,
        },
        "source_shot_inventory": inventory,
        "generation_shot_map": mappings,
        "eating_plan": {
            "minimum_nonconsecutive_events_for_30s": 3,
            "events": events,
        },
    }
    return "\n".join(blocks), story


class PromptLintRegressionTests(unittest.TestCase):
    def run_lint(
        self,
        txt: str,
        role_lock: dict | None = None,
        story_plan: dict | None = None,
        *,
        enforce_prompt_length: bool = True,
        stage: str | None = None,
    ):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            txt_path = root / "prompt.txt"
            txt_path.write_text(txt, encoding="utf-8")
            role_path = None
            if role_lock is not None:
                role_path = root / "role_lock.json"
                role_path.write_text(json.dumps(role_lock, ensure_ascii=False), encoding="utf-8")
            story_path = None
            if story_plan is not None:
                story_path = root / "story_plan.json"
                story_path.write_text(json.dumps(story_plan, ensure_ascii=False), encoding="utf-8")
            return lint(
                txt_path,
                role_path,
                min_prompt_chars=3000 if enforce_prompt_length else None,
                max_prompt_chars=4000 if enforce_prompt_length else None,
                story_plan_path=story_path,
                stage=stage,
                enforce_prompt_length=True if enforce_prompt_length else None,
            )

    def assert_has_code(self, issues, code: str):
        self.assertIn(code, {issue.code for issue in issues}, [issue.message for issue in issues])

    def test_correct_onscreen_man_offscreen_woman(self):
        self.assertEqual([], self.run_lint(make_txt(), make_role_lock()))

    def test_prompt_below_3000_is_blocked(self):
        text = make_txt().replace(GOOD_PROMPT, BASE_GOOD_PROMPT)
        issues = self.run_lint(text, make_role_lock())
        self.assert_has_code(issues, "PROMPT_TOO_SHORT")

    def test_prompt_length_is_not_enforced_by_default(self):
        text = make_txt().replace(GOOD_PROMPT, BASE_GOOD_PROMPT)
        issues = self.run_lint(text, make_role_lock(), enforce_prompt_length=False)
        self.assertNotIn("PROMPT_TOO_SHORT", {issue.code for issue in issues})
        self.assertNotIn("PROMPT_TOO_LONG", {issue.code for issue in issues})

    def test_story_plan_can_explicitly_enable_both_prompt_limits(self):
        text, story = make_short_merge_story()
        text = text.replace(GOOD_PROMPT, BASE_GOOD_PROMPT)
        story["prompt_length_contract"] = {
            "enabled": True,
            "minimum_non_whitespace_characters": 3000,
            "maximum_non_whitespace_characters": 4000,
        }
        issues = self.run_lint(text, story_plan=story, enforce_prompt_length=False)
        self.assert_has_code(issues, "PROMPT_TOO_SHORT")

    def test_min_or_max_without_explicit_length_contract_is_blocked(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            txt_path = Path(temp_dir) / "prompt.txt"
            txt_path.write_text(make_txt(), encoding="utf-8")
            issues = lint(txt_path, min_prompt_chars=3000)
        self.assert_has_code(issues, "PROMPT_LENGTH_CONTRACT_INVALID")

    def test_declared_prompt_count_must_match_computed_count(self):
        text = make_txt().replace(
            f"主体非空白字符数：{len(''.join(GOOD_PROMPT.split()))}",
            "主体非空白字符数：3000",
        )
        issues = self.run_lint(text, make_role_lock())
        self.assert_has_code(issues, "PROMPT_CHAR_COUNT_MISMATCH")

    def test_wrong_speaker_is_blocked(self):
        issues = self.run_lint(make_txt(speaker="男生"), make_role_lock())
        self.assert_has_code(issues, "DIALOGUE_SPEAKER_MISMATCH")

    def test_offscreen_character_visibility_leak_is_blocked(self):
        issues = self.run_lint(make_txt(extra_prompt="女生的侧脸进入画面并清晰可见。"), make_role_lock())
        self.assert_has_code(issues, "OFFSCREEN_CHARACTER_VISIBLE")

    def test_internal_product_label_and_status_word_are_blocked(self):
        issues = self.run_lint(make_txt(extra_prompt="本镜审核写PASS。", product="V2 person_eating"), make_role_lock())
        self.assert_has_code(issues, "PRODUCT_STATE_LABEL_NOT_CHINESE")
        self.assert_has_code(issues, "STRUCTURE_RESULT_MISREPRESENTED_AS_CONTENT_AUDIT")

    def test_missing_bite_chain_is_blocked_when_source_requires_it(self):
        role = make_role_lock(required_actions=[{"shot_id": "S001", "kind": "bite_chain"}])
        issues = self.run_lint(make_txt(), role)
        self.assert_has_code(issues, "SOURCE_ACTION_OMITTED")

    def test_abstract_emotion_without_six_layers_is_blocked(self):
        text = make_txt().replace(GOOD_PROMPT, "0.00–1.00秒：男生自然开心地展示礼物。禁止出现任何字幕和水印。")
        issues = self.run_lint(text, make_role_lock())
        self.assert_has_code(issues, "PERFORMANCE_DETAIL_MISSING")

    def test_generic_speech_placeholder_is_blocked(self):
        text = make_txt().replace(
            "采用轻微川渝年轻情侣城市口音，声音中低，平翘舌咬字偏松，起音轻，语速稍快，重音落在追问，尾音带一点笑气。",
            "声音中低，起音轻，语速稍快，重音落在礼物名称，尾音带笑，沿用原片生活口语节奏。",
        )
        issues = self.run_lint(text, make_role_lock())
        self.assert_has_code(issues, "GENERIC_SPEECH_PLACEHOLDER")
        self.assert_has_code(issues, "ACCENT_PLAN_MISSING")

    def test_undisclosed_creative_accent_proposal_is_blocked(self):
        role = make_role_lock()
        role["speech_plan"]["disclosed_to_user"] = False
        issues = self.run_lint(make_txt(), role)
        self.assert_has_code(issues, "ACCENT_PROPOSAL_NOT_DISCLOSED")

    def test_butter_crisp_bare_product_requires_physical_microstructure(self):
        issues = self.run_lint(make_txt(product="黄油脆丝棒完整未破、手持展示"), make_role_lock())
        self.assert_has_code(issues, "PRODUCT_MICROSTRUCTURE_RULE_MISSING")

    def test_butter_crisp_box_requires_confirmed_dimensions_and_ratios(self):
        issues = self.run_lint(
            make_txt(product="黄油脆丝棒未开封零售外盒", scale_mode="physical_consistency"), make_role_lock()
        )
        self.assert_has_code(issues, "PACKAGE_DIMENSION_LOCK_MISSING")

    def test_butter_crisp_requires_a_unique_scale_mode_before_compilation(self):
        issues = self.run_lint(make_txt(product="黄油脆丝棒完整未破、手持展示"), make_role_lock())
        self.assert_has_code(issues, "SCALE_MODE_UNDECLARED")

    def test_butter_crisp_physical_mode_requires_visible_product_scale_lock(self):
        issues = self.run_lint(
            make_txt(product="黄油脆丝棒完整未破、手持展示", scale_mode="physical_consistency"), make_role_lock()
        )
        self.assert_has_code(issues, "PRODUCT_PACKAGE_SCALE_LOCK_MISSING")

    def test_bare_product_does_not_require_an_invisible_box_ratio(self):
        extra = (
            "黄油脆丝棒六面是实体片状脆丝覆盖层，每块有自身厚度和侧边厚度，形成前后遮挡、"
            "重叠缝隙与不规则窄缝；少量碎片轮廓凸出并打破干净外轮廓。禁止平面贴图、"
            "印刷图案、浅浮雕、压花或光滑橙色基底。单根约12 × 2.5 × 1 cm，正面目标4.8:1，"
            "成品视觉长宽比4:1–5:1，侧面厚宽比约0.40。"
        )
        issues = self.run_lint(
            make_txt(
                product="黄油脆丝棒完整未破、手持展示",
                scale_mode="physical_consistency",
                extra_prompt=extra,
            ),
            make_role_lock(),
        )
        self.assertNotIn("PRODUCT_PACKAGE_SCALE_LOCK_MISSING", {issue.code for issue in issues})

    def test_relative_pixel_resize_rejects_physical_projection_ratio_in_the_same_prompt(self):
        issues = self.run_lint(
            make_txt(
                product="黄油脆丝棒完整未破、手持展示",
                scale_mode="relative_pixel_resize",
                extra_prompt="本轮以唯一原始批准帧为尺寸事实源，产品与盒面投影比为0.80。",
            ),
            make_role_lock(),
        )
        self.assert_has_code(issues, "SCALE_MODE_COLLISION")

    def test_butter_crisp_structure_and_box_dimensions_can_pass(self):
        extra = (
            "黄油脆丝棒六面是实体片状脆丝覆盖层，每块有自身厚度和侧边厚度，形成前后遮挡、"
            "重叠缝隙与不规则窄缝；少量碎片轮廓凸出并打破干净外轮廓。禁止平面贴图、"
            "印刷图案、浅浮雕、压花或光滑橙色基底。外盒严格15×15×4.5 cm，正面1:1正方形，"
            "厚度约为正面边长30%，是一只扁方盒。单根约12 × 2.5 × 1 cm，正面目标4.8:1，"
            "成品视觉长宽比4:1–5:1，侧面厚宽比约0.40；同平面、"
            "可比朝向时单根与15 cm盒面边长比例为12:15=0.80，跨景深时先做透视校正。"
        )
        issues = self.run_lint(
            make_txt(
                product="黄油脆丝棒完整未破手持与未开封零售外盒",
                extra_prompt=extra,
                scale_mode="physical_consistency",
            ),
            make_role_lock(),
        )
        blocked = {issue.code for issue in issues}
        self.assertNotIn("PRODUCT_MICROSTRUCTURE_RULE_MISSING", blocked)
        self.assertNotIn("PACKAGE_DIMENSION_LOCK_MISSING", blocked)
        self.assertNotIn("SCALE_MODE_UNDECLARED", blocked)
        self.assertNotIn("PRODUCT_PACKAGE_SCALE_LOCK_MISSING", blocked)
        self.assertNotIn("SCALE_MODE_COLLISION", blocked)

    def test_butter_crisp_break_requires_state_specific_geometry_and_hand_scale(self):
        issues = self.run_lint(
            make_txt(product="黄油脆丝棒双手掰断、两段断面展示", scale_mode="physical_consistency"),
            make_role_lock(),
        )
        blocked = {issue.code for issue in issues}
        expected = {
            "PROJECTION_MODE_UNDECLARED",
            "PROJECTION_SOURCE_UNDECLARED",
            "BREAK_SOURCE_SINGLE_INSTANCE_MISSING",
            "BREAK_FRACTION_MISSING",
            "BROKEN_LENGTH_CONSERVATION_MISSING",
            "BROKEN_WIDTH_THICKNESS_LOCK_MISSING",
            "BROKEN_SEGMENT_ASPECT_RULE_MISSING",
            "HAND_PRODUCT_SCALE_ANCHOR_MISSING",
            "BREAK_CAMERA_DEPTH_LOCK_MISSING",
            "FRACTURE_SURFACE_LOCK_MISSING",
            "FRACTURE_VIEW_MODE_MISSING",
        }
        self.assertTrue(expected.issubset(blocked), blocked)

    def test_butter_crisp_mid_break_state_locks_can_pass(self):
        extra = (
            "黄油脆丝棒六面是实体片状脆丝覆盖层，每片有自身厚度，形成前后遮挡、重叠缝隙；"
            "少量碎片轮廓凸出并打破干净外轮廓。禁止平面贴图、印刷图案、浅浮雕、压花或"
            "光滑橙色基底。掰断前是同一根约12 × 2.5 × 1 cm完整脆丝棒，正面目标4.8:1，"
            "完整体成品视觉长宽比4:1–5:1，侧面厚宽比约0.40；同平面、可比朝向时与15 cm"
            "盒面边长比例为12:15=0.80，跨景深先透视校正。断裂位置break_fraction为0.50，"
            "两段长度之和约等于原来的12 cm，左右段宽度都继续约2.5 cm、厚度继续约1 cm，"
            "每段正面长宽比2.2:1–2.6:1。严格继承原始首帧的双手和断前投影尺寸，产品与"
            "拇指指腹同一焦平面，手到镜头距离不变，断后不向镜头额外推进。两个断面来自"
            "同一断裂点，断口不规则、疏松多孔。fracture_view_mode为"
            "long_face_with_end_visible，约6 × 2.5 cm长外表面朝镜头，断面只在上端小端面。"
            "执行source_pixel_lock：脸框0.97–1.03，手掌框0.95–1.05，手脸比例±5%，"
            "腕部中心位移不超过画幅2%；只允许产品与接触指尖局部变化，肩线、裁切和相机不变。"
        )
        issues = self.run_lint(
            make_txt(
                product="黄油脆丝棒双手掰断、两段断面展示",
                extra_prompt=extra,
                scale_mode="physical_consistency",
                projection_mode="source_pixel_lock",
                projection_source="原视频 source/original.mp4 00:25.200 真实帧",
            ),
            make_role_lock(),
        )
        blocked = {issue.code for issue in issues}
        break_codes = {
            "BREAK_SOURCE_SINGLE_INSTANCE_MISSING",
            "BREAK_FRACTION_MISSING",
            "BROKEN_LENGTH_CONSERVATION_MISSING",
            "BROKEN_WIDTH_THICKNESS_LOCK_MISSING",
            "BROKEN_SEGMENT_ASPECT_RULE_MISSING",
            "HAND_PRODUCT_SCALE_ANCHOR_MISSING",
            "BREAK_CAMERA_DEPTH_LOCK_MISSING",
            "FRACTURE_SURFACE_LOCK_MISSING",
            "FRACTURE_VIEW_MODE_MISSING",
        }
        self.assertTrue(break_codes.isdisjoint(blocked), blocked)

    def test_butter_crisp_break_rejects_failed_generated_projection_source(self):
        issues = self.run_lint(
            make_txt(
                product="黄油脆丝棒双手掰断、两段断面展示",
                scale_mode="physical_consistency",
                projection_mode="source_pixel_lock",
                projection_source="脆丝棒双手掰断_修复测试候选_未批准.png",
            ),
            make_role_lock(),
        )
        self.assert_has_code(issues, "FAILED_FRAME_USED_AS_PROJECTION_SOURCE")

    def test_source_pixel_lock_rejects_camera_axis_push_language(self):
        issues = self.run_lint(
            make_txt(
                product="黄油脆丝棒双手掰断、两段断面展示",
                scale_mode="physical_consistency",
                projection_mode="source_pixel_lock",
                projection_source="原视频 source/original.mp4 00:25.200 真实帧",
                extra_prompt="双手把两段交替前推并紧贴镜头。",
            ),
            make_role_lock(),
        )
        self.assert_has_code(issues, "UNAUTHORIZED_CAMERA_AXIS_ADVANCE")

    def test_short_adjacent_source_shots_can_merge_without_deletion(self):
        txt, story = make_short_merge_story()
        self.assertEqual([], self.run_lint(txt, story_plan=story))

    def test_short_source_shot_cannot_remain_unmerged(self):
        txt, story = make_short_merge_story()
        story["source_shot_inventory"] = story["source_shot_inventory"][:1]
        story["generation_shot_map"][0]["source_shot_ids"] = ["SRC001"]
        story["generation_shot_map"][0]["generation_duration_seconds"] = 3.0
        issues = self.run_lint(txt, story_plan=story)
        self.assert_has_code(issues, "SHORT_SOURCE_SHOT_NOT_MERGED")

    def test_every_source_shot_must_remain_mapped(self):
        txt, story = make_short_merge_story()
        story["generation_shot_map"][0]["source_shot_ids"] = ["SRC001"]
        issues = self.run_lint(txt, story_plan=story)
        self.assert_has_code(issues, "SOURCE_SHOT_COVERAGE_INCOMPLETE")

    def test_source_shot_merge_must_keep_inventory_order(self):
        txt, story = make_short_merge_story()
        story["source_shot_inventory"].insert(
            1,
            {
                "source_shot_id": "SRC009",
                "source_start": "00:01.800",
                "source_end": "00:02.000",
                "duration_seconds": 0.2,
                "action": "中间不可跳过动作",
                "source_first_frame": "work/source/SRC009.png",
                "approved_delivery_image": "work/approved/SRC009.png",
            },
        )
        issues = self.run_lint(txt, story_plan=story)
        self.assert_has_code(issues, "NONADJACENT_SOURCE_SHOT_MERGE")

    def test_inserted_generation_shot_requires_unique_add_id(self):
        txt, story = make_short_merge_story()
        description = "根据原片停顿节奏与新版口播口感转折新增一次非连续吃食证明"
        txt += "\n" + make_txt(
            sid="S002",
            source_time="新增镜头",
            source_ids="新增镜头",
            duration=4.0,
            description=description,
            image_field="ADD001=work/approved/ADD001.png",
        )
        story["generation_shot_map"].append(
            {
                "shot_id": "S002",
                "origin": "inserted_eating",
                "source_shot_ids": [],
                "generation_duration_seconds": 4.0,
                "shot_description": description,
                "revised_script_anchor": "新版口播口感句之后",
                "insertion_rationale": "原片吃食不足，只补缺少的一次",
                "rhythm_anchor": "语义转折后的切点",
                "source_reference_shot_ids": ["SRC001"],
                "source_reference_frame": "work/source/SRC001.png",
            }
        )
        issues = self.run_lint(txt, story_plan=story)
        self.assert_has_code(issues, "INSERTED_SHOT_ID_INVALID")

    def test_valid_inserted_add_has_no_fabricated_source_timecode(self):
        txt, story = make_short_merge_story()
        description = "根据原片 SRC002 的停顿节奏与新版口播口感转折新增一个独立分镜"
        txt += "\n" + make_txt(
            sid="S002",
            source_time="新增镜头",
            source_ids="新增镜头",
            duration=4.0,
            description=description,
            image_field="ADD001=work/approved/ADD001.png",
        )
        story["generation_shot_map"].append(
            {
                "shot_id": "S002",
                "origin": "inserted_break",
                "inserted_shot_id": "ADD001",
                "source_shot_ids": [],
                "generation_duration_seconds": 4.0,
                "shot_description": description,
                "revised_script_anchor": "新版口播口感句之后",
                "insertion_rationale": "增加一个独立产品证据镜头",
                "rhythm_anchor": "口感句后半拍停顿",
                "source_reference_shot_ids": ["SRC002"],
                "source_reference_frame": "work/source/SRC002.png",
            }
        )
        issues = self.run_lint(txt, story_plan=story)
        self.assertNotIn("SOURCE_TIMECODE_MAP_MISMATCH", {issue.code for issue in issues})
        self.assertNotIn("INSERTED_SHOT_ID_INVALID", {issue.code for issue in issues})

    def test_eating_can_finish_and_speech_resume_without_forced_swallow(self):
        txt, story = make_short_merge_story(include_eating=True)
        issues = self.run_lint(txt, story_plan=story)
        self.assertNotIn("EATING_PHASE_EVIDENCE_MISSING", {issue.code for issue in issues})
        self.assertNotIn("EATING_TEMPLATE_PHASE_FORCED", {issue.code for issue in issues})
        self.assertNotIn("CHEWING_SPEECH_CONFLICT", {issue.code for issue in issues})

    def test_unobserved_swallow_cannot_be_forced_into_source_eating_timeline(self):
        txt, story = make_short_merge_story(include_eating=True)
        txt = txt.replace(
            "1.80–2.40秒：闭口咀嚼结束后口型恢复",
            "1.80–2.40秒：人物完成吞咽后口型恢复",
        )
        issues = self.run_lint(txt, story_plan=story)
        self.assert_has_code(issues, "EATING_TEMPLATE_PHASE_FORCED")

    def test_bite_or_closed_chew_cannot_overlap_full_onscreen_speech(self):
        issues = self.run_lint(
            make_txt(extra_prompt="\n1.00–2.00秒：男生在闭口咀嚼期间开口说出完整口播。"),
            make_role_lock(),
        )
        self.assert_has_code(issues, "CHEWING_SPEECH_CONFLICT")

    def test_30s_video_with_three_distributed_source_eating_events_needs_no_insert(self):
        txt, story = make_30s_three_eating_fixture()
        self.assertEqual([], self.run_lint(txt, story_plan=story))

    def test_30s_video_below_three_eating_events_is_blocked(self):
        txt, story = make_30s_three_eating_fixture()
        story["eating_plan"]["events"] = story["eating_plan"]["events"][:2]
        issues = self.run_lint(txt, story_plan=story)
        self.assert_has_code(issues, "EATING_SHOT_QUOTA_MISSING")

    def test_30s_video_with_three_source_events_must_not_add_a_fourth(self):
        txt, story = make_30s_three_eating_fixture()
        story["eating_plan"]["events"].append(
            {
                "event_id": "EAT-I04",
                "origin": "inserted",
                "generation_shot_id": "S006",
                "source_shot_ids": [],
                "narrative_section": "多余补吃",
                "rhythm_anchor": "无必要补吃点",
                "revised_script_anchor": "新版口播末句",
                "required_phases": ["approach"],
                "speech_resume_after": "closed_chew_end",
            }
        )
        issues = self.run_lint(txt, story_plan=story)
        self.assert_has_code(issues, "UNNECESSARY_EATING_SHOT_INSERTION")

    def test_eating_events_must_not_use_adjacent_generation_shots(self):
        txt, story = make_30s_three_eating_fixture()
        story["eating_plan"]["events"][1]["generation_shot_id"] = "S002"
        issues = self.run_lint(txt, story_plan=story)
        self.assert_has_code(issues, "EATING_EVENTS_NOT_DISTRIBUTED")

    def test_prompt_core_fact_cannot_disappear_during_compression(self):
        text = make_txt().replace("核心动作：抬眼、看向", "核心动作：倒立旋转")
        issues = self.run_lint(text, make_role_lock())
        self.assert_has_code(issues, "PROMPT_CORE_FACT_OMITTED")

    def test_repeated_long_sentence_padding_is_blocked(self):
        repeated = "逐帧保证人物、产品、动作和镜头关系稳定，不得突然变形、跳帧、悬空或改变身份。"
        issues = self.run_lint(make_txt(extra_prompt=f"\n{repeated}\n{repeated}"), make_role_lock())
        self.assert_has_code(issues, "PROMPT_PADDING_DETECTED")

    def test_negative_constraint_padding_is_blocked(self):
        restrictions = "\n".join(
            f"禁止错误形态{index:02d}进入画面，不得改变本轮锁定对象。"
            for index in range(80)
        )
        issues = self.run_lint(make_txt(extra_prompt="\n" + restrictions), make_role_lock())
        self.assert_has_code(issues, "NEGATIVE_CONSTRAINT_OVERLOAD")

    def test_narrative_six_layer_format_is_required(self):
        text = make_txt().replace("【原片叙事复原】", "【零散六层描述】")
        issues = self.run_lint(text, make_role_lock())
        self.assert_has_code(issues, "NARRATIVE_FORMAT_MISSING")

    def test_emotional_causality_is_required_for_person_shot(self):
        prompt = GOOD_PROMPT
        start = prompt.index("【原片叙事复原】") + len("【原片叙事复原】")
        end = prompt.index("【原片逐时动作】")
        flat = prompt[:start] + "\n人物表情自然开心，动作真实。\n" + prompt[end:]
        text = make_txt().replace(GOOD_PROMPT, flat)
        issues = self.run_lint(text, make_role_lock())
        self.assert_has_code(issues, "EMOTIONAL_CAUSALITY_MISSING")

    def test_six_layer_audit_fields_cannot_leak_into_prompt(self):
        issues = self.run_lint(make_txt(extra_prompt="\n情绪层status=observed，confidence=0.92。"), make_role_lock())
        self.assert_has_code(issues, "SIX_LAYER_AUDIT_LEAKED_INTO_PROMPT")

    def test_user_delivery_directory_must_contain_only_one_docx(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            (root / "最终交付.docx").write_bytes(b"docx fixture")
            self.assert_has_code(lint_delivery_directory(root), "FINAL_DOCX_INVALID")
            (root / "对齐表.json").write_text("{}", encoding="utf-8")
            self.assert_has_code(lint_delivery_directory(root), "USER_DELIVERY_ARTIFACT_LEAK")

    def test_valid_editable_docx_with_real_image_relationship_passes(self):
        from docx import Document
        from docx.shared import Inches

        one_pixel_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            image_path = root / "frame.png"
            image_path.write_bytes(one_pixel_png)
            document = Document()
            document.add_paragraph("准确秒数：原片0.000–4.000秒｜生成镜内0.000–4.000秒")
            document.add_paragraph("分镜描述：按原片节奏与新版口播重排的可编辑描述")
            document.add_paragraph("口播稿：这是可编辑的新版口播正文")
            document.add_paragraph("即梦可复制 Prompt：从第一帧开始按镜内秒数执行动作。")
            document.add_picture(str(image_path), width=Inches(0.4))
            docx_path = root / "最终交付.docx"
            document.save(docx_path)
            image_path.unlink()
            self.assertEqual([], lint_delivery_directory(root))

    def test_real_docx_without_editable_labels_or_image_is_blocked(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            document = Document()
            document.add_paragraph("只有一句普通文字。")
            document.save(root / "最终交付.docx")
            issues = lint_delivery_directory(root)
            self.assert_has_code(issues, "FINAL_DOCX_IMAGE_MISSING")
            self.assert_has_code(issues, "FINAL_DOCX_EDITABLE_TEXT_MISSING")

    def test_lint_stages_require_their_declared_inputs(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            txt_path = Path(temp_dir) / "prompt.txt"
            txt_path.write_text(make_txt(), encoding="utf-8")
            text_branch = lint(txt_path, stage="text_branch")
            postexport = lint(txt_path, stage="full_delivery_postexport")
        self.assertTrue(sum(issue.code == "LINT_STAGE_INPUT_MISSING" for issue in text_branch) >= 2)
        self.assertTrue(sum(issue.code == "LINT_STAGE_INPUT_MISSING" for issue in postexport) >= 2)

    def test_user_delivery_directory_requires_a_docx(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            self.assert_has_code(lint_delivery_directory(Path(temp_dir)), "FINAL_DOCX_MISSING")


if __name__ == "__main__":
    unittest.main(verbosity=2)
