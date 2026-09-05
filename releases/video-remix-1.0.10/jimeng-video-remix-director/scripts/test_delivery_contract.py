#!/usr/bin/env python3
"""Regression tests for source coverage, eating rhythm, break proof and package art."""

from __future__ import annotations

import hashlib
import copy
import json
import re
import struct
import subprocess
import sys
import tempfile
import zipfile
import zlib
from pathlib import Path

from docx import Document

from pipeline import (
    canonical_input_hashes,
    normalized_prompt_length_contract,
    validate_break_plan,
    validate_eating_plan,
    validate_package_artwork,
    validate_revised_script_coverage,
    validate_source_shot_contract,
)
from align_exports import parse_docx_body
from export_jimeng_docx import (
    SNAPSHOT_PAYLOAD_FIELDS,
    delivery_asset_responsibility,
    digest,
    digest_text,
    exact_time_label,
    format_break_occurrence,
    format_package_face,
    format_performance_layers,
    target_frame_caption,
    validate_compile_snapshot,
    validate_unit_delivery_asset_binding,
)


LAYER_KEYS = (
    "emotion_trigger",
    "gaze",
    "facial_microreaction",
    "body_hand_preparation",
    "breath_pause",
    "voice_speech",
)


def codes(issues: list[dict]) -> set[str]:
    return {str(item.get("code")) for item in issues}


def six_layers() -> dict:
    return {
        key: {
            "status": "not_applicable",
            "source_timecode": None,
            "source_reference_frame": None,
            "observable_evidence": "本测试分镜不包含需要该层证明的人物表演。",
            "confidence": 1.0,
            "gap_reason": None,
        }
        for key in LAYER_KEYS
    }


def write_json(path: Path, value: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_test_png(path: Path, rgb: tuple[int, int, int]) -> None:
    """Write a dependency-free RGB PNG with stable, distinct bytes."""
    width, height = 8, 8
    scanline = b"\x00" + bytes(rgb) * width
    raw = scanline * height

    def chunk(kind: bytes, payload: bytes) -> bytes:
        return struct.pack(">I", len(payload)) + kind + payload + struct.pack(">I", zlib.crc32(kind + payload) & 0xFFFFFFFF)

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def build_delivery_proof_fixture(root: Path) -> dict:
    """Build a minimal immutable compile + DOCX proof package for adversarial audit tests."""
    project_dir = root / "delivery-proof-project"
    compile_id = "20260822T000000-abc123"
    prompt_contract = {
        "enabled": False,
        "minimum_non_whitespace_characters": 0,
        "maximum_non_whitespace_characters": 0,
    }
    source_image_one = project_dir / "assets" / "SRC001-lift.png"
    source_image_two = project_dir / "assets" / "SRC001-hold.png"
    inserted_image = project_dir / "assets" / "ADD001.png"
    product_reference = project_dir / "assets" / "product-reference.png"
    write_test_png(source_image_one, (220, 150, 90))
    write_test_png(source_image_two, (175, 105, 45))
    write_test_png(inserted_image, (90, 170, 220))
    write_test_png(product_reference, (240, 205, 90))

    source_layers = six_layers()
    inserted_layers = six_layers()
    source_unit_value = {
        "source_shot_id": "SRC001",
        "source_timecode": {"start": 0.0, "end": 2.0, "duration": 2.0},
        "generation_timecode": {"start": 0.0, "end": 2.0, "duration": 2.0},
        "storyboard_description": "人物先在桌前举起原片产品，镜头保持稳定。",
        "script_text": "先看原片产品",
        "source_first_frame": "assets/SRC001-lift.png",
        "delivery_asset_ids": ["ASSET-SRC001-LIFT", "ASSET-SRC001-HOLD"],
        "delivery_asset_roles": {
            "ASSET-SRC001-LIFT": "动作关键状态1：人物将产品举入镜头中心。",
            "ASSET-SRC001-HOLD": "动作关键状态2：人物稳定手持并让产品正面可见。",
        },
        "source_performance_layers": source_layers,
    }
    inserted_unit_value = {
        "inserted_shot_id": "ADD001",
        "generation_timecode": {"start": 2.0, "end": 4.0, "duration": 2.0},
        "storyboard_description": "按新版口播节奏补入一镜独立产品近景。",
        "script_text": "再看新增近景",
        "delivery_asset_ids": ["ASSET-ADD001"],
        "insertion_rationale": "新版口播需要一个独立近景证据。",
        "rhythm_anchor": "承接上一句结束后的自然切点。",
        "source_reference_shot_ids": ["SRC001"],
        "source_reference_frame": "assets/SRC001.png",
        "source_performance_layers": inserted_layers,
    }
    shot = {
        "id": "S001",
        "title": "原片展示与新增近景",
        "timecode": {"start": 0.0, "end": 4.0, "duration": 4.0},
        "visual_type": "person_product_showcase",
        "narrative_role": "visual_proof",
        "script_segment_ids": ["T001"],
        "risk": {"level": "low", "reasons": []},
        "source_units": [source_unit_value],
        "inserted_units": [inserted_unit_value],
        "audio": {"delivery_mode": "voiceover", "script_text": "先看原片产品再看新增近景"},
        "asset_links": {
            "source_first_frame": "assets/SRC001-lift.png",
            "selected_beauty_keyframe": "assets/SRC001-hold.png",
            "approved_generation_first_frame": "assets/ADD001.png",
            "product_references": ["assets/product-reference.png"],
            "avatar_reference": None,
        },
        "product_state": {"state": "whole", "package_artwork": {"visible_faces": []}},
    }
    reuse_plan = {
        "inventory": [
            {
                "asset_id": "ASSET-SRC001-LIFT",
                "path": "assets/SRC001-lift.png",
                "sha256": digest(source_image_one),
                "approval_status": "user_approved",
                "user_approval": {"status": "user_approved", "display_receipt_id": "gallery-proof-001", "approved_at": "2026-08-24T12:01:00+08:00", "asset_sha256": digest(source_image_one)},
                "source_shot_ids": ["SRC001"],
                "responsibility": "动作关键状态1：人物将产品举入镜头中心。",
            },
            {
                "asset_id": "ASSET-SRC001-HOLD",
                "path": "assets/SRC001-hold.png",
                "sha256": digest(source_image_two),
                "approval_status": "user_approved",
                "user_approval": {"status": "user_approved", "display_receipt_id": "gallery-proof-001", "approved_at": "2026-08-24T12:01:00+08:00", "asset_sha256": digest(source_image_two)},
                "source_shot_ids": ["SRC001"],
                "responsibility": "动作关键状态2：人物稳定手持并让产品正面可见。",
            },
            {
                "asset_id": "ASSET-ADD001",
                "path": "assets/ADD001.png",
                "sha256": digest(inserted_image),
                "approval_status": "user_approved",
                "user_approval": {"status": "user_approved", "display_receipt_id": "gallery-proof-001", "approved_at": "2026-08-24T12:01:00+08:00", "asset_sha256": digest(inserted_image)},
                "inserted_shot_ids": ["ADD001"],
                "responsibility": "新增近景：产品完整居中并承接口播切点。",
            },
        ],
        "shot_decisions": [
            {
                "shot_id": "S001",
                "selected_asset_ids": ["ASSET-SRC001-LIFT", "ASSET-SRC001-HOLD", "ASSET-ADD001"],
            }
        ],
        "summary": {"reused_frame_count": 3, "new_generation_count": 0, "expected_word_image_count": 3},
    }
    canonical_values = {
        "project.json": {
            "project_id": "delivery-proof-test",
            "product_mode": "replace_product",
            "prompt_length_contract": prompt_contract,
        },
        "library/product_bible.json": {"profile_id": "butter-crisp-v1", "version": 1},
        "library/product_library.json": {},
        "library/style_bible.json": {"profile_id": "ugc-test-v1", "version": 1},
        "library/correction_memory.json": {},
        "library/knowledge_index.json": {},
        "library/avatar_library.json": {},
        "planning/story_plan.json": {"break_plan": {"occurrences": []}},
        "planning/asset_reuse_plan.json": reuse_plan,
        "source/source_manifest.json": {"sha256": "source-video-sha256"},
        "shots/shot_manifest.json": {"shots": [shot]},
    }
    assert set(canonical_values) == set(SNAPSHOT_PAYLOAD_FIELDS)
    for relative_path, value in canonical_values.items():
        write_json(project_dir / relative_path, value)

    prompt = "严格保留原片节奏并执行新增近景。"
    prompt_markdown = f"# S001\n\n```text\n{prompt}\n```\n"
    prompt_path = project_dir / "prompts" / "S001.md"
    prompt_path.parent.mkdir(parents=True, exist_ok=True)
    prompt_path.write_text(prompt_markdown, encoding="utf-8")
    history_dir = project_dir / "prompts" / "history" / compile_id
    history_dir.mkdir(parents=True, exist_ok=True)
    (history_dir / "S001.md").write_text(prompt_markdown, encoding="utf-8")

    pack_shot = {
        "shot_id": "S001",
        "title": shot["title"],
        "timecode": shot["timecode"],
        "visual_type": shot["visual_type"],
        "narrative_role": shot["narrative_role"],
        "delivery_mode": "voiceover",
        "script_segment_ids": ["T001"],
        "risk": shot["risk"],
        "product_profile": "butter-crisp-v1",
        "product_version": 1,
        "source_shot_ids": ["SRC001"],
        "inserted_shot_ids": ["ADD001"],
        "style_profile": "ugc-test-v1",
        "style_version": 1,
        "source_first_frame": "assets/SRC001-lift.png",
        "selected_beauty_keyframe": "assets/SRC001-hold.png",
        "approved_generation_first_frame": "assets/ADD001.png",
        "product_references": ["assets/product-reference.png"],
        "prompt_file": "prompts/S001.md",
        "source_units": [source_unit_value],
        "inserted_units": [inserted_unit_value],
        "prompt_sha256": digest_text(prompt),
        "prompt_file_sha256": digest(prompt_path),
        "prompt_non_whitespace_characters": len(re.sub(r"\s+", "", prompt)),
    }
    input_hashes = canonical_input_hashes(project_dir)
    pack = {
        "schema_version": "1.1",
        "project_id": "delivery-proof-test",
        "compile_id": compile_id,
        "source_sha256": "source-video-sha256",
        "canonical_input_hashes": input_hashes,
        "prompt_length_contract": prompt_contract,
        "history_dir": f"prompts/history/{compile_id}",
        "shots": [pack_shot],
    }
    pack_path = project_dir / "prompts" / "generation_pack.json"
    write_json(pack_path, pack)
    write_json(history_dir / "generation_pack.json", pack)
    snapshot = {
        "schema_version": "1.1",
        "compile_id": compile_id,
        "canonical_input_hashes": input_hashes,
        "prompt_length_contract": prompt_contract,
    }
    for relative_path, snapshot_field in SNAPSHOT_PAYLOAD_FIELDS.items():
        snapshot[snapshot_field] = canonical_values[relative_path]
    snapshot_path = history_dir / "input_snapshot.json"
    write_json(snapshot_path, snapshot)

    asset_records = {
        "ASSET-SRC001-LIFT": reuse_plan["inventory"][0],
        "ASSET-SRC001-HOLD": reuse_plan["inventory"][1],
        "ASSET-ADD001": reuse_plan["inventory"][2],
    }
    image_paths = {
        "ASSET-SRC001-LIFT": source_image_one,
        "ASSET-SRC001-HOLD": source_image_two,
        "ASSET-ADD001": inserted_image,
    }

    document = Document()
    document.add_paragraph("逐分镜执行稿")
    document.add_heading("当前结构结论", level=1)
    document.add_paragraph("每个动作镜头至少绑定一张批准目标帧；同一动作镜头可按动作顺序绑定多张。")
    document.add_heading("生成段总览", level=1)
    document.add_paragraph("S001｜SRC001 + ADD001｜00:00.000–00:04.000")
    document.add_heading("S001｜SRC001 + ADD001｜00:00.000–00:04.000", level=1)
    document.add_paragraph("连续生成片段：原片展示与新增近景｜4.000 秒")
    document.add_paragraph("对齐状态：动作镜头 2 个｜批准目标帧 3 张")
    document.add_heading("动作镜头对应", level=2)
    for kind, unit in (("source", source_unit_value), ("inserted", inserted_unit_value)):
        item_id = unit.get("source_shot_id") or unit.get("inserted_shot_id")
        document.add_heading(f"{item_id}｜{exact_time_label(kind, unit)}", level=3)
        document.add_paragraph(f"准确秒数：{exact_time_label(kind, unit)}")
        if kind == "inserted":
            document.add_paragraph(f"新增原因：{unit['insertion_rationale']}")
            document.add_paragraph(f"节奏锚点：{unit['rhythm_anchor']}")
            document.add_paragraph(f"源片表演依据：{'、'.join(unit['source_reference_shot_ids'])}")
        document.add_paragraph(f"分镜描述：{unit['storyboard_description']}")
        document.add_paragraph(f"口播稿：{unit['script_text']}")

    document.add_heading("目标帧与职责", level=2)
    for unit in (source_unit_value, inserted_unit_value):
        item_id = unit.get("source_shot_id") or unit.get("inserted_shot_id")
        for asset_id in unit["delivery_asset_ids"]:
            asset = asset_records[asset_id]
            responsibility = delivery_asset_responsibility(unit, asset)
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).paragraphs[0].add_run().add_picture(str(image_paths[asset_id]))
            table.cell(0, 0).add_paragraph(target_frame_caption(item_id, asset, responsibility))
    document.add_heading("可复制Prompt原文", level=2)
    document.add_paragraph(prompt)
    docx_path = project_dir / "exports" / "proof.docx"
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)

    def export_unit(kind: str, unit: dict) -> dict:
        delivery_assets = []
        item_id = unit.get("source_shot_id") or unit.get("inserted_shot_id")
        for asset_id in unit["delivery_asset_ids"]:
            asset = asset_records[asset_id]
            image_path = image_paths[asset_id]
            responsibility = delivery_asset_responsibility(unit, asset)
            delivery_assets.append(
                {
                    "asset_id": asset_id,
                    "path": str(image_path.relative_to(project_dir)),
                    "sha256": digest(image_path),
                    "responsibility": responsibility,
                    "display_caption": target_frame_caption(item_id, asset, responsibility),
                }
            )
        value = {
            "generation_timecode": unit["generation_timecode"],
            "exact_time_label": exact_time_label(kind, unit),
            "storyboard_description": unit["storyboard_description"],
            "script_text": unit["script_text"],
            "source_performance_layers": unit["source_performance_layers"],
            "performance_layers_text": format_performance_layers(unit),
            "eating_occurrences": [],
            "eating_occurrence_texts": [],
            "break_occurrences": [],
            "break_occurrence_texts": [],
            "delivery_assets": delivery_assets,
        }
        if kind == "source":
            value.update({"source_shot_id": "SRC001", "source_timecode": unit["source_timecode"]})
        else:
            value.update(
                {
                    "inserted_shot_id": "ADD001",
                    "insertion_rationale": unit["insertion_rationale"],
                    "rhythm_anchor": unit["rhythm_anchor"],
                    "source_reference_shot_ids": unit["source_reference_shot_ids"],
                    "source_reference_frame": unit["source_reference_frame"],
                }
            )
        return value

    parsed = parse_docx_body(docx_path)
    prompt_count = len(re.sub(r"\s+", "", prompt))
    export_manifest = {
        "schema_version": "2.0",
        "project_id": "delivery-proof-test",
        "compile_id": compile_id,
        "canonical_input_hashes": input_hashes,
        "prompt_length_contract": prompt_contract,
        "generation_pack_sha256": digest(pack_path),
        "input_snapshot_path": str(snapshot_path.relative_to(project_dir)),
        "input_snapshot_sha256": digest(snapshot_path),
        "source_sha256": "source-video-sha256",
        "docx_sha256": digest(docx_path),
        "unit_order": ["SRC001", "ADD001"],
        "reused_frame_count": 3,
        "new_generation_count": 0,
        "expected_word_image_count": 3,
        "expected_body_image_occurrence_count": 3,
        "body_image_relationship_count": parsed["all_body_image_relationship_count"],
        "continuity_boundary_reference_count": 0,
        "embedded_media_count": parsed["embedded_media_count"],
        "prompt_non_whitespace_characters": {"S001": prompt_count},
        "compile_snapshot_validated": True,
        "shots": [
            {
                "shot_id": "S001",
                "source_first_frame": "assets/SRC001-lift.png",
                "approved_generation_first_frame": "assets/ADD001.png",
                "source_units": [export_unit("source", source_unit_value)],
                "inserted_units": [export_unit("inserted", inserted_unit_value)],
                "selected_asset_ids": ["ASSET-SRC001-LIFT", "ASSET-SRC001-HOLD", "ASSET-ADD001"],
                "continuity_boundary_references": [],
                "avatar_reference": None,
                "product_references": ["assets/product-reference.png"],
                "package_faces": [],
                "package_face_texts": [],
                "package_face_word_texts": [],
                "prompt_sha256": digest_text(prompt),
                "prompt_file_sha256": digest(prompt_path),
                "prompt_non_whitespace_characters": prompt_count,
            }
        ],
    }
    manifest_path = project_dir / "review" / "proof.manifest.json"
    write_json(manifest_path, export_manifest)
    return {
        "project_dir": project_dir,
        "docx": docx_path,
        "manifest": manifest_path,
        "manifest_value": export_manifest,
        "pack": pack,
        "pack_path": pack_path,
        "snapshot": snapshot,
        "snapshot_path": snapshot_path,
        "shot": shot,
        "images": [source_image_one, source_image_two, inserted_image],
    }


def source_unit(source_id: str, start: float, end: float, asset_id: str, frame: str) -> dict:
    duration = end - start
    return {
        "source_shot_id": source_id,
        "source_timecode": {"start": start, "end": end, "duration": duration},
        "generation_timecode": {"start": 0.0, "end": duration, "duration": duration},
        "storyboard_description": f"{source_id} 的可见动作、构图和节奏描述。",
        "script_text": "对应新版口播",
        "source_first_frame": frame,
        "delivery_asset_ids": [asset_id],
        "source_performance_layers": six_layers(),
    }


def test_source_shot_contract(root: Path) -> None:
    frame_paths = []
    for index in range(1, 4):
        path = root / f"frame-{index}.jpg"
        path.write_bytes(f"frame-{index}".encode())
        frame_paths.append(path)
    source = {
        "duration": 8.0,
        "frame_rate": 30,
        "source_shots": [
            {"id": "SRC001", "start_frame": 0, "end_frame": 60, "timecode": {"start": 0.0, "end": 2.0, "duration": 2.0}, "storyboard_description": "第一原片分镜"},
            {"id": "SRC002", "start_frame": 60, "end_frame": 120, "timecode": {"start": 2.0, "end": 4.0, "duration": 2.0}, "storyboard_description": "第二原片分镜"},
            {"id": "SRC003", "start_frame": 120, "end_frame": 240, "timecode": {"start": 4.0, "end": 8.0, "duration": 4.0}, "storyboard_description": "第三原片分镜"},
        ]
    }
    first = source_unit("SRC001", 0.0, 2.0, "A1", str(frame_paths[0]))
    second = source_unit("SRC002", 2.0, 4.0, "A2", str(frame_paths[1]))
    second["generation_timecode"] = {"start": 2.0, "end": 4.0, "duration": 2.0}
    third = source_unit("SRC003", 4.0, 8.0, "A3", str(frame_paths[2]))
    shots = [
        {"id": "S001", "timecode": {"start": 0.0, "end": 4.0, "duration": 4.0}, "merge_reason": "相邻短镜合并至4秒，全部保留", "source_units": [first, second]},
        {"id": "S002", "timecode": {"start": 4.0, "end": 8.0, "duration": 4.0}, "source_units": [third]},
    ]
    reuse = {
        "inventory": [
            {
                "asset_id": f"A{index}",
                "path": str(path),
                "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
                "approval_status": "approved",
                "source_shot_ids": [f"SRC{index:03d}"],
                "responsibility": f"{f'SRC{index:03d}'} 的批准目标帧",
            }
            for index, path in enumerate(frame_paths, 1)
        ],
        "shot_decisions": [
            {"shot_id": "S001", "selected_asset_ids": ["A1", "A2"]},
            {"shot_id": "S002", "selected_asset_ids": ["A3"]},
        ],
    }
    project = {
        "project_rules": {
            "preserve_every_source_shot": True,
            "require_at_least_one_approved_image_per_source_shot": True,
            "require_at_least_one_approved_image_per_inserted_shot": True,
            "require_structured_six_layer_evidence": True,
            "require_frame_accurate_source_timeline": True,
            "minimum_generation_clip_seconds": 4.0,
        }
    }
    issues: list[dict] = []
    validate_source_shot_contract(root, project, source, shots, reuse, issues)
    assert not issues, issues

    extra_frame = root / "frame-1-action-state-2.jpg"
    extra_frame.write_bytes(b"frame-1-action-state-2")
    multi_asset = copy.deepcopy(shots)
    multi_asset[0]["source_units"][0]["delivery_asset_ids"] = ["A1", "A1B"]
    multi_asset[0]["source_units"][0]["delivery_asset_roles"] = {
        "A1": "动作状态1：产品进入画面。",
        "A1B": "动作状态2：产品在镜头中心稳定展示。",
    }
    multi_reuse = copy.deepcopy(reuse)
    multi_reuse["inventory"].append(
        {
            "asset_id": "A1B",
            "path": str(extra_frame),
            "sha256": hashlib.sha256(extra_frame.read_bytes()).hexdigest(),
            "approval_status": "approved",
            "source_shot_ids": ["SRC001"],
            "responsibility": "动作状态2：产品在镜头中心稳定展示。",
        }
    )
    multi_reuse["shot_decisions"][0]["selected_asset_ids"] = ["A1", "A1B", "A2"]
    issues = []
    validate_source_shot_contract(root, project, source, multi_asset, multi_reuse, issues)
    assert not issues, issues

    zero_asset = copy.deepcopy(shots)
    zero_asset[0]["source_units"][0]["delivery_asset_ids"] = []
    issues = []
    validate_source_shot_contract(root, project, source, zero_asset, reuse, issues)
    assert "SOURCE_SHOT_APPROVED_IMAGE_MISSING" in codes(issues)

    cross_unit_reuse = copy.deepcopy(shots)
    cross_unit_reuse[0]["source_units"][1]["delivery_asset_ids"] = ["A1"]
    cross_reuse_plan = copy.deepcopy(reuse)
    cross_reuse_plan["shot_decisions"][0]["selected_asset_ids"] = ["A1", "A1"]
    issues = []
    validate_source_shot_contract(root, project, source, cross_unit_reuse, cross_reuse_plan, issues)
    assert {
        "DELIVERY_FRAME_DUPLICATED",
        "DELIVERY_FRAME_PROVENANCE_MISMATCH",
    }.intersection(codes(issues))

    missing_layers = copy.deepcopy(shots)
    missing_layers[0]["source_units"][0].pop("source_performance_layers")
    issues = []
    validate_source_shot_contract(root, project, source, missing_layers, reuse, issues)
    assert "SIX_LAYER_EVIDENCE_MISSING" in codes(issues)

    gapped_source = copy.deepcopy(source)
    gapped_source["source_shots"][1]["timecode"] = {"start": 2.1, "end": 4.0, "duration": 1.9}
    issues = []
    validate_source_shot_contract(root, project, gapped_source, shots, reuse, issues)
    assert "SOURCE_TIMELINE_GAP" in codes(issues)

    broken = [dict(shots[0], source_units=[first]), shots[1]]
    issues = []
    validate_source_shot_contract(root, project, source, broken, reuse, issues)
    assert "SOURCE_SHOT_COVERAGE_MISMATCH" in codes(issues)
    assert "GENERATION_CLIP_TOO_SHORT" not in codes(issues), "The group remains 4 seconds; coverage, not duration, is the blocker."


def eating_occurrence(occurrence_id: str, shot_id: str, origin: str) -> dict:
    value = {
        "id": occurrence_id,
        "shot_id": shot_id,
        "origin": origin,
        "generation_timecode": {"start": 0.2, "end": 1.2, "duration": 1.0},
        "rhythm_rationale": "在卖点口播之后插入独立吃食证据，并与前后吃食节奏分开。",
        "source_evidence": ["原片可见人物张口、咬合、产品离嘴"] if origin == "source" else [],
        "insertion_rationale": "原片不足三次，只补足缺少的一次" if origin == "inserted" else None,
        "appetite_evidence": {
            "bite_readability": "牙齿接触和产品离嘴清楚",
            "crisp_sound": "短促咔嚓和少量沙沙声",
            "product_state_change": "同一根形成自然咬口并缩短",
            "source_performance_basis": "继承原片送入口、视线和头部节奏",
        },
        "visible_swallow_required": False,
        "speech_after_bite": {
            "enabled": True,
            "start_trigger": "product_left_mouth",
            "mouth_speakable_evidence": "咬合结束、剩余产品已离嘴，嘴唇与下颌恢复可说状态。",
        },
    }
    if origin == "source":
        value["source_shot_id"] = "SRC001"
    else:
        value["inserted_shot_id"] = f"ADD{occurrence_id[-3:]}"
    return value


def test_eating_plan() -> None:
    shots = [{"id": f"S{index:03d}", "source_units": [], "inserted_units": []} for index in range(1, 6)]
    full_unit_time = {"start": 0.0, "end": 4.0, "duration": 4.0}
    shots[0]["source_units"] = [{"source_shot_id": "SRC001", "generation_timecode": full_unit_time}]
    shots[2]["inserted_units"] = [{"inserted_shot_id": "ADD002", "generation_timecode": full_unit_time}]
    shots[4]["inserted_units"] = [{"inserted_shot_id": "ADD003", "generation_timecode": full_unit_time}]
    story = {
        "eating_plan": {
            "source_duration_seconds": 30.0,
            "source_eating_occurrence_count": 1,
            "inserted_eating_occurrence_count": 2,
            "target_eating_occurrence_count": 3,
            "occurrences": [
                eating_occurrence("E001", "S001", "source"),
                eating_occurrence("E002", "S003", "inserted"),
                eating_occurrence("E003", "S005", "inserted"),
            ],
        }
    }
    project = {
        "project_rules": {
            "minimum_eating_occurrences_when_source_duration_gte_30": 3,
            "eating_occurrences_must_be_non_contiguous": True,
        }
    }
    issues: list[dict] = []
    validate_eating_plan(project, {"duration": 30.0}, story, shots, issues)
    assert not issues, issues

    story["eating_plan"]["inserted_eating_occurrence_count"] = 3
    issues = []
    validate_eating_plan(project, {"duration": 30.0}, story, shots, issues)
    assert "EATING_INSERT_COUNT_MISMATCH" in codes(issues)


def test_break_plan() -> None:
    project = {"project_name": "黄油脆丝棒", "product_profile": "butter-crisp-v1", "project_rules": {}}
    product = {"name": "达尔顿黄油脆丝棒", "profile_id": "butter-crisp-v1"}
    shots = [
        {
            "id": "S001",
            "visual_type": "person_product_showcase",
            "character": {"present": True},
            "source_units": [],
            "inserted_units": [{"inserted_shot_id": "ADD101", "generation_timecode": {"start": 0.0, "end": 4.0, "duration": 4.0}, "storyboard_description": "人物双手掰开同一根脆丝棒，展示互补橙金断面。"}],
            "action_beats": [
                {
                    "id": "BEAT-B001",
                    "start": 0.5,
                    "end": 2.5,
                    "action": "双手掰开同一根脆丝棒并脆裂",
                    "product_change": "同一根折断成两段互补橙金断面，断点落下少量碎屑",
                    "foley_cue": "断裂出现的同一帧同步一次短促咔嚓",
                }
            ],
            "product_state": {"count": 1, "state": "breaking"},
            "audio": {"foley": "断裂帧同步一次咔嚓与少量碎屑声"},
        },
        {
            "id": "S003",
            "visual_type": "product_showcase",
            "character": {"present": False, "hands_only": True},
            "source_units": [],
            "inserted_units": [{"inserted_shot_id": "ADD103", "generation_timecode": {"start": 0.0, "end": 4.0, "duration": 4.0}, "storyboard_description": "无人出镜纯手掰断同一根脆丝棒，停半拍展示断面。"}],
            "action_beats": [
                {
                    "id": "BEAT-B002",
                    "start": 0.2,
                    "end": 2.0,
                    "action": "纯手掰断同一根脆丝棒并脆裂",
                    "product_change": "同一根形成两段互补橙金断面，断点落下少量碎屑",
                    "foley_cue": "断裂出现的同一帧同步一次短促咔嚓",
                }
            ],
            "product_state": {"count": 1, "state": "breaking"},
            "audio": {"foley": "脆裂帧同步短促咔嚓和碎屑声"},
        },
    ]
    proof = {
        "single_snap": True,
        "fracture_visible": True,
        "material_conservation_locked": True,
        "crumbs": {"minimum": 3, "maximum": 8},
        "foley": "一次短促清楚的咔嚓与少量碎屑声",
        "complementary_orange_gold_fracture": "两个断面来自同一断点，均为橙金至焦糖橙且轮廓互补。",
        "same_stick_two_piece_conservation": "断后仅两段，长度之和与断前同一根一致。",
        "sound_sync": "咔嚓严格落在断裂出现的同一帧，碎屑声紧随重力下落。",
    }
    story = {
        "break_plan": {
            "occurrences": [
                {
                    "id": "B001",
                    "shot_id": "S001",
                    "mode": "person_present",
                    "origin": "inserted",
                    "inserted_shot_id": "ADD101",
                    "action_beat_id": "BEAT-B001",
                    "generation_timecode": {"start": 1.0, "end": 2.0, "duration": 1.0},
                    "rhythm_rationale": "人物口播在断裂声前后停半拍。",
                    "insertion_rationale": "按人物出镜节奏加入目标产品脆断证明。",
                    "crisp_proof": {**proof, "action_beat_id": "BEAT-B001"},
                },
                {
                    "id": "B002",
                    "shot_id": "S003",
                    "mode": "hands_only_product",
                    "origin": "inserted",
                    "inserted_shot_id": "ADD103",
                    "action_beat_id": "BEAT-B002",
                    "generation_timecode": {"start": 0.5, "end": 1.5, "duration": 1.0},
                    "rhythm_rationale": "独立纯手部特写承担硬性酥脆证据。",
                    "insertion_rationale": "满足无人出镜纯手单根展示的硬性证明。",
                    "crisp_proof": {**proof, "action_beat_id": "BEAT-B002"},
                },
            ]
        }
    }
    issues: list[dict] = []
    validate_break_plan(project, product, story, shots, issues)
    assert not issues, issues
    exported_break_text = format_break_occurrence(story["break_plan"]["occurrences"][0])
    for expected in (
        "生成镜内 1.000–2.000 秒",
        "新增依据 按人物出镜节奏加入目标产品脆断证明",
        "互补橙金断面=两个断面来自同一断点",
        "同一根两段守恒=断后仅两段",
        "音画同步=咔嚓严格落在断裂出现的同一帧",
    ):
        assert expected in exported_break_text

    metadata_only_shots = copy.deepcopy(shots)
    metadata_only_shots[0]["action_beats"] = [{"action": "展示产品", "product_change": "保持完整"}]
    issues = []
    validate_break_plan(project, product, story, metadata_only_shots, issues)
    assert "BREAK_ACTION_BEAT_MISSING" in codes(issues)

    story["break_plan"]["occurrences"] = story["break_plan"]["occurrences"][:1]
    issues = []
    validate_break_plan(project, product, story, shots, issues)
    assert "HANDS_ONLY_BREAK_SHOWCASE_MISSING" in codes(issues)


def test_package_artwork(root: Path) -> None:
    masters = {}
    for face_name in ("front", "side", "top"):
        master = root / f"{face_name}-master.png"
        master.write_bytes(f"approved-{face_name}-master".encode())
        masters[face_name] = master
    crop = root / "front-candidate-crop.png"
    crop.write_bytes(b"candidate-visible-face")
    parent = root / "approved-delivery.png"
    parent.write_bytes(b"approved-delivery-image")
    projection_manifest = root / "front.projection.json"
    write_json(
        projection_manifest,
        {
            "schema_version": "package-master-projection-v1.0",
            "face": "front",
            "projection_method": "homography",
            "candidate": {"path": str(parent), "sha256": hashlib.sha256(parent.read_bytes()).hexdigest(), "size": [1080, 1920]},
            "master": {"path": str(masters["front"]), "sha256": hashlib.sha256(masters["front"].read_bytes()).hexdigest(), "size": [1024, 1024]},
            "visible_mask": None,
            "target_quad_tl_tr_br_bl": [[10, 10], [710, 10], [710, 710], [10, 710]],
            "output": {"path": str(parent), "sha256": hashlib.sha256(parent.read_bytes()).hexdigest(), "size": [1080, 1920]},
            "model_redraw_used": False,
        },
    )
    project = {"project_rules": {"package_artwork_policy": "preserve_master_projection"}}
    product = {
        "package_artwork": {
            "policy": "preserve_master_projection",
            "minimum_legible_face_area_ratio": 0.08,
            "face_masters": {key: str(value) for key, value in masters.items()},
        }
    }
    front_face = {
        "box_id": "BOX001",
        "face": "front",
        "visibility_state": "occluded",
        "visible_extent": "partial",
        "master_reference": str(masters["front"]),
        "expected_visible_regions": ["左上品牌", "中央主标题"],
        "expected_visible_polygon": [[10, 10], [710, 10], [710, 710], [10, 710]],
        "visible_area_ratio": 0.4,
        "legibility_required": True,
        "occluded_or_offframe_regions": ["右下产品图自然出框"],
        "natural_crop_or_occlusion": True,
        "projection_method": "homography",
        "qa_evidence": {
            "candidate_face_crop": str(crop),
            "candidate_face_crop_sha256": hashlib.sha256(crop.read_bytes()).hexdigest(),
            "delivery_asset_id": "A-BOX",
            "parent_image_sha256": hashlib.sha256(parent.read_bytes()).hexdigest(),
            "crop_rect_xywh": [10, 10, 700, 700],
            "master_sha256": hashlib.sha256(masters["front"].read_bytes()).hexdigest(),
            "projection_manifest": str(projection_manifest),
            "projection_manifest_sha256": hashlib.sha256(projection_manifest.read_bytes()).hexdigest(),
            "visible_region_checkpoints": [
                {"id": "左上品牌", "status": "matched"},
                {"id": "中央主标题", "status": "matched"},
            ],
            "text_legibility": "matched",
            "orientation": "matched",
            "cross_edge_registration": "not_applicable",
            "cross_edge_registration_reason": "本机位没有可见跨棱印刷。",
            "occlusion_scope": "matched",
            "model_redraw_detected": False,
            "unexpected_missing_region": False,
        },
        "qa_status": "approved",
    }
    hidden_faces = [
        {
            "box_id": "BOX001",
            "face": face_name,
            "visibility_state": "hidden",
            "visible_extent": "none",
            "master_reference": str(masters[face_name]),
            "not_applicable_reason": "源片锁定机位下该面完全位于盒体背后。",
        }
        for face_name in ("side", "top")
    ]
    shot = {
        "id": "S001",
        "source_units": [{"source_shot_id": "SRC001", "delivery_asset_ids": ["A-BOX"]}],
        "inserted_units": [],
        "product_state": {
            "packaging": "retail_box",
            "package_artwork": {
                "artwork_scaled_or_relaid_out": False,
                "visible_faces": [front_face, *hidden_faces],
            },
        },
    }
    reuse = {"inventory": [{"asset_id": "A-BOX", "path": str(parent), "approval_status": "approved", "source_shot_ids": ["SRC001"]}]}
    issues: list[dict] = []
    validate_package_artwork(root, project, product, [shot], issues, reuse)
    assert not issues, issues
    exported_package_text = format_package_face(front_face)
    for expected in (
        "BOX001/front",
        "visibility_state=occluded",
        "visible_extent=partial",
        "polygon=",
        "area_ratio=0.4",
        "legibility_required=true",
        f"master={masters['front']}",
        "expected_checkpoints=左上品牌、中央主标题",
        "parent_asset=A-BOX",
        f"parent_sha256={hashlib.sha256(parent.read_bytes()).hexdigest()}",
    ):
        assert expected in exported_package_text

    shot["product_state"]["package_artwork"]["artwork_scaled_or_relaid_out"] = True
    issues = []
    validate_package_artwork(root, project, product, [shot], issues, reuse)
    assert "PACKAGE_ARTWORK_RECOMPOSED" in codes(issues)

    shot["product_state"]["package_artwork"]["artwork_scaled_or_relaid_out"] = False
    shot["product_state"]["package_artwork"]["visible_faces"][0]["qa_evidence"]["visible_region_checkpoints"][1]["status"] = "missing"
    issues = []
    validate_package_artwork(root, project, product, [shot], issues, reuse)
    assert "PACKAGE_ARTWORK_FRAGMENT_MISSING" in codes(issues)


def test_exported_six_layer_gap_reason() -> None:
    layers = {
        key: {
            "status": "template_supplement" if key == "breath_pause" else "not_visible",
            "source_timecode": None,
            "source_reference_frame": "source/SRC001.jpg",
            "observable_evidence": "原片机位没有拍到胸腹起伏，无法直接观察呼吸节奏。",
            "confidence": 0.8,
            "gap_reason": "原片无可见呼吸证据，只按相邻停顿节奏补一处轻微停顿。" if key == "breath_pause" else None,
        }
        for key in (
            "emotion_trigger",
            "gaze",
            "facial_microreaction",
            "body_hand_preparation",
            "breath_pause",
            "voice_speech",
        )
    }
    text = format_performance_layers({"source_performance_layers": layers})
    assert "可观察证据 原片机位没有拍到胸腹起伏" in text
    assert "模板补缺原因 原片无可见呼吸证据" in text


def test_inserted_unit_and_script_coverage(root: Path) -> None:
    source_frame = root / "source-frame.jpg"
    added_frame = root / "added-frame.jpg"
    source_frame.write_bytes(b"source")
    added_frame.write_bytes(b"added")
    source = {
        "duration": 4.0,
        "frame_rate": 30,
        "source_shots": [
            {"id": "SRC001", "start_frame": 0, "end_frame": 120, "timecode": {"start": 0.0, "end": 4.0, "duration": 4.0}, "storyboard_description": "原片人物展示产品"}
        ]
    }
    source_record = source_unit("SRC001", 0.0, 4.0, "A-SRC", str(source_frame))
    source_record["script_text"] = "原片口播"
    inserted_record = {
        "inserted_shot_id": "ADD001",
        "generation_timecode": {"start": 0.0, "end": 4.0, "duration": 4.0},
        "storyboard_description": "根据原片节奏和新版口播新增的非连续吃食证明镜头",
        "script_text": "新增口播",
        "delivery_asset_ids": ["A-ADD"],
        "insertion_rationale": "原片达到30秒但只出现两次吃食，只补缺少的一次",
        "rhythm_anchor": "卖点句结束后的自然切点",
        "source_reference_shot_ids": ["SRC001"],
        "source_reference_frame": str(source_frame),
        "source_performance_layers": six_layers(),
    }
    shots = [
        {
            "id": "S001",
            "timecode": {"start": 0.0, "end": 4.0, "duration": 4.0},
            "source_units": [source_record],
            "inserted_units": [],
            "audio": {"delivery_mode": "voiceover", "script_text": "原片口播"},
        },
        {
            "id": "S002",
            "timecode": {"start": 4.0, "end": 8.0, "duration": 4.0},
            "source_units": [],
            "inserted_units": [inserted_record],
            "audio": {"delivery_mode": "voiceover", "script_text": "新增口播"},
        },
    ]
    reuse = {
        "inventory": [
            {
                "asset_id": "A-SRC",
                "path": str(source_frame),
                "sha256": hashlib.sha256(source_frame.read_bytes()).hexdigest(),
                "approval_status": "approved",
                "source_shot_ids": ["SRC001"],
                "responsibility": "SRC001 原片人物展示目标帧",
            },
            {
                "asset_id": "A-ADD",
                "path": str(added_frame),
                "sha256": hashlib.sha256(added_frame.read_bytes()).hexdigest(),
                "approval_status": "approved",
                "inserted_shot_ids": ["ADD001"],
                "responsibility": "ADD001 新增吃食证明目标帧",
            },
        ],
        "shot_decisions": [
            {"shot_id": "S001", "selected_asset_ids": ["A-SRC"]},
            {"shot_id": "S002", "selected_asset_ids": ["A-ADD"]},
        ],
    }
    project = {
        "project_rules": {
            "preserve_every_source_shot": True,
            "require_at_least_one_approved_image_per_source_shot": True,
            "require_at_least_one_approved_image_per_inserted_shot": True,
            "require_revised_script_full_coverage": True,
            "require_structured_six_layer_evidence": True,
            "require_frame_accurate_source_timeline": True,
            "minimum_generation_clip_seconds": 4.0,
        }
    }
    issues: list[dict] = []
    validate_source_shot_contract(root, project, source, shots, reuse, issues)
    assert not issues, issues

    story = {
        "subtitle_script": {
            "provided_by_user": True,
            "text": "原片口播新增口播",
            "effective_characters": 8,
        }
    }
    issues = []
    validate_revised_script_coverage(project, story, shots, issues)
    assert not issues, issues

    inserted_record["script_text"] = "无"
    issues = []
    validate_revised_script_coverage(project, story, shots, issues)
    assert "WORD_SCRIPT_COVERAGE_MISMATCH" in codes(issues)


def test_delivery_proof_adversaries(root: Path) -> None:
    fixture = build_delivery_proof_fixture(root)
    project_dir = fixture["project_dir"]
    docx_path = fixture["docx"]
    base_manifest = fixture["manifest_value"]
    aligner = Path(__file__).resolve().with_name("align_exports.py")

    def run_alignment(candidate: Path, name: str, should_pass: bool) -> dict:
        manifest_value = copy.deepcopy(base_manifest)
        manifest_value["docx_sha256"] = digest(candidate)
        manifest_value["embedded_media_count"] = parse_docx_body(candidate)["embedded_media_count"]
        manifest_path = project_dir / "review" / f"{name}.manifest.json"
        write_json(manifest_path, manifest_value)
        result = subprocess.run(
            [
                sys.executable,
                str(aligner),
                "--project-dir",
                str(project_dir),
                "--docx",
                str(candidate),
                "--export-manifest",
                str(manifest_path),
                "--require-docx",
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        expected_code = 0 if should_pass else 2
        assert result.returncode == expected_code, result.stdout + result.stderr
        return json.loads((project_dir / "review" / "alignment_manifest.json").read_text(encoding="utf-8"))

    valid_alignment = run_alignment(docx_path, "valid", True)
    assert valid_alignment["summary"]["status"] == "aligned"
    assert valid_alignment["global_checks"]["compile_snapshot_valid"] is True
    assert valid_alignment["global_checks"]["docx_owner_frame_relationship_count_matches_assets"] is True
    assert valid_alignment["shots"][0]["checks"]["selected_asset_ids_match_manifest"] is True
    assert valid_alignment["shots"][0]["checks"]["script_aligned"] is True
    assert valid_alignment["shots"][0]["checks"]["character_count_aligned"] is True

    # Delete only the second SRC001 action-state drawing.  The first SRC001
    # drawing remains, proving the audit rejects a missing member of an ordered
    # multi-image unit instead of merely checking that the unit has some image.
    missing_one = project_dir / "exports" / "src-missing-one-action-state.docx"
    missing_document = Document(docx_path)
    source_hold_caption = (
        "SRC001｜ASSET-SRC001-HOLD｜已批准｜职责："
        "动作关键状态2：人物稳定手持并让产品正面可见。"
    )
    removed = 0
    for table in missing_document.tables:
        if source_hold_caption not in "\n".join(
            cell.text for row in table.rows for cell in row.cells
        ):
            continue
        for drawing in list(table._element.xpath(".//w:drawing")):
            drawing.getparent().remove(drawing)
            removed += 1
    assert removed == 1
    missing_document.save(missing_one)
    missing_alignment = run_alignment(missing_one, "src-missing-one-action-state", False)
    source_card = next(
        unit
        for shot_result in missing_alignment["shots"]
        for unit in shot_result["units"]
        if unit["unit_id"] == "SRC001"
    )
    assert source_card["checks"]["body_card_has_at_least_one_image_relationship"] is True
    assert source_card["checks"]["body_card_image_count_matches_owner_assets"] is False
    assert source_card["checks"]["target_frame_captions_and_hashes_match"] is False

    # Reverse the two SRC001 action-state media payloads while keeping captions
    # and relationship positions fixed.  This is an order-sensitive negative.
    parsed = parse_docx_body(docx_path)
    source_part_one, source_part_two = parsed["units"]["SRC001"]["image_parts"]
    swapped = project_dir / "exports" / "reordered-src-action-states.docx"
    with zipfile.ZipFile(docx_path) as source_archive, zipfile.ZipFile(swapped, "w") as target_archive:
        source_bytes_one = source_archive.read(source_part_one)
        source_bytes_two = source_archive.read(source_part_two)
        for info in source_archive.infolist():
            if info.filename == source_part_one:
                payload = source_bytes_two
            elif info.filename == source_part_two:
                payload = source_bytes_one
            else:
                payload = source_archive.read(info.filename)
            target_archive.writestr(info, payload)
    swapped_alignment = run_alignment(swapped, "reordered-src-action-states", False)
    swapped_source = next(unit for unit in swapped_alignment["shots"][0]["units"] if unit["unit_id"] == "SRC001")
    swapped_add = next(unit for unit in swapped_alignment["shots"][0]["units"] if unit["unit_id"] == "ADD001")
    assert swapped_source["checks"]["target_frame_captions_and_hashes_match"] is False
    assert swapped_add["checks"]["target_frame_captions_and_hashes_match"] is True

    # Relabel one SRC image as if ADD001 owned it.  The actual image and OOXML
    # relationship stay untouched, so only strict caption-owner binding catches
    # the impersonation.
    wrong_owner = project_dir / "exports" / "wrong-target-frame-owner.docx"
    wrong_owner_document = Document(docx_path)
    wrong_owner_count = 0
    for table in wrong_owner_document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text != source_hold_caption:
                        continue
                    paragraph.text = source_hold_caption.replace("SRC001｜", "ADD001｜", 1)
                    wrong_owner_count += 1
    assert wrong_owner_count == 1
    wrong_owner_document.save(wrong_owner)
    wrong_owner_alignment = run_alignment(wrong_owner, "wrong-target-frame-owner", False)
    wrong_source = next(unit for unit in wrong_owner_alignment["shots"][0]["units"] if unit["unit_id"] == "SRC001")
    wrong_add = next(unit for unit in wrong_owner_alignment["shots"][0]["units"] if unit["unit_id"] == "ADD001")
    assert wrong_source["checks"]["body_card_image_count_matches_owner_assets"] is False
    assert wrong_add["checks"]["body_card_image_count_matches_owner_assets"] is False

    # Each editable fact is independently mandatory; images and canonical
    # Prompt remain present so none of these failures can hide behind a generic
    # "document has content" check.
    deleted_labels = {
        "missing-seconds": f"准确秒数：{exact_time_label('inserted', fixture['shot']['inserted_units'][0])}",
        "missing-description": f"分镜描述：{fixture['shot']['inserted_units'][0]['storyboard_description']}",
        "missing-script": f"口播稿：{fixture['shot']['inserted_units'][0]['script_text']}",
    }
    for name, exact_paragraph in deleted_labels.items():
        candidate = project_dir / "exports" / f"{name}.docx"
        document = Document(docx_path)
        matches = [paragraph for paragraph in document.paragraphs if paragraph.text == exact_paragraph]
        assert len(matches) == 1, exact_paragraph
        matches[0]._element.getparent().remove(matches[0]._element)
        document.save(candidate)
        alignment = run_alignment(candidate, name, False)
        inserted_result = next(unit for unit in alignment["shots"][0]["units"] if unit["unit_id"] == "ADD001")
        assert inserted_result["checks"]["editable_labels_match"] is False

    project = json.loads((project_dir / "project.json").read_text(encoding="utf-8"))
    shots = [fixture["shot"]]
    _, _, valid_errors = validate_compile_snapshot(project_dir, project, shots, fixture["pack"])
    assert not valid_errors, valid_errors

    # A canonical shot edit after compile must make both the input hash/snapshot
    # and the per-shot payload stale.  Exporter calls this validator before it
    # writes any Word bytes.
    shot_manifest_path = project_dir / "shots" / "shot_manifest.json"
    original_shot_manifest = shot_manifest_path.read_bytes()
    changed_shot_manifest = json.loads(original_shot_manifest)
    changed_shot_manifest["shots"][0]["source_units"][0]["storyboard_description"] += "（编译后篡改）"
    write_json(shot_manifest_path, changed_shot_manifest)
    _, _, stale_errors = validate_compile_snapshot(
        project_dir, project, changed_shot_manifest["shots"], fixture["pack"]
    )
    assert any("canonical inputs changed after compile" in value for value in stale_errors)
    assert any("generation_pack.source_units is stale or mixed" in value for value in stale_errors)
    assert any("input_snapshot.shot_manifest differs" in value for value in stale_errors)
    shot_manifest_path.write_bytes(original_shot_manifest)

    # Missing snapshot payload is a hard failure even when copied hash claims
    # still match.  This closes the previous "hash declaration only" loophole.
    snapshot_path = fixture["snapshot_path"]
    original_snapshot = snapshot_path.read_bytes()
    incomplete_snapshot = json.loads(original_snapshot)
    incomplete_snapshot.pop("shot_manifest")
    write_json(snapshot_path, incomplete_snapshot)
    _, _, snapshot_errors = validate_compile_snapshot(project_dir, project, shots, fixture["pack"])
    assert any("input_snapshot.shot_manifest missing" in value for value in snapshot_errors)
    snapshot_path.write_bytes(original_snapshot)

    mixed_pack = copy.deepcopy(fixture["pack"])
    mixed_pack["shots"][0]["inserted_units"][0]["script_text"] = "混入另一版口播"
    _, _, mixed_errors = validate_compile_snapshot(project_dir, project, shots, mixed_pack)
    assert any("active generation_pack differs" in value for value in mixed_errors)
    assert any("generation_pack.inserted_units is stale or mixed" in value for value in mixed_errors)

    inventory = {
        item["asset_id"]: item
        for item in json.loads((project_dir / "planning" / "asset_reuse_plan.json").read_text(encoding="utf-8"))["inventory"]
    }
    multi_asset_unit = copy.deepcopy(fixture["shot"]["source_units"][0])
    resolved_assets, multi_asset_errors = validate_unit_delivery_asset_binding(
        project_dir, "S001", "SRC001", multi_asset_unit, inventory
    )
    assert [item["asset_id"] for item in resolved_assets] == ["ASSET-SRC001-LIFT", "ASSET-SRC001-HOLD"]
    assert not multi_asset_errors, multi_asset_errors

    zero_asset_unit = copy.deepcopy(multi_asset_unit)
    zero_asset_unit["delivery_asset_ids"] = []
    _, zero_asset_errors = validate_unit_delivery_asset_binding(
        project_dir, "S001", "SRC001", zero_asset_unit, inventory
    )
    assert any("at least one approved target-frame asset_id" in value for value in zero_asset_errors)

    duplicate_asset_unit = copy.deepcopy(multi_asset_unit)
    duplicate_asset_unit["delivery_asset_ids"] = ["ASSET-SRC001-LIFT", "ASSET-SRC001-LIFT"]
    _, duplicate_asset_errors = validate_unit_delivery_asset_binding(
        project_dir, "S001", "SRC001", duplicate_asset_unit, inventory
    )
    assert any("delivery_asset_ids contains a duplicate" in value for value in duplicate_asset_errors)

    # Disabled means both bounds are truly off even if stale numeric decoration
    # exists; enabled means both positive bounds must be present and ordered.
    assert normalized_prompt_length_contract(
        {
            "prompt_length_contract": {
                "enabled": False,
                "minimum_non_whitespace_characters": 3000,
                "maximum_non_whitespace_characters": 4000,
            }
        }
    ) == {
        "enabled": False,
        "minimum_non_whitespace_characters": 0,
        "maximum_non_whitespace_characters": 0,
    }
    try:
        normalized_prompt_length_contract(
            {"prompt_length_contract": {"enabled": True, "minimum_non_whitespace_characters": 3000}}
        )
    except ValueError:
        pass
    else:
        raise AssertionError("An enabled Prompt length contract without both bounds must fail")


def main() -> None:
    with tempfile.TemporaryDirectory(prefix="remix-contract-test-") as temp_dir:
        root = Path(temp_dir)
        test_source_shot_contract(root)
        test_package_artwork(root)
        test_inserted_unit_and_script_coverage(root)
        test_delivery_proof_adversaries(root)
    test_eating_plan()
    test_break_plan()
    test_exported_six_layer_gap_reason()
    print("DELIVERY CONTRACT TESTS PASSED")


if __name__ == "__main__":
    main()
