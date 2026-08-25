#!/usr/bin/env python3
"""Run a deterministic end-to-end test of project creation, planning, lint and compile."""

from __future__ import annotations

import json
import hashlib
import copy
import re
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

from PIL import Image
from docx import Document


SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

from init_project import initialize_project  # noqa: E402
from pipeline import compile_project, lint_project, verify_prompt_delivery  # noqa: E402
from align_exports import parse_docx_body  # noqa: E402


def read_json(path: Path):
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def write_json(path: Path, value) -> None:
    with path.open("w", encoding="utf-8") as handle:
        json.dump(value, handle, ensure_ascii=False, indent=2)
        handle.write("\n")


def main() -> int:
    with tempfile.TemporaryDirectory(prefix="jimeng-skill-test-") as temp_dir:
        root = Path(temp_dir)
        project_dir = initialize_project(
            name="榴莲大福测试",
            output=root,
            product_profile="durian-daifuku-v1",
            style_profile="ugc-food-review-v1",
            project_id="test-project",
            execution_tier="full_delivery",
        )
        assert (project_dir / "planning" / "workflow_state.json").is_file()
        assert (project_dir / "planning" / "skill_update_candidates.json").is_file()
        assert (project_dir / "review" / "alignment_manifest.json").is_file()

        source_file = project_dir / "source" / "original.mp4"
        source_frame = project_dir / "source" / "test-source.jpg"
        beauty_frame = project_dir / "source" / "test-beauty.jpg"
        first_frame = project_dir / "source" / "test-approved.jpg"
        added_frame = project_dir / "source" / "test-added-approved.jpg"
        product_reference = project_dir / "source" / "product-reference.jpg"
        package_crop = project_dir / "source" / "package-front-crop.png"
        package_masters = {
            face: project_dir / "source" / f"package-{face}-master.png"
            for face in ("front", "side", "top")
        }
        package_projection_manifest = project_dir / "review" / "S002-front.projection.json"
        source_file.write_bytes(b"test-video-placeholder")
        for path, color in (
            (source_frame, (210, 180, 150)),
            (beauty_frame, (200, 170, 140)),
            (first_frame, (190, 160, 130)),
            (added_frame, (160, 190, 135)),
            (product_reference, (230, 210, 180)),
        ):
            Image.new("RGB", (180, 320), color).save(path)
        Image.new("RGB", (120, 120), (245, 180, 80)).save(package_crop)
        for index, path in enumerate(package_masters.values()):
            Image.new("RGB", (400, 400), (230 - index * 15, 170 + index * 10, 70)).save(path)
        write_json(
            package_projection_manifest,
            {
                "schema_version": "package-master-projection-v1.0",
                "face": "front",
                "projection_method": "homography",
                "candidate": {"path": str(added_frame), "sha256": hashlib.sha256(added_frame.read_bytes()).hexdigest(), "size": [180, 320]},
                "master": {"path": str(package_masters["front"]), "sha256": hashlib.sha256(package_masters["front"].read_bytes()).hexdigest(), "size": [400, 400]},
                "visible_mask": None,
                "target_quad_tl_tr_br_bl": [[10, 10], [170, 10], [170, 250], [10, 250]],
                "output": {"path": str(added_frame), "sha256": hashlib.sha256(added_frame.read_bytes()).hexdigest(), "size": [180, 320]},
                "model_redraw_used": False,
            },
        )

        project = read_json(project_dir / "project.json")
        project["source_video"] = "source/original.mp4"
        project["generation_mode"] = "text_to_video"
        project["product_profile"] = "self-test-product-v1"
        project["project_rules"]["packaging_visible"] = True
        write_json(project_dir / "project.json", project)

        product_bible = read_json(project_dir / "library" / "product_bible.json")
        product_bible["profile_id"] = "self-test-product-v1"
        product_bible["package_artwork"] = {
            "policy": "preserve_master_projection",
            "minimum_legible_face_area_ratio": 0.08,
            "face_masters": {
                face: f"source/package-{face}-master.png" for face in ("front", "side", "top")
            },
        }
        write_json(project_dir / "library" / "product_bible.json", product_bible)
        product_library = read_json(project_dir / "library" / "product_library.json")
        product_library["products"][0]["id"] = "self-test-product-v1"
        write_json(project_dir / "library" / "product_library.json", product_library)

        source_manifest = read_json(project_dir / "source" / "source_manifest.json")
        source_manifest.update(
            {
                "source_video": "source/original.mp4",
                "sha256": "test-sha256",
                "duration": 4.0,
                "width": 1080,
                "height": 1920,
                "frame_rate": 30.0,
                "video_first_frame": "source/test-source.jpg",
                "source_shots": [
                    {
                        "id": "SRC001",
                        "timecode": {"start": 0.0, "end": 4.0, "duration": 4.0},
                        "start_frame": 0,
                        "end_frame": 120,
                        "storyboard_description": "人物在暖色家庭餐厅双手掰开原食物并短距离展示断面。",
                    }
                ],
            }
        )
        write_json(project_dir / "source" / "source_manifest.json", source_manifest)

        source_script = "你看这个冰皮，轻轻一拉就能看出它有多软糯。"
        added_script = "再把横截面举近一点，馅料层次也看得很清楚。"
        full_script = source_script + added_script
        story_plan = read_json(project_dir / "planning" / "story_plan.json")
        story_plan.update(
            {
                "status": "reviewed",
                "subtitle_script": {
                    "provided_by_user": True,
                    "path": None,
                    "text": full_script,
                    "effective_characters": len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", full_script)),
                    "language": "zh-CN",
                },
                "source_style_assessment": {
                    "delivery_style": "voiceover_dominant",
                    "observed_voiceover_ratio": 1.0,
                    "observed_on_screen_speech_ratio": 0.0,
                    "observed_silent_ratio": 0.0,
                    "notes": ["原片人物只展示产品，不直接讲话"],
                },
                "narrative_logic": {
                    "hook": "用拉伸动作快速建立软糯卖点",
                    "product_promise": "冰皮柔软而不形成细丝",
                    "visual_proof": "双手掰开并短距离拉伸",
                    "eating_experience": "由后续人物吃产品镜头承接",
                    "closing_payoff": "以满足表情确认口感",
                },
                "story_arc": {
                    "opening_emotional_hook": "人物带着想证明软糯卖点的期待先把动作递到镜头前",
                    "desire_build": "外皮形变与缓慢拉开让期待逐步绷紧",
                    "proof_turn": "拉带和馅料出现后，人物从专注转为被结果击中的惊喜",
                    "sensory_payoff": "横截面近景把柔软和馅料层次兑现为可见证据",
                    "closing_impulse": "人物带着确认后的满足把视线送回观众",
                },
                "delivery_strategy": {
                    "mode": "voiceover_dominant",
                    "rationale": "保持原片人物只展示产品的风格，让字幕稿作为画外音覆盖动作证据。",
                    "voiceover_target_ratio": 1.0,
                    "on_screen_speech_target_ratio": 0.0,
                    "silent_target_ratio": 0.0,
                },
                "visual_mix_targets": {
                    "product_showcase": {"min_ratio": 0.0, "max_ratio": 0.0},
                    "person_product_showcase": {"min_ratio": 1.0, "max_ratio": 1.0},
                    "person_eating": {"min_ratio": 0.0, "max_ratio": 0.0},
                },
                "pacing": {
                    "opening_hook_seconds": 3.0,
                    "target_average_shot_seconds": 4.0,
                    "minimum_generation_clip_seconds": 4.0,
                    "maximum_single_shot_seconds": 5.0,
                    "maximum_on_screen_chars_per_second": 5.0,
                    "maximum_voiceover_chars_per_second": 5.5,
                    "rhythm_notes": ["单镜头自测"],
                },
                "segments": [
                    {
                        "id": "T001",
                        "text": source_script,
                        "delivery_mode": "voiceover",
                        "delivery_rationale": "画面动作承担证明，台词无需人物口型。",
                        "assigned_shots": ["S001"],
                    },
                    {
                        "id": "T002",
                        "text": added_script,
                        "delivery_mode": "voiceover",
                        "delivery_rationale": "新增横截面镜跟随上一镜动作节奏，以画外音补充层次卖点。",
                        "assigned_shots": ["S002"],
                    },
                ],
            }
        )
        write_json(project_dir / "planning" / "story_plan.json", story_plan)

        manifest = {
            "schema_version": "1.0",
            "version": 1,
            "source_analysis_status": "reviewed",
            "shots": [
                {
                    "id": "S001",
                    "title": "掰开并短距离拉伸",
                    "visual_type": "person_product_showcase",
                    "narrative_role": "hook",
                    "script_segment_ids": ["T001"],
                    "scene_rationale": "家庭餐厅环境延续原片生活感，让人物双手和断面成为唯一视觉重点。",
                    "timecode": {"start": 0.0, "end": 4.0, "duration": 4.0},
                    "source_units": [
                        {
                            "source_shot_id": "SRC001",
                            "source_timecode": {"start": 0.0, "end": 4.0, "duration": 4.0},
                            "generation_timecode": {"start": 0.0, "end": 4.0, "duration": 4.0},
                            "storyboard_description": "人物在暖色家庭餐厅双手掰开原食物并短距离展示断面，镜头固定。",
                            "script_text": source_script,
                            "source_first_frame": "source/test-source.jpg",
                            "delivery_asset_ids": ["FRAME-S001-APPROVED"],
                            "source_performance_layers": {
                                key: {
                                    "status": "observed" if key != "voice_speech" else "audible",
                                    "source_timecode": {"start": 0.0, "end": 4.0, "duration": 4.0},
                                    "source_reference_frame": "source/test-source.jpg",
                                    "observable_evidence": "原片可见人物视线、面部和手部动作" if key != "voice_speech" else "原片画外音节奏可听",
                                    "confidence": 0.95,
                                    "gap_reason": None,
                                }
                                for key in (
                                    "emotion_trigger",
                                    "gaze",
                                    "facial_microreaction",
                                    "body_hand_preparation",
                                    "breath_pause",
                                    "voice_speech",
                                )
                            },
                        }
                    ],
                    "purpose": "展示薄冰皮的柔软和短距离延展",
                    "source_facts": ["人物双手在胸前掰开原食物", "暖色家庭背景", "镜头基本固定"],
                    "source_locks": ["保持人物、双手位置、构图、机位、景深、暖色光和动作方向"],
                    "allowed_changes": ["只替换人物手中的食物主体"],
                    "scene": {
                        "location": "暖色家庭餐厅",
                        "background": ["浅米色墙面", "虚化木质家具", "暖黄色落地灯"],
                        "foreground": ["人物双手和榴莲大福"]
                    },
                    "character": {
                        "present": True,
                        "identity": "年轻亚洲女性，深色长发，浅色居家上衣",
                        "position": "画面中央偏上",
                        "gaze": "先看食物，完成拉伸后抬眼看镜头",
                        "micro_expressions": ["掰开前专注", "看到拉伸后眼睛略微睁大", "结尾轻轻点头"]
                    },
                    "emotion": {
                        "start": "专注期待",
                        "trigger": "双手拉开后看见内部榴莲馅形成真实拉带",
                        "inferred_intention": "向镜头证明大福柔软、满馅并值得期待",
                        "progression": ["拉开时短暂停顿", "看到真实拉带后自然惊喜"],
                        "end": "满足确认",
                        "narrative_payoff": "观众通过她从专注到惊喜的变化直观看到产品卖点得到证明",
                        "evidence_basis": ["双手相反方向施力", "视线落在拉带", "眼睛略微睁大", "结尾轻点头"],
                        "intensity": "natural"
                    },
                    "action_beats": [
                        {
                            "start": 0.0,
                            "end": 1.2,
                            "actor": "人物双手",
                            "action": "从两侧轻轻捏住大福并缓慢向相反方向施力",
                            "expression": "低头专注观察",
                            "product_change": "完整外皮先产生柔软形变",
                            "camera_response": "保持固定中近景并把焦点落在双手"
                        },
                        {
                            "start": 1.2,
                            "end": 3.1,
                            "actor": "人物双手",
                            "action": "继续拉开至约4厘米并停留半秒",
                            "expression": "眼睛略微睁大，嘴角出现轻微惊喜",
                            "product_change": "形成2至4条短而宽的奶白糯米皮拉带并少量露出金黄榴莲馅",
                            "camera_response": "轻微转焦到拉带和不规则断面"
                        },
                        {
                            "start": 3.1,
                            "end": 4.0,
                            "actor": "人物双手和视线",
                            "action": "停止拉伸并让拉带轻微回缩，同时抬眼看镜头",
                            "expression": "满足地轻轻点头",
                            "product_change": "拉带弯曲并自然搭回馅料边缘",
                            "camera_response": "保持产品清晰并让人物脸部仍可辨认"
                        }
                    ],
                    "product_state": {
                        "profile": "self-test-product-v1",
                        "state": "stretched",
                        "count": "1",
                        "packaging": "none",
                        "shot_specific_traits": ["本镜头拉伸距离约4厘米"]
                    },
                    "camera": {
                        "shot_size": "正面中近景",
                        "angle": "与双手基本平齐",
                        "movement": "基本固定，仅在断面出现时轻微转焦",
                        "focus": "双手、白色冰皮拉带和少量榴莲馅",
                        "lens_feel": "真实手机近距离镜头"
                    },
                    "lighting": {
                        "source": "室内暖光和原片柔和正面光",
                        "temperature": "warm",
                        "notes": ["白色冰皮不能因暖光变黄", "人物和食物光向一致"]
                    },
                    "audio": {
                        "delivery_mode": "voiceover",
                        "delivery_rationale": "字幕稿解释拉伸卖点，人物画面专注展示动作。",
                        "script_text": source_script,
                        "speech_timing": None,
                        "speech_capacity": {
                            "segment_count": 1,
                            "effective_characters": 19,
                            "speakable_seconds": 4.0,
                            "characters_per_second": 4.75,
                            "excluded_intervals": []
                        },
                        "voice_direction": "年轻女性自然画外音，像向熟人分享；前半句轻快，软糯二字轻微加重",
                        "foley": ["手指接触糯米皮的轻微软糯摩擦声", "安静室内底噪"],
                        "music": "低音量轻快无歌词音乐"
                    },
                    "continuity": ["人物服装和发型与前一镜一致", "产品外皮颜色保持奶白"],
                    "hard_constraints": ["人物不讲话", "无字幕", "无水印", "无包装", "只替换食物主体"],
                    "prohibited": ["蜘蛛网细丝", "黄色拉丝", "大量爆浆", "改变人物和构图"],
                    "asset_links": {
                        "source_first_frame": "source/test-source.jpg",
                        "beauty_keyframe_candidates": ["source/test-beauty.jpg"],
                        "selected_beauty_keyframe": "source/test-beauty.jpg",
                        "approved_generation_first_frame": "source/test-approved.jpg",
                        "product_references": ["source/product-reference.jpg"],
                        "avatar_reference": None
                    },
                    "risk": {"level": "high", "reasons": ["双手掰开", "产品拉伸", "不规则横截面"]},
                    "status": "reviewed"
                }
            ]
        }
        added_shot = copy.deepcopy(manifest["shots"][0])
        added_layers = copy.deepcopy(manifest["shots"][0]["source_units"][0]["source_performance_layers"])
        added_shot.update(
            {
                "id": "S002",
                "title": "新增横截面近距证明",
                "narrative_role": "evidence",
                "script_segment_ids": ["T002"],
                "timecode": {"start": 4.0, "end": 8.0, "duration": 4.0},
                "source_units": [],
                "inserted_units": [
                    {
                        "inserted_shot_id": "ADD001",
                        "generation_timecode": {"start": 0.0, "end": 4.0, "duration": 4.0},
                        "storyboard_description": "承接上一镜动作节奏，人物把同一枚已掰开的两半大福举近镜头，短暂停留展示横截面层次。",
                        "script_text": added_script,
                        "delivery_asset_ids": ["FRAME-S002-APPROVED"],
                        "insertion_rationale": "新版口播增加横截面层次卖点，原片只有一次较远掰开展示，因此补一枚有源片依据的近距证明镜。",
                        "rhythm_anchor": "紧接 SRC001 掰开动作完成后的自然切点，承接第二句画外音。",
                        "source_reference_shot_ids": ["SRC001"],
                        "source_reference_frame": "source/test-source.jpg",
                        "source_performance_layers": added_layers,
                    }
                ],
                "purpose": "按新版口播补充横截面层次的近距视觉证据",
            }
        )
        added_shot["audio"].update(
            {
                "script_text": added_script,
                "delivery_rationale": "新增近距证明镜承接上一镜动作，画外音说明横截面层次。",
                "speech_capacity": {
                    "segment_count": 1,
                    "effective_characters": len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", added_script)),
                    "speakable_seconds": 4.0,
                    "characters_per_second": len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", added_script)) / 4.0,
                    "excluded_intervals": [],
                },
            }
        )
        added_shot["asset_links"].update(
            {
                "approved_generation_first_frame": "source/test-added-approved.jpg",
                "source_first_frame": "source/test-source.jpg",
            }
        )
        front_face = {
            "box_id": "BOX001",
            "face": "front",
            "visibility_state": "occluded",
            "visible_extent": "partial",
            "master_reference": "source/package-front-master.png",
            "expected_visible_regions": ["左上品牌", "中央主标题"],
            "expected_visible_polygon": [[10, 10], [170, 10], [170, 250], [10, 250]],
            "visible_area_ratio": 0.4,
            "legibility_required": True,
            "occluded_or_offframe_regions": ["右下产品图自然出框"],
            "natural_crop_or_occlusion": True,
            "projection_method": "homography",
            "qa_evidence": {
                "candidate_face_crop": "source/package-front-crop.png",
                "candidate_face_crop_sha256": hashlib.sha256(package_crop.read_bytes()).hexdigest(),
                "delivery_asset_id": "FRAME-S002-APPROVED",
                "parent_image_sha256": hashlib.sha256(added_frame.read_bytes()).hexdigest(),
                "crop_rect_xywh": [10, 10, 160, 240],
                "master_sha256": hashlib.sha256(package_masters["front"].read_bytes()).hexdigest(),
                "projection_manifest": "review/S002-front.projection.json",
                "projection_manifest_sha256": hashlib.sha256(package_projection_manifest.read_bytes()).hexdigest(),
                "visible_region_checkpoints": [
                    {"id": "左上品牌", "status": "matched"},
                    {"id": "中央主标题", "status": "matched"},
                ],
                "text_legibility": "matched",
                "orientation": "matched",
                "cross_edge_registration": "not_applicable",
                "cross_edge_registration_reason": "本机位没有可见跨棱印刷。",
                "occlusion_scope": "matched",
                "model_redraw_detected": False,
                "unexpected_missing_region": False,
            },
            "qa_status": "approved",
        }
        hidden_faces = [
            {
                "box_id": "BOX001",
                "face": face,
                "visibility_state": "hidden",
                "visible_extent": "none",
                "master_reference": f"source/package-{face}-master.png",
                "not_applicable_reason": "源片锁定机位下该面完全位于盒体背后。",
            }
            for face in ("side", "top")
        ]
        added_shot["product_state"]["packaging"] = "retail_box"
        added_shot["product_state"]["package_artwork"] = {
            "artwork_scaled_or_relaid_out": False,
            "visible_faces": [front_face, *hidden_faces],
        }
        manifest["shots"].append(added_shot)
        for shot_index, shot_value in enumerate(manifest["shots"], 1):
            shot_value["emotion"].update(
                {
                    "persona_drive": "怕观众只听卖点不相信，所以急着用手里的真实变化证明",
                    "primary_emotion": "挖到产品证据后的兴奋",
                    "secondary_emotions": ["拉开前的期待绷紧", "看到层次后的惊喜"],
                    "undertone": "带一点怕观众不信的认真较真",
                    "residue": "证明完成后仍舍不得收住的满足和分享冲动",
                    "emotion_vocabulary": ["期待绷紧", "眼睛发亮", "惊喜上扬", "满足落稳", "急着分享"],
                    "commercial_turn": "先用动作勾起期待，再以形变/断面证明，最后落到可信的分享感",
                    "creative_enhancement": {"status": "none", "terms": [], "observable_execution": []},
                }
            )
            beat_details = [
                ("手指接触柔软外皮", ["期待绷紧"], "视线锁在产品、眉心轻收", "画外音起音轻快", "继续拉开"),
                ("拉带和馅料进入视野", ["被惊喜击中"], "双眼短促放大、眉峰提起", "卖点词重音落稳", "停止受力并展示"),
                ("产品变化已经得到证明", ["满足确认", "急着分享"], "眉心舒展、抬眼看回镜头", "句尾带笑意回落", "把证据留给下一镜"),
            ]
            for beat_index, (beat, details) in enumerate(zip(shot_value["action_beats"], beat_details), 1):
                trigger, emotion_terms, visible_change, voice_change, next_action = details
                beat.update(
                    {
                        "id": f"S{shot_index:03d}-B{beat_index:02d}",
                        "trigger": trigger,
                        "emotion_terms": emotion_terms,
                        "visible_change": visible_change,
                        "voice_change": voice_change,
                        "next_action": next_action,
                    }
                )
        write_json(project_dir / "shots" / "shot_manifest.json", manifest)

        reuse_plan = read_json(project_dir / "planning" / "asset_reuse_plan.json")
        reuse_plan.update(
            {
                "status": "reviewed",
                "contract_binding": {
                    "bundle_release_id": "video-remix-1.0.9",
                    "prompt_authoring_contract": "narrative-six-layer-v1",
                    "product_profile": "self-test-product-v1",
                },
                "scope": {
                    "current_project": ".",
                    "historical_packages": [str(project_dir)],
                    "requested_operations": ["product_replace", "reuse_frames"],
                },
                "inventory": [
                    {
                        "asset_id": "FRAME-S001-APPROVED",
                        "asset_type": "approved_frame",
                        "library_layer": "scene_shot",
                        "path": "source/test-approved.jpg",
                        "sha256": hashlib.sha256(first_frame.read_bytes()).hexdigest(),
                        "width": 180,
                        "height": 320,
                        "approval_status": "user_approved",
                        "user_approval": {
                            "status": "user_approved", "display_receipt_id": "gallery-self-test-001",
                            "approved_at": "2026-08-24T12:01:00+08:00",
                            "asset_sha256": hashlib.sha256(first_frame.read_bytes()).hexdigest(),
                        },
                        "responsibility": "SRC001 动作关键状态：双手掰开并展示大福横截面。",
                        "rights_status": "cleared",
                        "source_project": "test-project",
                        "origin_bundle_release_id": "video-remix-1.0.9",
                        "origin_product_profile": "self-test-product-v1",
                        "source_shot_ids": ["SRC001"],
                        "avatar_ids": [],
                        "product_ids": ["self-test-product-v1"],
                        "product_states": ["stretched"],
                        "has_source_subtitles": False,
                        "has_watermark": False,
                        "is_composite_or_contact_sheet": False,
                    },
                    {
                        "asset_id": "FRAME-S002-APPROVED",
                        "asset_type": "approved_frame",
                        "library_layer": "scene_shot",
                        "path": "source/test-added-approved.jpg",
                        "sha256": hashlib.sha256(added_frame.read_bytes()).hexdigest(),
                        "width": 180,
                        "height": 320,
                        "approval_status": "user_approved",
                        "user_approval": {
                            "status": "user_approved", "display_receipt_id": "gallery-self-test-001",
                            "approved_at": "2026-08-24T12:01:00+08:00",
                            "asset_sha256": hashlib.sha256(added_frame.read_bytes()).hexdigest(),
                        },
                        "responsibility": "ADD001 动作关键状态：近距举起两半大福展示馅料层次。",
                        "rights_status": "cleared",
                        "source_project": "test-project",
                        "origin_bundle_release_id": "video-remix-1.0.9",
                        "origin_product_profile": "self-test-product-v1",
                        "source_shot_ids": [],
                        "inserted_shot_ids": ["ADD001"],
                        "avatar_ids": [],
                        "product_ids": ["self-test-product-v1"],
                        "product_states": ["stretched"],
                        "has_source_subtitles": False,
                        "has_watermark": False,
                        "is_composite_or_contact_sheet": False,
                    },
                ],
                "shot_decisions": [
                    {
                        "shot_id": "S001",
                        "decision": "reuse",
                        "candidate_asset_ids": ["FRAME-S001-APPROVED"],
                        "selected_asset_ids": ["FRAME-S001-APPROVED"],
                        "required_avatar_ids": [],
                        "required_product_ids": ["self-test-product-v1"],
                        "required_product_states": ["stretched"],
                        "identity_review": "not_applicable",
                        "product_review": "matched",
                        "scene_action_review": "matched",
                        "allowed_deterministic_transforms": [],
                        "generation_reason": None,
                        "candidate_rejection_reasons": [],
                    },
                    {
                        "shot_id": "S002",
                        "decision": "reuse",
                        "candidate_asset_ids": ["FRAME-S002-APPROVED"],
                        "selected_asset_ids": ["FRAME-S002-APPROVED"],
                        "required_avatar_ids": [],
                        "required_product_ids": ["self-test-product-v1"],
                        "required_product_states": ["stretched"],
                        "identity_review": "not_applicable",
                        "product_review": "matched",
                        "scene_action_review": "matched",
                        "allowed_deterministic_transforms": [],
                        "generation_reason": None,
                        "candidate_rejection_reasons": [],
                    },
                ],
                "summary": {
                    "reused_frame_count": 2,
                    "new_generation_count": 0,
                    "rejected_asset_count": 0,
                    "expected_word_image_count": 2,
                },
                "gallery_receipt": {
                    "status": "user_approved", "display_receipt_id": "gallery-self-test-001",
                    "displayed_at": "2026-08-24T12:00:00+08:00", "approved_at": "2026-08-24T12:01:00+08:00",
                    "asset_refs": [
                        {"shot_id": "S001", "asset_id": "FRAME-S001-APPROVED", "sha256": hashlib.sha256(first_frame.read_bytes()).hexdigest()},
                        {"shot_id": "S002", "asset_id": "FRAME-S002-APPROVED", "sha256": hashlib.sha256(added_frame.read_bytes()).hexdigest()},
                    ],
                },
            }
        )
        write_json(project_dir / "planning" / "asset_reuse_plan.json", reuse_plan)

        lint = lint_project(project_dir)
        assert lint["counts"]["ERROR"] == 0, lint
        compiled = compile_project(project_dir)
        assert compiled["shot_count"] == 2
        assert compiled["execution_tier"] == "full_delivery"
        assert compiled["delivery_status"] == "authorized"
        assert (project_dir / "prompts" / "S001.md").is_file()
        assert (project_dir / "prompts" / "S002.md").is_file()
        generation_pack = read_json(project_dir / "prompts" / "generation_pack.json")
        assert generation_pack["compile_id"] == compiled["compile_id"]
        assert generation_pack["compile_input_sha256"]
        assert generation_pack["canonical_input_hashes"]
        for shot_entry in generation_pack["shots"]:
            assert shot_entry["compile_id"] == generation_pack["compile_id"]
            assert shot_entry["compile_input_sha256"] == generation_pack["compile_input_sha256"]
            assert shot_entry["canonical_input_hashes"] == generation_pack["canonical_input_hashes"]
            assert shot_entry["shot_input_sha256"]
            assert isinstance(shot_entry["source_units"], list)
            assert isinstance(shot_entry["inserted_units"], list)
            assert shot_entry["source_units"] or shot_entry["inserted_units"]
            assert shot_entry["prompt_sha256"]
            assert shot_entry["prompt_file_sha256"]

        prompt_text = (project_dir / "prompts" / "S001.md").read_text(encoding="utf-8")
        assert "口播在本镜头作为后期画外音" in prompt_text
        assert "人物展示产品" in prompt_text
        assert "【原片叙事复原】" in prompt_text
        assert "【原片逐时动作】" in prompt_text
        assert "【最小纠错附录】" in prompt_text
        assert (project_dir / "review" / "shot_cards.md").is_file()
        prompt_match = re.search(r"```text\s*\n(.*?)\n```", prompt_text, re.S)
        assert prompt_match
        body = prompt_match.group(1).strip()
        assert "observed" not in body
        assert "not_visible" not in body

        # Prompt length is owned solely by project.json.prompt_length_contract.
        # The default project keeps that gate disabled, so a post-compile padding
        # rewrite would correctly make the compile snapshot stale and must never
        # be part of the successful export path.
        prompt_path = project_dir / "prompts" / "S001.md"
        docx_path = project_dir / "exports" / "test.docx"
        exported = subprocess.run(
            [
                sys.executable,
                str(SCRIPT_DIR / "export_jimeng_docx.py"),
                "--project-dir",
                str(project_dir),
                "--out",
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        assert exported.returncode == 0, exported.stdout + exported.stderr
        export_manifest = read_json(project_dir / "review" / "test.manifest.json")
        assert export_manifest["embedded_media_count"] == 2
        assert export_manifest["reused_frame_count"] == 2
        assert export_manifest["prompt_length_contract"] == {
            "enabled": False,
            "minimum_non_whitespace_characters": 0,
            "maximum_non_whitespace_characters": 0,
        }
        deprecated_length_flags = subprocess.run(
            [
                sys.executable,
                str(SCRIPT_DIR / "export_jimeng_docx.py"),
                "--project-dir",
                str(project_dir),
                "--out",
                str(project_dir / "exports" / "wrong-cli-length-contract.docx"),
                "--min-prompt-chars",
                "3000",
                "--max-prompt-chars",
                "4000",
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        assert deprecated_length_flags.returncode != 0
        assert "project.json.prompt_length_contract" in (deprecated_length_flags.stdout + deprecated_length_flags.stderr)

        aligned = subprocess.run(
            [
                sys.executable,
                str(SCRIPT_DIR / "align_exports.py"),
                "--project-dir",
                str(project_dir),
                "--docx",
                str(docx_path),
                "--require-docx",
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        assert aligned.returncode == 0, aligned.stdout + aligned.stderr + json.dumps(
            read_json(project_dir / "review" / "alignment_manifest.json"), ensure_ascii=False, indent=2
        )
        alignment = read_json(project_dir / "review" / "alignment_manifest.json")
        assert alignment["summary"]["status"] == "aligned"
        assert alignment["shots"][0]["checks"]["prompt_text_aligned"] is True
        assert alignment["shots"][0]["checks"]["frame_aligned"] is True
        assert next(item for item in alignment["shots"] if item["shot_id"] == "S002")["checks"]["package_artwork_aligned"] is True
        assert alignment["summary"]["unit_count"] == 2
        assert not (project_dir / "exports" / "完整逐分镜Prompt.txt").exists(), "Final aligner is an audit, not a TXT exporter."

        def audit_wrong_docx(candidate: Path, name: str) -> dict:
            candidate_manifest = json.loads(json.dumps(export_manifest))
            candidate_manifest["docx_sha256"] = hashlib.sha256(candidate.read_bytes()).hexdigest()
            candidate_manifest_path = project_dir / "review" / f"{name}.manifest.json"
            write_json(candidate_manifest_path, candidate_manifest)
            result = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT_DIR / "align_exports.py"),
                    "--project-dir",
                    str(project_dir),
                    "--docx",
                    str(candidate),
                    "--export-manifest",
                    str(candidate_manifest_path),
                    "--require-docx",
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            assert result.returncode == 2, result.stdout + result.stderr
            return read_json(project_dir / "review" / "alignment_manifest.json")

        # A valid OPC Word file with the approved image only on its cover must
        # still fail: cover media cannot substitute for the SRC001 body card.
        wrong_docx = project_dir / "exports" / "wrong-cover-only.docx"
        wrong_document = Document()
        wrong_document.add_picture(str(first_frame))
        wrong_document.add_heading("S001", level=1)
        wrong_document.add_paragraph(body)
        wrong_document.save(wrong_docx)
        rejected_alignment = audit_wrong_docx(wrong_docx, "wrong-cover-only")
        rejected_codes = {item.get("code") for item in rejected_alignment.get("blockers", [])}
        assert "UNIT_ALIGNMENT_FAILED" in rejected_codes

        # Remove ADD001's owner-frame drawing while leaving its media part in
        # the package.  Global media presence must not pass.
        missing_body_image_docx = project_dir / "exports" / "wrong-missing-body-image.docx"
        missing_body_document = Document(docx_path)
        removed_drawings = 0
        for table in missing_body_document.tables:
            if (
                "ADD001｜FRAME-S002-APPROVED｜已批准｜职责："
                "ADD001 动作关键状态：近距举起两半大福展示馅料层次。"
            ) not in "\n".join(cell.text for row in table.rows for cell in row.cells):
                continue
            for drawing in list(table._element.xpath(".//w:drawing")):
                drawing.getparent().remove(drawing)
                removed_drawings += 1
        assert removed_drawings == 1
        missing_body_document.save(missing_body_image_docx)
        missing_image_alignment = audit_wrong_docx(missing_body_image_docx, "wrong-missing-body-image")
        add_result = next(
            unit
            for shot_result in missing_image_alignment["shots"]
            for unit in shot_result["units"]
            if unit["unit_id"] == "ADD001"
        )
        assert add_result["checks"]["target_frame_captions_and_hashes_match"] is False
        assert add_result["checks"]["body_card_has_at_least_one_image_relationship"] is False

        # Swap ADD001's relationship target bytes with SRC001's image.  Captions
        # and asset IDs remain correct, so only per-card image hashing can catch it.
        parsed = parse_docx_body(docx_path)
        source_media = parsed["units"]["SRC001"]["image_parts"][0]
        inserted_media = parsed["units"]["ADD001"]["image_parts"][0]
        swapped_docx = project_dir / "exports" / "wrong-swapped-body-image.docx"
        with zipfile.ZipFile(docx_path) as source_archive, zipfile.ZipFile(swapped_docx, "w") as target_archive:
            source_bytes = source_archive.read(source_media)
            for info in source_archive.infolist():
                data = source_bytes if info.filename == inserted_media else source_archive.read(info.filename)
                target_archive.writestr(info, data)
        swapped_alignment = audit_wrong_docx(swapped_docx, "wrong-swapped-body-image")
        swapped_add = next(
            unit
            for shot_result in swapped_alignment["shots"]
            for unit in shot_result["units"]
            if unit["unit_id"] == "ADD001"
        )
        assert swapped_add["checks"]["target_frame_captions_and_hashes_match"] is False

        # Delete one editable spoken-script paragraph.  The image and Prompt are
        # untouched; unit-level script alignment must still block delivery.
        missing_script_docx = project_dir / "exports" / "wrong-missing-script.docx"
        missing_script_document = Document(docx_path)
        removed_scripts = 0
        for paragraph in list(missing_script_document.paragraphs):
            if paragraph.text == f"口播稿：{added_script}":
                paragraph._element.getparent().remove(paragraph._element)
                removed_scripts += 1
        assert removed_scripts == 1
        missing_script_document.save(missing_script_docx)
        missing_script_alignment = audit_wrong_docx(missing_script_docx, "wrong-missing-script")
        missing_script_shot = next(item for item in missing_script_alignment["shots"] if item["shot_id"] == "S002")
        assert missing_script_shot["checks"]["script_aligned"] is False

        # Hidden faces still require explicit front/side/top inventory rows.
        # Removing the editable top-face row must fail even though the visible
        # front-face image and all global package media remain present.
        missing_package_face_docx = project_dir / "exports" / "wrong-missing-package-face.docx"
        missing_package_document = Document(docx_path)
        removed_package_rows = 0
        for paragraph in list(missing_package_document.paragraphs):
            if paragraph.text.startswith("包装盒面：BOX001/top｜"):
                paragraph._element.getparent().remove(paragraph._element)
                removed_package_rows += 1
        assert removed_package_rows == 1
        missing_package_document.save(missing_package_face_docx)
        missing_package_alignment = audit_wrong_docx(missing_package_face_docx, "wrong-missing-package-face")
        missing_package_shot = next(item for item in missing_package_alignment["shots"] if item["shot_id"] == "S002")
        assert missing_package_shot["checks"]["package_artwork_aligned"] is False

        # A prompt edited after compile must block export even when the text
        # remains visually equivalent (here only an extra trailing newline).
        compiled_prompt_bytes = prompt_path.read_bytes()
        prompt_path.write_bytes(compiled_prompt_bytes + b"\n")
        stale_docx = project_dir / "exports" / "stale.docx"
        stale_export = subprocess.run(
            [
                sys.executable,
                str(SCRIPT_DIR / "export_jimeng_docx.py"),
                "--project-dir",
                str(project_dir),
                "--out",
                str(stale_docx),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        assert stale_export.returncode != 0
        assert "stale or mixed compile snapshot" in (stale_export.stdout + stale_export.stderr)
        assert not stale_docx.exists()
        prompt_path.write_bytes(compiled_prompt_bytes)

        candidates = read_json(project_dir / "planning" / "skill_update_candidates.json")
        candidates["candidates"] = [
            {
                "candidate_id": "RULE-TEST-001",
                "category": "alignment",
                "observed_problem": "测试问题",
                "proposed_rule": "每次导出后核对唯一事实源哈希。",
                "scope": "cross_project",
                "evidence": ["review/alignment_manifest.json"],
                "target_skill": "jimeng-video-remix-director",
                "target_resource": "references/workflow.md",
                "risk_level": "medium",
                "interaction_surfaces": ["prompt_compile", "docx_alignment"],
                "regression_case_ids": ["GOLDEN-DOCX-001"],
                "replaces": [],
                "rollback_trigger": "唯一事实源哈希或 Word 对齐回退。",
                "status": "new",
                "user_approved": False,
            }
        ]
        write_json(project_dir / "planning" / "skill_update_candidates.json", candidates)
        reviewed = subprocess.run(
            [sys.executable, str(SCRIPT_DIR / "review_skill_candidates.py"), "--project-dir", str(project_dir)],
            capture_output=True,
            text=True,
            check=False,
        )
        assert reviewed.returncode == 0, reviewed.stdout + reviewed.stderr
        assert (project_dir / "review" / "skill_update_report.md").is_file()

        # prompt_only skips generation-frame and candidate-image QA, but it may
        # not skip any canonical narrative/performance/product Prompt inputs.
        project = read_json(project_dir / "project.json")
        project["execution_tier"] = "prompt_only"
        project["status"] = "analyzed"
        write_json(project_dir / "project.json", project)
        workflow = read_json(project_dir / "planning" / "workflow_state.json")
        workflow["execution_tier"] = "prompt_only"
        workflow["prompt_delivery_authorized"] = False
        workflow["compile_receipt"] = None
        write_json(project_dir / "planning" / "workflow_state.json", workflow)
        shot_manifest = read_json(project_dir / "shots" / "shot_manifest.json")
        for shot in shot_manifest["shots"]:
            shot.setdefault("asset_links", {})["approved_generation_first_frame"] = None
            for unit in [*(shot.get("source_units") or []), *(shot.get("inserted_units") or [])]:
                unit["delivery_asset_ids"] = []
        write_json(project_dir / "shots" / "shot_manifest.json", shot_manifest)
        reuse_plan = read_json(project_dir / "planning" / "asset_reuse_plan.json")
        reuse_plan["inventory"] = []
        reuse_plan["shot_decisions"] = []
        reuse_plan["summary"] = {
            "reused_frame_count": 0,
            "new_generation_count": 0,
            "rejected_asset_count": 0,
            "expected_word_image_count": 0,
        }
        reuse_plan["gallery_receipt"] = None
        write_json(project_dir / "planning" / "asset_reuse_plan.json", reuse_plan)

        prompt_only_lint = lint_project(project_dir)
        assert prompt_only_lint["counts"]["ERROR"] == 0, prompt_only_lint
        prompt_only_compile = compile_project(project_dir)
        assert prompt_only_compile["execution_tier"] == "prompt_only"
        assert prompt_only_compile["delivery_status"] == "authorized"
        assert (project_dir / "prompts" / "canonical_prompt_only.md").is_file()
        assert verify_prompt_delivery(project_dir)["status"] == "authorized"

        # Copying the visible receipt fields into a minimal hand-written
        # workflow cannot impersonate the compiler-owned state machine.
        workflow_path = project_dir / "planning" / "workflow_state.json"
        compiled_workflow = read_json(workflow_path)
        forged_workflow = {
            "execution_tier": "prompt_only",
            "compile_receipt": copy.deepcopy(compiled_workflow["compile_receipt"]),
        }
        write_json(workflow_path, forged_workflow)
        forged_verification = verify_prompt_delivery(project_dir)
        assert forged_verification["status"] == "blocked"
        assert "PROMPT_DELIVERY_WORKFLOW_CLAIM_MISMATCH" in {
            item["code"] for item in forged_verification["errors"]
        }
        write_json(workflow_path, compiled_workflow)
        assert verify_prompt_delivery(project_dir)["status"] == "authorized"

        rogue_prompt = project_dir / "prompts" / "总控_手写_prompt_only.md"
        rogue_prompt.write_text("这是一段绕过编译器的手写 Prompt。\n", encoding="utf-8")
        bypass_lint = lint_project(project_dir, write_report=False)
        assert "NON_CANONICAL_PROMPT_BYPASS" in {item["code"] for item in bypass_lint["issues"]}
        bypass_verification = verify_prompt_delivery(project_dir)
        assert bypass_verification["status"] == "blocked"
        assert "NON_CANONICAL_PROMPT_BYPASS" in {item["code"] for item in bypass_verification["errors"]}
        rogue_prompt.unlink()

        # The strict full-delivery path is unchanged: the same project cannot
        # compile without independently approved generation frames.
        project = read_json(project_dir / "project.json")
        project["execution_tier"] = "full_delivery"
        project["generation_mode"] = "image_to_video"
        write_json(project_dir / "project.json", project)
        workflow = read_json(project_dir / "planning" / "workflow_state.json")
        workflow["execution_tier"] = "full_delivery"
        write_json(project_dir / "planning" / "workflow_state.json", workflow)
        strict_lint = lint_project(project_dir, write_report=False)
        strict_codes = {item["code"] for item in strict_lint["issues"]}
        assert "missing_approved_generation_first_frame" in strict_codes
        assert strict_lint["counts"]["ERROR"] > 0

    print("SELF TEST PASSED")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
