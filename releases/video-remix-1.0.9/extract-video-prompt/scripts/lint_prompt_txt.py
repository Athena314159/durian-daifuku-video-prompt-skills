#!/usr/bin/env python3
"""Lint the mechanical and role-lock rules of a compiled video-prompt TXT.

The script deliberately does not claim that content review is complete. It can
find encodable omissions and contradictions; a human still has to compare the
source video, audio, revised script, and generated pixels.
"""

from __future__ import annotations

import argparse
import json
import posixpath
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from validate_text_handoff import validate_text_handoff


SHOT_RE = re.compile(r"(?m)^S\d{3}｜.*$")
PROMPT_RE = re.compile(
    r"【完整Prompt(?:｜主体非空白字符数：\d+)?】\s*(.*?)(?=\n【(?:原片动作对应|内容审核记录)】|\n={10,}|\Z)",
    re.S,
)
PROMPT_COUNT_RE = re.compile(r"【完整Prompt｜主体非空白字符数：(\d+)】")
SCRIPT_RE = re.compile(r"【口播稿】\s*(.*?)(?=\n【完整Prompt|\n={10,}|\Z)", re.S)
FIELD_RE = re.compile(
    r"(?m)^(?P<name>原片时间|源分镜ID|独立生成时长|分镜描述|人物位置|声音方式|产品形态|尺度模式|投影模式|投影事实源|生成首帧|分镜图|核心主体|核心动作|核心产品|适用表演层)：(?P<value>.*)$"
)
TIME_RE = re.compile(r"(?m)^\s*0\.00(?:0)?\s*(?:秒)?\s*[–—-]")
SOURCE_TIME_RE = re.compile(
    r"(?m)^原片时间：(?P<start>\d{1,2}:\d{2}(?:\.\d{1,3})?)[–—-](?P<end>\d{1,2}:\d{2}(?:\.\d{1,3})?)\s*$"
)
TIMED_ACTION_LINE_RE = re.compile(
    r"(?m)^\s*\d+(?:\.\d+)?\s*(?:秒)?\s*[–—-]\s*\d+(?:\.\d+)?\s*(?:秒)?\s*[:：].*$"
)
DIALOGUE_RE = re.compile(r"(?m)^\s*([^：\n]+)：[“\"](.*?)[”\"]\s*$")
DURATION_RE = re.compile(r"(?m)^独立生成时长：([0-9]+(?:\.[0-9]+)?)秒")
SPEECH_SECONDS_RE = re.compile(r"(?m)^实际可说时段：([0-9]+(?:\.[0-9]+)?)秒")
SPEECH_RATE_RE = re.compile(r"(?m)^计划语速：([0-9]+(?:\.[0-9]+)?)字/秒")
SEGMENT_COUNT_RE = re.compile(r"(?m)^本镜句段数：([0-9]+)")
STATUS_RE = re.compile(r"(?i)(?<![A-Za-z])(PASS|FAIL|ERROR)(?![A-Za-z])")
INTERNAL_PRODUCT_RE = re.compile(
    r"(?i)(?<![A-Za-z0-9_])(?:V[1-5]|whole|bitten|person_eating|hand_tear|cut_open|two_halves)(?![A-Za-z0-9_])"
)
GENERIC_SPEECH_PLACEHOLDERS = (
    "沿用原片生活口语节奏",
    "沿用原片生活口语风格",
    "沿用原片口语节奏",
    "保持原片口语感",
)
ACCENT_SPECIFIC_TERMS = (
    "北方",
    "东北",
    "京腔",
    "山东",
    "河南",
    "川渝",
    "四川",
    "重庆",
    "江浙",
    "吴语",
    "粤语",
    "广府",
    "闽南",
    "台湾",
    "港式",
    "湖南",
    "湖北",
    "平翘舌",
    "前后鼻音",
    "卷舌",
    "儿化",
    "鼻音",
)

NO_TEXT_GROUPS = (
    ("禁止", "严禁", "不得", "不出现", "不能出现"),
    ("字幕", "对白文字", "自动字幕"),
    ("水印",),
)
PERFORMANCE_GROUPS = {
    "视频核心": ("核心", "叙事作用", "真正要表达", "情绪回报点", "情绪落点"),
    "视线轨迹": ("视线", "目光", "看向", "抬眼", "回看"),
    "五官微反应": ("眉峰", "眉眼", "眼睑", "瞳孔", "鼻翼", "嘴角", "下巴", "喉结"),
    "身体或手部准备": ("肩", "重心", "手指", "手腕", "掌根", "指腹", "身体", "上身"),
    "呼吸或停顿": ("呼吸", "吸气", "呼气", "换气", "停顿", "半拍", "四分之一秒", "鼻息"),
    "讲话细节": ("语速", "重音", "尾音", "起音", "音量", "口吻", "口语", "连读", "断句", "音色"),
}
BITE_CHAIN_GROUPS = {
    "接近口部": ("送向", "接近", "靠近"),
    "张口或咬合": ("张嘴", "张口", "咬合", "咬入"),
    "产品离嘴或撤回": ("离嘴", "撤回", "抽离", "移开"),
}
EATING_PHASE_GROUPS = {
    "approach": ("送向口部", "接近口部", "靠近嘴", "送到嘴边"),
    "open_mouth": ("张嘴", "张口"),
    "bite_contact": ("牙齿接触", "咬合", "咬入"),
    "crisp_fracture": ("脆断", "咬断", "清脆断裂", "咔嚓"),
    "withdraw": ("产品离嘴", "离嘴", "撤回", "抽离", "移开嘴边"),
    "closed_chew": ("闭口咀嚼", "闭嘴咀嚼", "嘴唇闭合咀嚼"),
    "swallow": ("吞咽", "咽下", "喉结轻动"),
    "post_eating_reaction": ("吃后反应", "满足地笑", "满意地点头", "闭眼回味", "嗯～", "嗯~"),
    "onscreen_speech_resume": ("马上说", "立即说", "开口说", "接着说", "恢复口播"),
}
ALLOWED_EATING_PHASES = frozenset(EATING_PHASE_GROUPS)
MOUTH_BUSY_TERMS = ("咬合", "咬入", "咬断", "闭口咀嚼", "闭嘴咀嚼", "嘴唇闭合咀嚼")
ONSCREEN_SPEECH_TERMS = ("屏内说", "开口说", "说出", "接着说", "马上说", "立即说", "同步口播")
MOUTH_SPEECH_CONFLICT_PATTERNS = (
    re.compile(r"(?:咬合|咬入|咬断|闭口咀嚼|闭嘴咀嚼|嘴唇闭合咀嚼)(?:中|时|期间).*?(?:屏内说|开口说|说出|口播)"),
    re.compile(r"(?:一边|边).*?(?:咬|咀嚼).*?(?:一边|边|同时).*?(?:说|口播)"),
)
CORE_FACT_FIELDS = ("核心主体", "核心动作", "核心产品")
VISIBLE_TERMS = ("出现", "出镜", "进入画面", "露出", "可见", "纳入画面", "展示")
BODY_TERMS = ("脸", "侧脸", "嘴", "头发", "头部", "身体", "影子", "倒影", "自拍")
NEGATIVE_TERMS = ("不出现", "不出镜", "不得", "禁止", "严禁", "不能", "不可", "避免", "不生成", "绝不")
NEGATIVE_PROMPT_MARKERS = ("禁止", "严禁", "不得", "不要", "避免", "不能", "不可", "不出现", "不生成", "绝不")
NARRATIVE_PROMPT_HEADERS = (
    "【生成目标与叙事职责】",
    "【口播原文与声源】",
    "【原片叙事复原】",
    "【原片逐时动作】",
    "【产品与动作物理】",
    "【摄影、灯光与声音】",
    "【最小纠错附录】",
)
EMOTION_TRIGGER_TERMS = ("触发", "听到", "看到", "咬下", "接到", "感到", "当", "因为")
CHARACTER_INTENTION_TERMS = ("意图", "想要", "想把", "想让", "试图", "为了", "此刻想")
EMOTION_ARC_TERMS = ("从", "转为", "逐渐", "随即", "随后", "最终", "落到")
NARRATIVE_PAYOFF_TERMS = ("让观众", "观众感受到", "镜头落点", "情绪落点", "承接下一拍", "叙事落点")
SIX_LAYER_AUDIT_TOKENS = ("observed", "audible", "not_visible", "not_applicable", "template_supplement", "gap_reason", "confidence")
BUTTER_CRISP_PRODUCT_TERMS = ("黄油脆丝棒", "脆丝棒")
BUTTER_CRISP_BARE_STATE_TERMS = ("完整未破", "开袋并露出", "手持", "摆盘", "掰断", "断面", "咬食", "咬口", "碎屑")
BUTTER_CRISP_MATERIAL_GROUPS = {
    "实体片状覆盖层": ("实体片状", "片状脆丝", "片状碎片", "实体脆丝碎片", "实体材料层"),
    "独立厚度": ("自身厚度", "独立厚度", "侧边厚度"),
    "遮挡与缝隙": ("前后遮挡", "互相遮挡", "重叠缝隙", "不规则窄缝", "暗缝"),
    "轮廓凸出": ("轮廓凸出", "凸出并打破", "打破长边轮廓", "打破干净外轮廓"),
    "平面图案排除": ("平面贴图", "印刷图案", "浅浮雕", "压花", "光滑橙色基底"),
}
BUTTER_CRISP_BOX_STATE_TERMS = ("零售外盒", "外盒", "三盒", "两盒", "一盒", "盒体", "包装盒")
BUTTER_CRISP_BOX_DIMENSION_TERMS = ("15 × 15 × 4.5", "15×15×4.5", "15*15*4.5")
BUTTER_CRISP_BOX_RATIO_TERMS = ("1:1", "正方形正面", "正面近似正方形")
BUTTER_CRISP_BOX_DEPTH_TERMS = ("30%", "0.3边长", "0.3 边长", "厚度约为正面边长")
BUTTER_CRISP_PRODUCT_LENGTH_TERMS = ("12 cm", "12cm", "12 厘米", "约12厘米", "约 12 厘米")
BUTTER_CRISP_PRODUCT_DIMENSION_TERMS = ("12 × 2.5 × 1 cm", "12×2.5×1 cm", "12*2.5*1 cm", "12厘米×2.5厘米×1厘米")
BUTTER_CRISP_PRODUCT_TARGET_ASPECT_TERMS = ("4.8:1", "4.8：1", "12:2.5")
BUTTER_CRISP_PRODUCT_ASPECT_TERMS = ("4:1–5:1", "4:1-5:1", "4:1至5:1", "4:1~5:1", "4:1—5:1")
BUTTER_CRISP_PRODUCT_PACKAGE_RATIO_TERMS = ("0.80", "80%", "12:15")
BUTTER_CRISP_PERSPECTIVE_TERMS = ("同平面", "同一平面", "同距离", "透视校正", "跨景深")
BUTTER_CRISP_BREAK_STATE_TERMS = ("掰断", "两段断面", "断面展示")
BUTTER_CRISP_PROJECTION_MODES = {"source_pixel_lock", "intentional_depth_move"}
BUTTER_CRISP_FAILED_PROJECTION_SOURCE_TERMS = ("失败", "未批准", "候选", "generated", "生成结果")
BUTTER_CRISP_CAMERA_AXIS_ADVANCE_TERMS = (
    "推近镜头",
    "紧贴镜头",
    "占据前景",
    "交替前推",
    "贴到镜头前",
    "极近景",
)
BUTTER_CRISP_SOURCE_PIXEL_LOCKS = {
    "HAND_FACE_PROJECTION_LOCK_MISSING": (
        "脸、手掌和手脸比例的投影容差",
        ("脸框0.97–1.03", "脸框0.97-1.03", "手脸比例±5%", "手脸比例误差不超过5%"),
    ),
    "PALM_SCALE_LOCK_MISSING": (
        "手掌投影尺寸容差",
        ("手掌框0.95–1.05", "手掌框0.95-1.05", "双手联合框0.95–1.05", "双手联合框0.95-1.05"),
    ),
    "WRIST_POSITION_LOCK_MISSING": (
        "腕部位置容差",
        ("腕部中心位移不超过画幅2%", "腕部中心位移≤画幅2%", "腕部位置误差不超过画幅2%"),
    ),
    "LOCAL_CONTACT_EDIT_LOCK_MISSING": (
        "只允许产品和接触指尖局部变化",
        ("只允许产品与接触指尖", "仅允许产品和接触指尖", "拇指和食指末节局部"),
    ),
}
BUTTER_CRISP_BREAK_LOCKS = {
    "BREAK_SOURCE_SINGLE_INSTANCE_MISSING": (
        "同一根来源",
        ("同一根", "一根完整产品", "一根完整脆丝棒"),
    ),
    "BREAK_FRACTION_MISSING": (
        "断裂位置或 break_fraction",
        ("break_fraction", "断裂位置", "近中点断裂", "接近中部"),
    ),
    "BROKEN_LENGTH_CONSERVATION_MISSING": (
        "两段长度守恒",
        ("两段长度之和", "左右段长度之和", "两段总长度"),
    ),
    "BROKEN_WIDTH_THICKNESS_LOCK_MISSING": (
        "断后宽度和厚度保持",
        ("宽度都继续约2.5", "宽度继续约2.5", "宽度保持2.5", "宽度保持 2.5"),
    ),
    "BROKEN_SEGMENT_ASPECT_RULE_MISSING": (
        "按断裂位置计算的半段长宽比",
        ("2.2:1–2.6:1", "2.2:1-2.6:1", "按实际分数计算", "按断裂位置计算"),
    ),
    "HAND_PRODUCT_SCALE_ANCHOR_MISSING": (
        "同平面手指或断前完整棒尺度锚点",
        ("拇指指腹", "食指末节", "断前投影尺寸", "原始首帧的双手"),
    ),
    "BREAK_CAMERA_DEPTH_LOCK_MISSING": (
        "断裂前后相机和双手深度锁",
        ("不向镜头额外推进", "不得为了突出断面", "手到镜头距离"),
    ),
    "FRACTURE_SURFACE_LOCK_MISSING": (
        "同一断裂点及不规则断面",
        ("同一断裂点", "两个断面来自同一", "断口不规则"),
    ),
    "FRACTURE_VIEW_MODE_MISSING": (
        "唯一断面展示朝向 fracture_view_mode",
        ("fracture_view_mode", "长外表面朝镜头", "断面直接朝镜头"),
    ),
}


@dataclass(frozen=True)
class Issue:
    code: str
    message: str


def split_shots(text: str) -> list[tuple[str, str]]:
    matches = list(SHOT_RE.finditer(text))
    return [
        (
            match.group(0).strip(),
            text[match.start() : matches[i + 1].start()] if i + 1 < len(matches) else text[match.start() :],
        )
        for i, match in enumerate(matches)
    ]


def shot_id(title: str) -> str:
    return title.split("｜", 1)[0]


def field(block: str, name: str) -> str:
    for match in FIELD_RE.finditer(block):
        if match.group("name") == name:
            return match.group("value").strip()
    return ""


def normalize_text(value: str) -> str:
    return re.sub(r"[\s，。！？!?；;、“”\"：:…]+", "", value)


def parse_timecode(value: str) -> float:
    """Parse MM:SS(.mmm) into seconds."""
    match = re.fullmatch(r"(?P<minutes>\d{1,3}):(?P<seconds>\d{2})(?:\.(?P<fraction>\d{1,3}))?", value.strip())
    if not match:
        raise ValueError(f"非法时间码：{value}")
    seconds = int(match.group("seconds"))
    if seconds >= 60:
        raise ValueError(f"非法时间码：{value}")
    fraction = match.group("fraction") or "0"
    milliseconds = int(fraction.ljust(3, "0"))
    return int(match.group("minutes")) * 60 + seconds + milliseconds / 1000


def source_time_range(block: str) -> tuple[float, float] | None:
    match = SOURCE_TIME_RE.search(block)
    if not match:
        return None
    try:
        return parse_timecode(match.group("start")), parse_timecode(match.group("end"))
    except ValueError:
        return None


def source_shot_ids(block: str) -> list[str]:
    value = field(block, "源分镜ID")
    if not value or value == "新增镜头":
        return []
    return re.findall(r"SRC\d+", value, flags=re.I)


def core_fact_tokens(value: str) -> list[str]:
    """Core fields are authored as compact atomic facts separated by punctuation."""
    return [
        token.strip()
        for token in re.split(r"[、,，;；/|]+", value)
        if token.strip() and token.strip() not in {"无", "无目标产品", "不适用"}
    ]


def repeated_prompt_sentences(prompt: str) -> list[str]:
    """Find exact long-sentence padding while ignoring short necessary refrains."""
    counts: dict[str, tuple[int, str]] = {}
    for sentence in re.split(r"[。！？!?\n]+", prompt):
        normalized = normalize_text(sentence)
        if len(normalized) < 24:
            continue
        count, _ = counts.get(normalized, (0, sentence.strip()))
        counts[normalized] = (count + 1, sentence.strip())
    return [original for count, original in counts.values() if count > 1]


def negative_prompt_ratio(prompt: str) -> tuple[float, int, int]:
    """Measure negative-clause characters without rewarding positive padding."""
    total = len(re.sub(r"\s+", "", prompt))
    negative = 0
    for clause in re.split(r"[。！？!?；;\n]+", prompt):
        compact = re.sub(r"\s+", "", clause)
        if compact and any(marker in compact for marker in NEGATIVE_PROMPT_MARKERS):
            negative += len(compact)
    return (negative / total if total else 0.0), negative, total


def section_between(prompt: str, start: str, end: str) -> str:
    if start not in prompt:
        return ""
    value = prompt.split(start, 1)[1]
    return value.split(end, 1)[0] if end in value else value


def timed_action_text(prompt: str) -> str:
    return "\n".join(match.group(0) for match in TIMED_ACTION_LINE_RE.finditer(prompt))


def phase_is_present(prompt: str, phase: str) -> bool:
    return any(term in prompt for term in EATING_PHASE_GROUPS.get(phase, ()))


def phase_is_positively_present(prompt: str, phase: str) -> bool:
    negative_markers = ("不补", "不写", "不出现", "不需要", "无需", "禁止", "不得", "没有", "未见", "不强制")
    for term in EATING_PHASE_GROUPS.get(phase, ()):
        start = 0
        while True:
            index = prompt.find(term, start)
            if index < 0:
                break
            context = prompt[max(0, index - 12) : index]
            if not any(marker in context for marker in negative_markers):
                return True
            start = index + len(term)
    return False


def spoken_char_count(value: str) -> int:
    """Count Han characters, letters and digits; punctuation is not speech capacity."""
    return len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", value))


def dialogue_chunks(value: str) -> list[str]:
    """Return spoken chunks that may be separated by performance narration."""
    raw_chunks = [normalize_text(part) for part in re.split(r"[，。！？!?；;、…]+", value) if normalize_text(part)]
    chunks: list[str] = []
    for chunk in raw_chunks:
        if len(chunk) > 2 and chunk[0] in "哇诶哎啊欸哦嗯":
            chunks.extend((chunk[0], chunk[1:]))
        else:
            chunks.append(chunk)
    return chunks


def script_lines(block: str) -> list[tuple[str, str]]:
    match = SCRIPT_RE.search(block)
    if not match or match.group(1).strip() == "无":
        return []
    return [(m.group(1).strip(), m.group(2).strip()) for m in DIALOGUE_RE.finditer(match.group(1))]


def has_no_text_rule(prompt: str) -> bool:
    return all(any(term in prompt for term in group) for group in NO_TEXT_GROUPS)


def segment_is_positive_visibility(segment: str, label: str) -> bool:
    if label not in segment:
        return False
    if any(term in segment for term in NEGATIVE_TERMS):
        return False
    return any(term in segment for term in VISIBLE_TERMS) and any(term in segment for term in BODY_TERMS)


def load_role_lock(path: Path | None) -> dict[str, Any] | None:
    if path is None:
        return None
    with path.open(encoding="utf-8-sig") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError("角色锁根节点必须是对象")
    return data


def load_story_plan(path: Path | None) -> dict[str, Any] | None:
    if path is None:
        return None
    with path.open(encoding="utf-8-sig") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError("故事计划根节点必须是对象")
    return data


def load_text_handoff(path: Path | None) -> dict[str, Any] | None:
    if path is None:
        return None
    with path.open(encoding="utf-8-sig") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError("text handoff 根节点必须是对象")
    return data


def resolve_prompt_length_contract(
    story_plan: dict[str, Any] | None,
    enforce_prompt_length: bool | None,
    min_prompt_chars: int | None,
    max_prompt_chars: int | None,
) -> tuple[list[Issue], int, int]:
    """Resolve one explicit on/off contract; partial limits are never accepted."""
    issues: list[Issue] = []
    contract = story_plan.get("prompt_length_contract") if isinstance(story_plan, dict) else None
    if contract is not None and not isinstance(contract, dict):
        return [Issue("PROMPT_LENGTH_CONTRACT_INVALID", "story_plan.prompt_length_contract 必须是对象")], 0, 0

    if isinstance(contract, dict):
        enabled = contract.get("enabled")
        if not isinstance(enabled, bool):
            issues.append(Issue("PROMPT_LENGTH_CONTRACT_INVALID", "prompt_length_contract.enabled 必须是布尔值"))
            return issues, 0, 0
        if enforce_prompt_length is not None and enforce_prompt_length != enabled:
            issues.append(
                Issue(
                    "PROMPT_LENGTH_CONTRACT_CONFLICT",
                    "命令行长度开关与 story_plan.prompt_length_contract.enabled 不一致",
                )
            )
        if enabled:
            minimum = contract.get("minimum_non_whitespace_characters")
            maximum = contract.get("maximum_non_whitespace_characters")
            if min_prompt_chars is not None and min_prompt_chars != minimum:
                issues.append(Issue("PROMPT_LENGTH_CONTRACT_CONFLICT", "命令行 Prompt 下限与 story_plan 不一致"))
            if max_prompt_chars is not None and max_prompt_chars != maximum:
                issues.append(Issue("PROMPT_LENGTH_CONTRACT_CONFLICT", "命令行 Prompt 上限与 story_plan 不一致"))
        else:
            minimum = maximum = 0
            if min_prompt_chars not in (None, 0) or max_prompt_chars not in (None, 0):
                issues.append(Issue("PROMPT_LENGTH_CONTRACT_CONFLICT", "长度契约已关闭，不得单独传入 min/max"))
    else:
        enabled = enforce_prompt_length is True
        if enabled:
            minimum = 3000 if min_prompt_chars is None else min_prompt_chars
            maximum = 4000 if max_prompt_chars is None else max_prompt_chars
        else:
            minimum = maximum = 0
            if min_prompt_chars not in (None, 0) or max_prompt_chars not in (None, 0):
                issues.append(
                    Issue(
                        "PROMPT_LENGTH_CONTRACT_INVALID",
                        "传入 min/max 前必须显式启用 --enforce-prompt-length；默认不强制长度",
                    )
                )

    if enabled:
        if (
            not isinstance(minimum, int)
            or isinstance(minimum, bool)
            or not isinstance(maximum, int)
            or isinstance(maximum, bool)
            or minimum <= 0
            or maximum <= 0
            or maximum < minimum
        ):
            issues.append(
                Issue(
                    "PROMPT_LENGTH_CONTRACT_INVALID",
                    "长度契约启用时必须同时提供正数下限和上限，且 max>=min",
                )
            )
            return issues, 0, 0
        return issues, int(minimum), int(maximum)
    return issues, 0, 0


def lint_stage_inputs(
    stage: str | None,
    story_plan_path: Path | None,
    text_handoff_path: Path | None,
    delivery_dir: Path | None,
) -> list[Issue]:
    if stage is None:
        return []
    issues: list[Issue] = []
    if stage == "text_branch":
        if story_plan_path is None:
            issues.append(Issue("LINT_STAGE_INPUT_MISSING", "text_branch 必须传 --story-plan"))
        if text_handoff_path is None:
            issues.append(Issue("LINT_STAGE_INPUT_MISSING", "text_branch 必须传 --text-handoff"))
        if delivery_dir is not None:
            issues.append(Issue("LINT_STAGE_INPUT_FORBIDDEN", "text_branch 不得检查或交付 DOCX 目录"))
    elif stage == "full_delivery_precompile":
        if story_plan_path is None:
            issues.append(Issue("LINT_STAGE_INPUT_MISSING", "full_delivery_precompile 必须传 --story-plan"))
        if text_handoff_path is not None:
            issues.append(Issue("LINT_STAGE_INPUT_FORBIDDEN", "full_delivery_precompile 只读归一化 story plan，不再依赖 text handoff"))
        if delivery_dir is not None:
            issues.append(Issue("LINT_STAGE_INPUT_FORBIDDEN", "full_delivery_precompile 尚未导出 DOCX，不得传 --delivery-dir"))
    elif stage == "full_delivery_postexport":
        if story_plan_path is None:
            issues.append(Issue("LINT_STAGE_INPUT_MISSING", "full_delivery_postexport 必须传 --story-plan"))
        if delivery_dir is None:
            issues.append(Issue("LINT_STAGE_INPUT_MISSING", "full_delivery_postexport 必须传 --delivery-dir"))
        if text_handoff_path is not None:
            issues.append(Issue("LINT_STAGE_INPUT_FORBIDDEN", "full_delivery_postexport 只验收归一化终稿，不再依赖 text handoff"))
    else:
        issues.append(Issue("LINT_STAGE_INVALID", f"未知 lint stage：{stage}"))
    return issues


def lint_text_handoff_against_story_plan(
    handoff: dict[str, Any],
    story_plan: dict[str, Any],
) -> list[Issue]:
    """Prove the versioned handoff carries the same locked SRC/ADD/S sets."""
    issues: list[Issue] = []
    inventory = story_plan.get("source_shot_inventory")
    plan_map = story_plan.get("generation_shot_map")
    if not isinstance(inventory, list) or not isinstance(plan_map, list):
        return issues
    plan_sources = {
        str(item.get("source_shot_id", "")).upper()
        for item in inventory
        if isinstance(item, dict) and str(item.get("source_shot_id", "")).strip()
    }
    handoff_sources = {
        str(item.get("source_shot_id", "")).upper()
        for item in handoff.get("source_units", [])
        if isinstance(item, dict) and str(item.get("source_shot_id", "")).strip()
    }
    plan_shots = {
        str(item.get("shot_id", ""))
        for item in plan_map
        if isinstance(item, dict) and str(item.get("shot_id", "")).strip()
    }
    handoff_shots = {
        str(item.get("shot_id", ""))
        for item in handoff.get("generation_shot_map", [])
        if isinstance(item, dict) and str(item.get("shot_id", "")).strip()
    }
    plan_inserted: set[str] = set()
    for item in plan_map:
        if not isinstance(item, dict):
            continue
        raw_inserted = item.get("inserted_shot_ids")
        if isinstance(raw_inserted, list):
            plan_inserted.update(str(value).upper() for value in raw_inserted if str(value).strip())
        elif str(item.get("inserted_shot_id", "")).strip():
            plan_inserted.add(str(item["inserted_shot_id"]).upper())
    handoff_inserted = {
        str(item.get("inserted_shot_id", "")).upper()
        for item in handoff.get("inserted_units", [])
        if isinstance(item, dict) and str(item.get("inserted_shot_id", "")).strip()
    }
    for label, expected, actual in (
        ("SRC", plan_sources, handoff_sources),
        ("ADD", plan_inserted, handoff_inserted),
        ("S", plan_shots, handoff_shots),
    ):
        if expected != actual:
            issues.append(
                Issue(
                    "TEXT_HANDOFF_STORY_PLAN_MISMATCH",
                    f"text handoff 与 story_plan 的 {label} 集合不一致；缺{sorted(expected - actual)}，多{sorted(actual - expected)}",
                )
            )
    return issues


def lint_shot(
    title: str,
    block: str,
    min_prompt_chars: int = 0,
    max_prompt_chars: int = 0,
) -> tuple[list[Issue], str, list[tuple[str, str]]]:
    sid = shot_id(title)
    issues: list[Issue] = []
    required = (
        "原片时间：",
        "源分镜ID：",
        "独立生成时长：",
        "分镜描述：",
        "人物位置：",
        "声音方式：",
        "产品形态：",
        "生成首帧：",
        "分镜图：",
        "核心主体：",
        "核心动作：",
        "核心产品：",
        "适用表演层：",
        "【口播稿】",
        "【原片动作对应】",
        "【内容审核记录】",
    )
    for marker in required:
        if marker not in block:
            issues.append(Issue("TXT_EXPORT_MISSING", f"{sid} 缺少 {marker}"))

    prompt_match = PROMPT_RE.search(block)
    if not prompt_match:
        issues.append(Issue("TXT_EXPORT_MISSING", f"{sid} 缺少完整 Prompt 主体"))
        return issues, "", script_lines(block)

    prompt = prompt_match.group(1).strip()
    char_count = len(re.sub(r"\s+", "", prompt))
    if min_prompt_chars > 0 and char_count < min_prompt_chars:
        issues.append(
            Issue(
                "PROMPT_TOO_SHORT",
                f"{sid} Prompt 只有 {char_count} 个非空白字符，当前交付契约至少需要 {min_prompt_chars} 个",
            )
        )
    if max_prompt_chars > 0 and char_count > max_prompt_chars:
        issues.append(
            Issue(
                "PROMPT_TOO_LONG",
                f"{sid} Prompt 有 {char_count} 个非空白字符，当前交付契约上限为 {max_prompt_chars} 个",
            )
        )
    count_match = PROMPT_COUNT_RE.search(block)
    if not count_match or int(count_match.group(1)) != char_count:
        stated = count_match.group(1) if count_match else "缺失"
        issues.append(Issue("PROMPT_CHAR_COUNT_MISMATCH", f"{sid} 标注字符数为 {stated}，程序实算为 {char_count}"))
    if not TIME_RE.search(prompt):
        issues.append(Issue("GENERATION_TIME_NOT_ZERO", f"{sid} 未找到从 0.00 秒开始的动作时间"))
    if not has_no_text_rule(prompt):
        issues.append(Issue("NO_TEXT_RULE_MISSING", f"{sid} Prompt 缺少禁止字幕和水印的硬规则"))

    description = field(block, "分镜描述")
    if (
        len(normalize_text(description)) < 12
        or not any(marker in description for marker in ("原片", "SRC"))
        or not any(marker in description for marker in ("新版口播", "新口播"))
    ):
        issues.append(Issue("SHOT_DESCRIPTION_MISSING", f"{sid} 分镜描述必须说明原片/SRC节奏、动作与新版口播落点"))
    declared_source_ids = field(block, "源分镜ID")
    declared_source_time = field(block, "原片时间")
    if declared_source_ids == "新增镜头":
        if declared_source_time != "新增镜头":
            issues.append(Issue("SOURCE_TIMECODE_MAP_MISMATCH", f"{sid} 新增 ADD 不得编造原片秒数，原片时间必须写“新增镜头”"))
    elif source_time_range(block) is None:
        issues.append(Issue("SOURCE_TIMECODE_MAP_MISMATCH", f"{sid} 原片时间缺失或格式无效"))
    if not declared_source_ids:
        issues.append(Issue("SOURCE_SHOT_INVENTORY_MISSING", f"{sid} 源分镜ID字段为空"))
    if not field(block, "分镜图"):
        issues.append(Issue("SOURCE_SHOT_EVIDENCE_MISSING", f"{sid} 分镜图字段为空，未证明每个分镜都已做图"))

    for name in (*CORE_FACT_FIELDS, "适用表演层"):
        value = field(block, name)
        if not value:
            issues.append(Issue("PROMPT_CORE_FACT_MISSING", f"{sid} 缺少{name}字段"))
            continue
        if name in CORE_FACT_FIELDS:
            normalized_prompt = normalize_text(prompt)
            omitted = [token for token in core_fact_tokens(value) if normalize_text(token) not in normalized_prompt]
            if omitted:
                issues.append(
                    Issue(
                        "PROMPT_CORE_FACT_OMITTED",
                        f"{sid} 可复制 Prompt 遗漏{name}：{'、'.join(omitted)}",
                    )
                )

    repeated = repeated_prompt_sentences(prompt)
    if repeated:
        excerpt = repeated[0][:72]
        issues.append(Issue("PROMPT_PADDING_DETECTED", f"{sid} 存在重复长句填充：{excerpt}"))

    missing_headers = [header for header in NARRATIVE_PROMPT_HEADERS if header not in prompt]
    if missing_headers:
        issues.append(
            Issue(
                "NARRATIVE_FORMAT_MISSING",
                f"{sid} 缺少叙事优先 Prompt 结构：{'、'.join(missing_headers)}",
            )
        )
    else:
        positions = [prompt.index(header) for header in NARRATIVE_PROMPT_HEADERS]
        if positions != sorted(positions):
            issues.append(Issue("NARRATIVE_FORMAT_MISSING", f"{sid} narrative-six-layer-v1 各章节顺序错误"))

    narrative = section_between(prompt, "【原片叙事复原】", "【原片逐时动作】")
    person_position = field(block, "人物位置")
    person_shot = bool(person_position) and not any(
        marker in person_position for marker in ("无人物", "纯产品", "静物", "不适用")
    )
    if person_shot:
        missing_reasoning: list[str] = []
        if not any(term in narrative for term in EMOTION_TRIGGER_TERMS):
            missing_reasoning.append("可见/可听触发")
        if not any(term in narrative for term in CHARACTER_INTENTION_TERMS):
            missing_reasoning.append("人物意图")
        if not any(term in narrative for term in EMOTION_ARC_TERMS):
            missing_reasoning.append("情绪转折")
        if not any(term in narrative for term in NARRATIVE_PAYOFF_TERMS):
            missing_reasoning.append("叙事落点")
        if missing_reasoning:
            issues.append(
                Issue(
                    "EMOTIONAL_CAUSALITY_MISSING",
                    f"{sid} 原片叙事复原缺少{'、'.join(missing_reasoning)}；六层不能只罗列表情器官",
                )
            )

    leaked_audit_tokens = [token for token in SIX_LAYER_AUDIT_TOKENS if token in prompt]
    if leaked_audit_tokens:
        issues.append(
            Issue(
                "SIX_LAYER_AUDIT_LEAKED_INTO_PROMPT",
                f"{sid} 把内部六层审计字段写入可复制 Prompt：{'、'.join(leaked_audit_tokens)}",
            )
        )

    negative_ratio, negative_chars, total_chars = negative_prompt_ratio(prompt)
    if negative_ratio > 0.15:
        issues.append(
            Issue(
                "NEGATIVE_CONSTRAINT_OVERLOAD",
                f"{sid} 限制句占 {negative_chars}/{total_chars}={negative_ratio:.1%}，超过15%；请改写为原片正向叙事和可拍摄动作，并压缩末尾纠错附录",
            )
        )

    timed_text = timed_action_text(prompt)
    for pattern in MOUTH_SPEECH_CONFLICT_PATTERNS:
        conflict = pattern.search(timed_text)
        if conflict:
            issues.append(
                Issue(
                    "CHEWING_SPEECH_CONFLICT",
                    f"{sid} 咬合/闭口咀嚼与吃食者屏内完整口播重叠：{conflict.group(0)[:80]}",
                )
            )
            break

    product_state = field(block, "产品形态")
    if INTERNAL_PRODUCT_RE.search(product_state):
        issues.append(Issue("PRODUCT_STATE_LABEL_NOT_CHINESE", f"{sid} 产品形态使用了内部英文或字母标签：{product_state}"))

    is_butter_crisp = any(term in product_state or term in prompt for term in BUTTER_CRISP_PRODUCT_TERMS)
    scale_mode = field(block, "尺度模式")
    has_bare_product = is_butter_crisp and any(
        term in product_state for term in BUTTER_CRISP_BARE_STATE_TERMS
    )
    has_retail_box = is_butter_crisp and any(
        term in product_state for term in BUTTER_CRISP_BOX_STATE_TERMS
    )
    if is_butter_crisp:
        if scale_mode not in {"physical_consistency", "relative_pixel_resize"}:
            issues.append(
                Issue(
                    "SCALE_MODE_UNDECLARED",
                    f"{sid} 黄油脆丝棒镜头必须声明唯一尺度模式：physical_consistency 或 relative_pixel_resize",
                )
            )
        if scale_mode == "physical_consistency" and has_bare_product:
            required_scale_groups = {
                "12 × 2.5 × 1 cm 单根三轴尺寸": BUTTER_CRISP_PRODUCT_DIMENSION_TERMS,
                "4.8:1 正面目标长宽比": BUTTER_CRISP_PRODUCT_TARGET_ASPECT_TERMS,
                "4:1–5:1 棒体长宽比": BUTTER_CRISP_PRODUCT_ASPECT_TERMS,
            }
            for label, words in required_scale_groups.items():
                if not any(word in prompt for word in words):
                    issues.append(
                        Issue(
                            "PRODUCT_PACKAGE_SCALE_LOCK_MISSING",
                            f"{sid} physical_consistency 缺少{label}硬约束",
                        )
                    )
        if scale_mode == "relative_pixel_resize":
            if any(term in prompt for term in BUTTER_CRISP_PRODUCT_PACKAGE_RATIO_TERMS):
                issues.append(
                    Issue(
                        "SCALE_MODE_COLLISION",
                        f"{sid} relative_pixel_resize 不得同时把 12:15 或 0.80 写入本轮画面尺寸硬约束",
                    )
                )
    if has_bare_product:
        for label, words in BUTTER_CRISP_MATERIAL_GROUPS.items():
            if not any(word in prompt for word in words):
                issues.append(
                    Issue(
                        "PRODUCT_MICROSTRUCTURE_RULE_MISSING",
                        f"{sid} 黄油脆丝棒裸产品 Prompt 缺少{label}规则，可能退化成表面图案",
                    )
                )
    is_butter_break = is_butter_crisp and any(
        term in product_state or term in prompt for term in BUTTER_CRISP_BREAK_STATE_TERMS
    )
    if is_butter_break:
        projection_mode = field(block, "投影模式")
        projection_source = field(block, "投影事实源")
        if projection_mode not in BUTTER_CRISP_PROJECTION_MODES:
            issues.append(
                Issue(
                    "PROJECTION_MODE_UNDECLARED",
                    f"{sid} 黄油脆丝棒掰断镜必须声明 source_pixel_lock 或 intentional_depth_move",
                )
            )
        if not projection_source:
            issues.append(Issue("PROJECTION_SOURCE_UNDECLARED", f"{sid} 黄油脆丝棒掰断镜缺少原视频投影事实源"))
        elif any(term.lower() in projection_source.lower() for term in BUTTER_CRISP_FAILED_PROJECTION_SOURCE_TERMS):
            issues.append(
                Issue(
                    "FAILED_FRAME_USED_AS_PROJECTION_SOURCE",
                    f"{sid} 投影事实源指向失败/未批准生成图；该类图片只能用于诊断",
                )
            )
        if projection_mode == "source_pixel_lock":
            for code, (label, words) in BUTTER_CRISP_SOURCE_PIXEL_LOCKS.items():
                if not any(word in prompt for word in words):
                    issues.append(Issue(code, f"{sid} source_pixel_lock 缺少{label}硬约束"))
            used_axis_terms = [term for term in BUTTER_CRISP_CAMERA_AXIS_ADVANCE_TERMS if term in prompt]
            if used_axis_terms:
                issues.append(
                    Issue(
                        "UNAUTHORIZED_CAMERA_AXIS_ADVANCE",
                        f"{sid} source_pixel_lock 禁止沿相机轴前推：{'、'.join(used_axis_terms)}",
                    )
                )
        for code, (label, words) in BUTTER_CRISP_BREAK_LOCKS.items():
            if not any(word in prompt for word in words):
                issues.append(Issue(code, f"{sid} 黄油脆丝棒掰断镜缺少{label}硬约束"))
    if has_retail_box:
        if not any(term in prompt for term in BUTTER_CRISP_BOX_DIMENSION_TERMS):
            issues.append(Issue("PACKAGE_DIMENSION_LOCK_MISSING", f"{sid} 黄油脆丝棒外盒缺少 15×15×4.5 cm 尺寸锁"))
        if not any(term in prompt for term in BUTTER_CRISP_BOX_RATIO_TERMS):
            issues.append(Issue("PACKAGE_DIMENSION_LOCK_MISSING", f"{sid} 黄油脆丝棒外盒缺少 1:1 正方形正面约束"))
        if not any(term in prompt for term in BUTTER_CRISP_BOX_DEPTH_TERMS):
            issues.append(Issue("PACKAGE_DIMENSION_LOCK_MISSING", f"{sid} 黄油脆丝棒外盒缺少约 0.3 边长盒厚约束"))

    status = STATUS_RE.search(block)
    if status:
        issues.append(
            Issue(
                "STRUCTURE_RESULT_MISREPRESENTED_AS_CONTENT_AUDIT",
                f"{sid} 使用了状态词 {status.group(1)}；请改写为具体人物、台词、动作或待看像素事实",
            )
        )

    lines = script_lines(block)
    duration_match = DURATION_RE.search(block)
    if not duration_match:
        issues.append(Issue("PACING_FIELDS_MISSING", f"{sid} 缺少可解析的独立生成时长"))

    segment_match = SEGMENT_COUNT_RE.search(block)
    if not segment_match:
        issues.append(Issue("PACING_FIELDS_MISSING", f"{sid} 缺少本镜句段数统计"))
    else:
        stated_segments = int(segment_match.group(1))
        if stated_segments != len(lines):
            issues.append(Issue("PACING_FIELDS_MISSING", f"{sid} 标注{stated_segments}个句段，口播稿实算{len(lines)}个"))
    if len(lines) > 3:
        issues.append(Issue("SCRIPT_SEGMENT_OVERLOAD", f"{sid} 同镜包含{len(lines)}个台词句段，超过3个"))

    if lines:
        speech_seconds_match = SPEECH_SECONDS_RE.search(block)
        speech_rate_match = SPEECH_RATE_RE.search(block)
        if not speech_seconds_match or not speech_rate_match:
            issues.append(Issue("PACING_FIELDS_MISSING", f"{sid} 有口播但缺实际可说时段或计划语速统计"))
        else:
            speech_seconds = float(speech_seconds_match.group(1))
            stated_rate = float(speech_rate_match.group(1))
            if speech_seconds <= 0:
                issues.append(Issue("SPEECH_WINDOW_INVALID", f"{sid} 实际可说时段必须大于0秒"))
            else:
                effective_chars = sum(spoken_char_count(text) for _, text in lines)
                calculated_rate = effective_chars / speech_seconds
                if abs(stated_rate - calculated_rate) > 0.06:
                    issues.append(Issue("PACING_FIELDS_MISSING", f"{sid} 标注语速{stated_rate:.2f}字/秒，按{effective_chars}字/{speech_seconds:.2f}秒实算{calculated_rate:.2f}字/秒"))
                delivery = field(block, "声音方式")
                on_screen = "on_screen_speech" in delivery or "屏内" in delivery
                threshold = 5.0 if on_screen else 5.5
                if calculated_rate > threshold + 0.01:
                    mode = "屏内口播" if on_screen else "画外音"
                    issues.append(Issue("SPEECH_RATE_EXCEEDED", f"{sid} {mode}{calculated_rate:.2f}字/秒，超过{threshold:.2f}字/秒上限"))
    for phrase in GENERIC_SPEECH_PLACEHOLDERS:
        if phrase in prompt:
            issues.append(
                Issue(
                    "GENERIC_SPEECH_PLACEHOLDER",
                    f"{sid} 使用了泛化讲话占位表达“{phrase}”；请改写为具体口音、语音和节奏方案",
                )
            )
    if lines and not any(term in prompt for term in ACCENT_SPECIFIC_TERMS):
        issues.append(Issue("ACCENT_PLAN_MISSING", f"{sid} 有口播，但 Prompt 没有具体地域口语方向或语音特征"))

    performance_groups = dict(PERFORMANCE_GROUPS)
    if not lines:
        performance_groups["无口播镜头的现场声音"] = ("声音", "环境底噪", "摩擦", "呼吸", "轻响")
        performance_groups.pop("讲话细节", None)
    for label, words in performance_groups.items():
        if not any(word in prompt for word in words):
            issues.append(Issue("PERFORMANCE_DETAIL_MISSING", f"{sid} 缺少可观察的{label}"))

    for label, text in lines:
        normalized_prompt = normalize_text(prompt)
        if not all(chunk in normalized_prompt for chunk in dialogue_chunks(text)):
            issues.append(Issue("SCRIPT_OMITTED", f"{sid} 的完整 Prompt 未叙事性写入 {label} 台词：{text}"))

    action_section = block.split("【原片动作对应】", 1)[-1].split("【内容审核记录】", 1)[0]
    if not re.search(r"\d{1,2}:\d{2}(?:\.\d+)?", action_section) or not any(
        term in action_section for term in ("对应", "复刻", "映射")
    ):
        issues.append(Issue("SOURCE_ACTION_OMITTED", f"{sid} 原片动作对应未写可核对的时间与映射事实"))

    audit_section = block.split("【内容审核记录】", 1)[-1]
    if not any(term in audit_section for term in ("人物", "角色")):
        issues.append(Issue("TXT_EXPORT_MISSING", f"{sid} 内容审核记录缺人物事实"))
    if lines and not any(term in audit_section for term in ("台词", "口型", "声音", "声源")):
        issues.append(Issue("TXT_EXPORT_MISSING", f"{sid} 内容审核记录缺台词/声源事实"))

    return issues, prompt, lines


def lint_role_lock(
    shots: list[tuple[str, str]],
    prompts: dict[str, str],
    all_script_lines: list[tuple[str, str, str]],
    role_lock: dict[str, Any],
) -> list[Issue]:
    issues: list[Issue] = []
    characters = role_lock.get("characters")
    dialogue = role_lock.get("dialogue")
    if not isinstance(characters, dict) or not isinstance(dialogue, list):
        return [Issue("SPEAKER_IDENTITY_UNCONFIRMED", "角色锁必须包含 characters 对象和 dialogue 数组")]

    if dialogue:
        speech_plan = role_lock.get("speech_plan")
        if not isinstance(speech_plan, dict) or not str(speech_plan.get("summary", "")).strip():
            issues.append(Issue("ACCENT_PLAN_MISSING", "角色锁有台词，但缺少 speech_plan 具体讲话方案"))
        else:
            summary = str(speech_plan.get("summary", ""))
            for phrase in GENERIC_SPEECH_PLACEHOLDERS:
                if phrase in summary:
                    issues.append(Issue("GENERIC_SPEECH_PLACEHOLDER", f"角色锁 speech_plan 使用了泛化占位表达“{phrase}”"))
            if not any(term in summary for term in ACCENT_SPECIFIC_TERMS):
                issues.append(Issue("ACCENT_PLAN_MISSING", "角色锁 speech_plan 缺具体地域口语方向或语音特征"))
            if speech_plan.get("source") == "creative_proposal" and speech_plan.get("disclosed_to_user") is not True:
                issues.append(Issue("ACCENT_PROPOSAL_NOT_DISCLOSED", "口音来源是创作提案，但角色锁没有记录已在对话中告知用户"))

    delivered = {(normalize_text(label), normalize_text(text)): sid for sid, label, text in all_script_lines}
    expected_texts: dict[str, str] = {}
    for item in dialogue:
        if not isinstance(item, dict):
            continue
        label = str(item.get("label", "")).strip()
        text = str(item.get("text", "")).strip()
        if not label or not text:
            issues.append(Issue("SPEAKER_IDENTITY_UNCONFIRMED", "角色锁 dialogue 条目缺 label 或 text"))
            continue
        key = (normalize_text(label), normalize_text(text))
        if key not in delivered:
            wrong = [actual_label for (actual_label, actual_text) in delivered if actual_text == normalize_text(text)]
            if wrong:
                issues.append(Issue("DIALOGUE_SPEAKER_MISMATCH", f"台词“{text}”应由{label}说，TXT 中标成了{wrong[0]}"))
            else:
                issues.append(Issue("SCRIPT_OMITTED", f"角色锁台词未出现在任何分镜口播稿中：{label}“{text}”"))
        expected_texts[normalize_text(text)] = normalize_text(label)

    for sid, actual_label, text in all_script_lines:
        expected_label = expected_texts.get(normalize_text(text))
        if expected_label is not None and normalize_text(actual_label) != expected_label:
            issues.append(Issue("DIALOGUE_SPEAKER_MISMATCH", f"{sid} 台词“{text}”说话人应为{expected_label}，实际为{actual_label}"))

    for character in characters.values():
        if not isinstance(character, dict):
            continue
        label = str(character.get("label", "")).strip()
        visibility = character.get("visibility")
        if not label:
            issues.append(Issue("SPEAKER_IDENTITY_UNCONFIRMED", "角色锁 characters 条目缺 label"))
            continue
        for title, block in shots:
            sid = shot_id(title)
            position = field(block, "人物位置")
            prompt = prompts.get(sid, "")
            if visibility == "offscreen_all":
                if label not in position or not any(term in position for term in ("镜头后", "镜外", "不出镜")):
                    issues.append(Issue("OFFSCREEN_CHARACTER_VISIBLE", f"{sid} 人物位置没有继承“{label}全程镜外”"))
                for segment in re.split(r"[。！？!?；;\n]", prompt):
                    if segment_is_positive_visibility(segment, label):
                        excerpt = segment.strip()[:80]
                        issues.append(Issue("OFFSCREEN_CHARACTER_VISIBLE", f"{sid} 安排全程镜外的{label}出镜：{excerpt}"))
                        break
            elif visibility == "onscreen_all":
                if label not in position or not any(term in position for term in ("画面中", "镜头中", "在画面")):
                    issues.append(Issue("SPEAKER_IDENTITY_CONFLICT", f"{sid} 人物位置没有继承“{label}全程画面中”"))

    for item in role_lock.get("required_actions", []):
        if not isinstance(item, dict):
            continue
        sid = str(item.get("shot_id", ""))
        kind = item.get("kind")
        prompt = prompts.get(sid, "")
        if kind == "bite_chain":
            for label, words in BITE_CHAIN_GROUPS.items():
                if not any(word in prompt for word in words):
                    issues.append(Issue("SOURCE_ACTION_OMITTED", f"{sid} 原片含咬食，但 Prompt 缺少咬食链环节：{label}"))

    return issues


def numeric_shot_index(sid: str) -> int | None:
    match = re.fullmatch(r"S(\d{3})", sid)
    return int(match.group(1)) if match else None


def lint_story_plan(
    shots: list[tuple[str, str]],
    prompts: dict[str, str],
    story_plan: dict[str, Any],
) -> list[Issue]:
    """Check source-shot preservation, timing policy and eating-event planning."""
    issues: list[Issue] = []
    blocks = {shot_id(title): block for title, block in shots}

    raw_source_duration = story_plan.get("source_duration_seconds")
    try:
        source_duration = float(raw_source_duration)
        if source_duration <= 0:
            raise ValueError
    except (TypeError, ValueError):
        source_duration = 0.0
        issues.append(Issue("PACING_FIELDS_MISSING", "story_plan 缺合法 source_duration_seconds"))

    policy = story_plan.get("generation_time_policy")
    min_duration = max_duration = 0.0
    if not isinstance(policy, dict):
        issues.append(Issue("PACING_FIELDS_MISSING", "story_plan 缺 generation_time_policy 对象"))
    else:
        try:
            min_duration = float(policy.get("min_duration_seconds"))
            max_duration = float(policy.get("max_duration_seconds"))
            if min_duration <= 0 or max_duration < min_duration:
                raise ValueError
        except (TypeError, ValueError):
            min_duration = max_duration = 0.0
            issues.append(Issue("PACING_FIELDS_MISSING", "generation_time_policy 缺合法生成镜最短/最长值"))
        for name in ("onscreen_speech_max_chars_per_second", "voiceover_max_chars_per_second"):
            try:
                if float(policy.get(name)) <= 0:
                    raise ValueError
            except (TypeError, ValueError):
                issues.append(Issue("PACING_FIELDS_MISSING", f"generation_time_policy 缺合法 {name}"))

    inventory = story_plan.get("source_shot_inventory")
    inventory_by_id: dict[str, dict[str, Any]] = {}
    inventory_order: list[str] = []
    if not isinstance(inventory, list) or not inventory:
        issues.append(Issue("SOURCE_SHOT_INVENTORY_MISSING", "story_plan 缺非空 source_shot_inventory"))
        inventory = []
    for position, item in enumerate(inventory):
        if not isinstance(item, dict):
            issues.append(Issue("SOURCE_SHOT_EVIDENCE_MISSING", f"source_shot_inventory 第{position + 1}项不是对象"))
            continue
        source_id = str(item.get("source_shot_id", "")).strip()
        if not re.fullmatch(r"SRC\d+", source_id, flags=re.I):
            issues.append(Issue("SOURCE_SHOT_INVENTORY_MISSING", f"source_shot_inventory 第{position + 1}项缺合法 source_shot_id"))
            continue
        source_id = source_id.upper()
        if source_id in inventory_by_id:
            issues.append(Issue("SOURCE_SHOT_INVENTORY_MISSING", f"源分镜 ID 重复：{source_id}"))
            continue
        normalized_item = dict(item)
        normalized_item["source_shot_id"] = source_id
        inventory_by_id[source_id] = normalized_item
        inventory_order.append(source_id)
        missing = [
            name
            for name in ("source_start", "source_end", "action", "source_first_frame", "approved_delivery_image")
            if not str(item.get(name, "")).strip()
        ]
        if missing:
            issues.append(Issue("SOURCE_SHOT_EVIDENCE_MISSING", f"{source_id} 缺字段：{'、'.join(missing)}"))
        try:
            start = parse_timecode(str(item.get("source_start", "")))
            end = parse_timecode(str(item.get("source_end", "")))
            declared_duration = float(item.get("duration_seconds"))
            if end <= start or abs((end - start) - declared_duration) > 0.002:
                raise ValueError
            normalized_item["_start_seconds"] = start
            normalized_item["_end_seconds"] = end
            normalized_item["_duration_seconds"] = declared_duration
        except (TypeError, ValueError):
            issues.append(Issue("SOURCE_SHOT_EVIDENCE_MISSING", f"{source_id} 时间码或 duration_seconds 不合法/不一致"))

    timed_inventory = [inventory_by_id[source_id] for source_id in inventory_order]
    if timed_inventory and all("_start_seconds" in item for item in timed_inventory):
        if abs(float(timed_inventory[0]["_start_seconds"])) > 0.002:
            issues.append(Issue("SOURCE_SHOT_COVERAGE_INCOMPLETE", "source_shot_inventory 没有从原片 00:00.000 开始"))
        for previous, current in zip(timed_inventory, timed_inventory[1:]):
            if abs(float(previous["_end_seconds"]) - float(current["_start_seconds"])) > 0.002:
                issues.append(
                    Issue(
                        "SOURCE_SHOT_COVERAGE_INCOMPLETE",
                        f"{previous['source_shot_id']} 与 {current['source_shot_id']} 之间存在缺口、重叠或顺序错误",
                    )
                )
        if source_duration > 0 and abs(float(timed_inventory[-1]["_end_seconds"]) - source_duration) > 0.002:
            issues.append(Issue("SOURCE_SHOT_COVERAGE_INCOMPLETE", "source_shot_inventory 末尾没有覆盖 source_duration_seconds"))

    generation_map = story_plan.get("generation_shot_map")
    if not isinstance(generation_map, list) or not generation_map:
        issues.append(Issue("SOURCE_SHOT_COVERAGE_INCOMPLETE", "story_plan 缺非空 generation_shot_map"))
        generation_map = []
    coverage: dict[str, list[str]] = {source_id: [] for source_id in inventory_by_id}
    mapped_shots: set[str] = set()
    mapping_sources_by_sid: dict[str, list[str]] = {}
    mapping_origin_by_sid: dict[str, str] = {}
    mapping_inserted_by_sid: dict[str, set[str]] = {}
    seen_inserted_ids: set[str] = set()
    inventory_index = {source_id: index for index, source_id in enumerate(inventory_order)}
    for position, item in enumerate(generation_map):
        if not isinstance(item, dict):
            issues.append(Issue("SOURCE_SHOT_COVERAGE_INCOMPLETE", f"generation_shot_map 第{position + 1}项不是对象"))
            continue
        sid = str(item.get("shot_id", "")).strip()
        origin = str(item.get("origin", "")).strip()
        if sid not in blocks:
            issues.append(Issue("SOURCE_SHOT_COVERAGE_INCOMPLETE", f"generation_shot_map 引用了不存在的 TXT 镜头：{sid or '空'}"))
            continue
        if sid in mapped_shots:
            issues.append(Issue("SOURCE_SHOT_COVERAGE_INCOMPLETE", f"生成镜重复映射：{sid}"))
        mapped_shots.add(sid)
        block = blocks[sid]
        raw_ids = item.get("source_shot_ids")
        mapped_ids = [str(value).upper() for value in raw_ids] if isinstance(raw_ids, list) else []
        mapping_sources_by_sid[sid] = mapped_ids
        mapping_origin_by_sid[sid] = origin
        raw_inserted_ids = item.get("inserted_shot_ids")
        if isinstance(raw_inserted_ids, list):
            inserted_ids = [str(value).strip().upper() for value in raw_inserted_ids if str(value).strip()]
        else:
            legacy_inserted_id = str(item.get("inserted_shot_id", "")).strip().upper()
            inserted_ids = [legacy_inserted_id] if legacy_inserted_id else []
        inserted = bool(inserted_ids)
        pure_inserted = inserted and not mapped_ids
        if not mapped_ids and not inserted_ids:
            issues.append(Issue("SOURCE_SHOT_COVERAGE_INCOMPLETE", f"{sid} 没有绑定任何 SRC/ADD"))
        if origin.startswith("inserted") and not inserted_ids:
            issues.append(Issue("INSERTED_SHOT_ID_INVALID", f"{sid} 标为新增镜头却没有 inserted_shot_ids"))
        if inserted:
            if not inserted_ids:
                issues.append(Issue("INSERTED_SHOT_ID_INVALID", f"{sid} 新增镜头必须绑定至少一个唯一 ADD数字 inserted_shot_ids"))
            valid_inserted: set[str] = set()
            for inserted_id in inserted_ids:
                if not re.fullmatch(r"ADD\d+", inserted_id) or inserted_id in seen_inserted_ids:
                    issues.append(Issue("INSERTED_SHOT_ID_INVALID", f"{sid} 含非法或重复 inserted_shot_id={inserted_id or '空'}"))
                else:
                    seen_inserted_ids.add(inserted_id)
                    valid_inserted.add(inserted_id)
            mapping_inserted_by_sid[sid] = valid_inserted

        txt_ids = [value.upper() for value in source_shot_ids(block)]
        if pure_inserted:
            if field(block, "源分镜ID") != "新增镜头":
                issues.append(Issue("SOURCE_TIMECODE_MAP_MISMATCH", f"{sid} 新增镜头的源分镜ID必须写“新增镜头”"))
        elif txt_ids != mapped_ids:
            issues.append(
                Issue(
                    "SOURCE_TIMECODE_MAP_MISMATCH",
                    f"{sid} TXT 源分镜ID {txt_ids} 与 story_plan {mapped_ids} 不一致",
                )
            )

        indexes: list[int] = []
        for source_id in mapped_ids:
            if source_id not in inventory_by_id:
                issues.append(Issue("SOURCE_SHOT_COVERAGE_INCOMPLETE", f"{sid} 引用了 inventory 不存在的 {source_id}"))
                continue
            coverage[source_id].append(sid)
            indexes.append(inventory_index[source_id])
        if indexes and indexes != list(range(indexes[0], indexes[0] + len(indexes))):
            issues.append(Issue("NONADJACENT_SOURCE_SHOT_MERGE", f"{sid} 合并的源分镜不相邻或顺序颠倒：{'、'.join(mapped_ids)}"))

        try:
            generation_duration = float(item.get("generation_duration_seconds"))
        except (TypeError, ValueError):
            generation_duration = -1.0
            issues.append(Issue("PACING_FIELDS_MISSING", f"{sid} generation_shot_map 缺合法 generation_duration_seconds"))
        duration_match = DURATION_RE.search(block)
        if duration_match and generation_duration >= 0 and abs(float(duration_match.group(1)) - generation_duration) > 0.01:
            issues.append(Issue("SOURCE_TIMECODE_MAP_MISMATCH", f"{sid} TXT 生成时长与 story_plan 不一致"))
        if min_duration > 0 and generation_duration >= 0:
            if generation_duration + 0.001 < min_duration or generation_duration - 0.001 > max_duration:
                issues.append(
                    Issue(
                        "SHOT_DURATION_OUT_OF_POLICY",
                        f"{sid} 生成时长{generation_duration:.3f}秒不在项目{min_duration:.3f}–{max_duration:.3f}秒范围",
                    )
                )

        short_sources = [
            source_id
            for source_id in mapped_ids
            if float(inventory_by_id[source_id].get("_duration_seconds", min_duration or 4.0)) < (min_duration or 4.0)
        ]
        if short_sources and (len(mapped_ids) < 2 or generation_duration + 0.001 < (min_duration or 4.0)):
            issues.append(
                Issue(
                    "SHORT_SOURCE_SHOT_NOT_MERGED",
                    f"{sid} 含不足{(min_duration or 4.0):.2f}秒源分镜但未与相邻镜合法合并：{'、'.join(short_sources)}",
                )
            )

        if mapped_ids:
            source_items = [inventory_by_id[value] for value in mapped_ids if value in inventory_by_id]
            if source_items and all("_start_seconds" in value for value in source_items):
                expected_start = float(source_items[0]["_start_seconds"])
                expected_end = float(source_items[-1]["_end_seconds"])
                actual_range = source_time_range(block)
                if actual_range is None or abs(actual_range[0] - expected_start) > 0.002 or abs(actual_range[1] - expected_end) > 0.002:
                    issues.append(Issue("SOURCE_TIMECODE_MAP_MISMATCH", f"{sid} 原片时间没有精确覆盖所列源分镜"))
            image_field = field(block, "分镜图")
            missing_image_ids = [source_id for source_id in mapped_ids if source_id not in image_field]
            if missing_image_ids:
                issues.append(Issue("SOURCE_SHOT_EVIDENCE_MISSING", f"{sid} 分镜图字段漏列：{'、'.join(missing_image_ids)}"))

        plan_description = str(item.get("shot_description", "")).strip()
        if not plan_description or normalize_text(plan_description) != normalize_text(field(block, "分镜描述")):
            issues.append(Issue("SHOT_DESCRIPTION_MISSING", f"{sid} TXT 分镜描述缺失或与 story_plan 不一致"))
        if not str(item.get("revised_script_anchor", "")).strip():
            issues.append(Issue("SHOT_DESCRIPTION_MISSING", f"{sid} 缺 revised_script_anchor，无法证明按新版口播重排节奏"))
        if len(mapped_ids) > 1 and not str(item.get("merge_reason", "")).strip():
            issues.append(Issue("SHOT_DESCRIPTION_MISSING", f"{sid} 合并多个源分镜却缺 merge_reason"))

    for source_id, owners in coverage.items():
        if len(owners) != 1:
            state = "未映射" if not owners else f"重复映射到 {'、'.join(owners)}"
            issues.append(Issue("SOURCE_SHOT_COVERAGE_INCOMPLETE", f"{source_id} {state}"))
    for sid in blocks:
        if sid not in mapped_shots:
            issues.append(Issue("SOURCE_SHOT_COVERAGE_INCOMPLETE", f"TXT 镜头 {sid} 未出现在 generation_shot_map"))

    eating_plan = story_plan.get("eating_plan")
    if source_duration >= 30.0 and not isinstance(eating_plan, dict):
        issues.append(Issue("EATING_PLAN_MISSING", "原视频达到30秒但缺 eating_plan"))
        return issues
    if not isinstance(eating_plan, dict):
        return issues
    canonical_occurrences = isinstance(eating_plan.get("occurrences"), list)
    events = eating_plan.get("occurrences") if canonical_occurrences else eating_plan.get("events")
    if not isinstance(events, list):
        if source_duration >= 30.0:
            issues.append(Issue("EATING_PLAN_MISSING", "eating_plan 缺 occurrences 数组"))
        return issues

    source_events: list[dict[str, Any]] = []
    inserted_events: list[dict[str, Any]] = []
    event_positions: list[tuple[int, str, str]] = []
    narrative_sections: set[str] = set()
    seen_event_ids: set[str] = set()
    for position, event in enumerate(events):
        if not isinstance(event, dict):
            issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"eating_plan 第{position + 1}项不是对象"))
            continue
        event_id = str(event.get("id") or event.get("event_id") or "").strip()
        if not event_id or event_id in seen_event_ids:
            issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"吃食事件 ID 缺失或重复：{event_id or '空'}"))
        seen_event_ids.add(event_id)
        origin = str(event.get("origin", "")).strip()
        if origin == "source":
            source_events.append(event)
        elif origin == "inserted":
            inserted_events.append(event)
        else:
            issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"{event_id or position + 1} origin 必须是 source 或 inserted"))
        sid = str(event.get("shot_id") or event.get("generation_shot_id") or "").strip()
        if sid not in prompts:
            issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"{event_id or position + 1} 引用不存在的生成镜 {sid or '空'}"))
            prompt = ""
        else:
            prompt = prompts[sid]
            index = numeric_shot_index(sid)
            if index is not None:
                event_positions.append((index, origin, event_id))
        narrative = str(event.get("narrative_section", "")).strip()
        rhythm = str(event.get("rhythm_rationale") or event.get("rhythm_anchor") or "").strip()
        script_anchor = str(event.get("revised_script_anchor", "")).strip()
        if narrative:
            narrative_sections.add(narrative)
        if not narrative or not rhythm or not script_anchor:
            issues.append(
                Issue(
                    "EATING_RHYTHM_ANCHOR_MISSING",
                    f"{event_id or sid} 缺 narrative_section、rhythm_anchor 或 revised_script_anchor",
                )
            )
        phases = event.get("required_phases")
        if not isinstance(phases, list) or not phases:
            issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"{event_id or sid} 缺 required_phases"))
            phases = []
        invalid_phases = [str(phase) for phase in phases if phase not in ALLOWED_EATING_PHASES]
        if invalid_phases:
            issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"{event_id or sid} 含非法阶段：{'、'.join(invalid_phases)}"))
        for phase in phases:
            if phase in ALLOWED_EATING_PHASES and not phase_is_present(prompt, str(phase)):
                issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"{sid} Prompt 遗漏吃食阶段 {phase}"))
        speech_after_bite = event.get("speech_after_bite")
        if canonical_occurrences:
            if not isinstance(speech_after_bite, dict):
                issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"{event_id or sid} 缺 speech_after_bite"))
        elif not str(event.get("speech_resume_after", "")).strip():
            issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"{event_id or sid} 缺 speech_resume_after"))
        raw_event_source_ids = event.get("source_shot_ids") or event.get("source_shot_id") or []
        if isinstance(raw_event_source_ids, str):
            raw_event_source_ids = [raw_event_source_ids]
        event_source_ids = [str(value).upper() for value in raw_event_source_ids] if isinstance(raw_event_source_ids, list) else []
        if origin == "source" and not event_source_ids:
            issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"{event_id or sid} 原片吃食事件缺 source_shot_ids"))
        if origin == "inserted" and event_source_ids:
            issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"{event_id or sid} 新增吃食事件不得冒充源分镜"))
        if origin == "source" and any(
            source_id not in mapping_sources_by_sid.get(sid, []) for source_id in event_source_ids
        ):
            issues.append(Issue("EATING_PHASE_EVIDENCE_MISSING", f"{event_id or sid} 的源吃食分镜没有映射到 {sid}"))
        if origin == "inserted":
            event_inserted_id = str(event.get("inserted_shot_id", "")).strip().upper()
            if not event_inserted_id or event_inserted_id not in mapping_inserted_by_sid.get(sid, set()):
                issues.append(Issue("INSERTED_EATING_UNIT_MISMATCH", f"{event_id or sid} 必须绑定该生成镜的准确 ADD inserted_shot_id"))
        if origin == "source":
            action_text = timed_action_text(prompt)
            for optional_phase in ("swallow", "post_eating_reaction"):
                if optional_phase not in phases and phase_is_positively_present(action_text, optional_phase):
                    issues.append(
                        Issue(
                            "EATING_TEMPLATE_PHASE_FORCED",
                            f"{sid} 原片 required_phases 没有 {optional_phase}，却在时间动作中强制补写",
                        )
                    )

    if source_duration >= 30.0:
        expected_insertions = max(0, 3 - len(source_events))
        expected_target = max(len(source_events), 3)
        if len(events) != expected_target or len(inserted_events) < expected_insertions:
            issues.append(Issue("EATING_SHOT_QUOTA_MISSING", f"原片有{len(source_events)}次吃食，最终必须恰为{expected_target}次"))
        if len(inserted_events) != expected_insertions:
            issues.append(
                Issue(
                    "UNNECESSARY_EATING_SHOT_INSERTION",
                    f"原片有{len(source_events)}次吃食，只应补{expected_insertions}次，实际补{len(inserted_events)}次",
                )
            )
        if canonical_occurrences:
            count_expectations = {
                "source_eating_occurrence_count": len(source_events),
                "inserted_eating_occurrence_count": len(inserted_events),
                "target_eating_occurrence_count": len(events),
            }
            for name, actual in count_expectations.items():
                value = eating_plan.get(name)
                if not isinstance(value, int) or isinstance(value, bool) or value != actual:
                    issues.append(Issue("EATING_PLAN_MISSING", f"{name} 必须等于 occurrences 实际计数 {actual}"))
            for left, right in zip(sorted(event_positions), sorted(event_positions)[1:]):
                if right[0] - left[0] <= 1 and "inserted" in {left[1], right[1]}:
                    issues.append(Issue("EATING_EVENTS_NOT_DISTRIBUTED", f"新增吃食 {left[2]} 与 {right[2]} 之间缺非吃食节奏拍"))
        else:
            minimum = eating_plan.get("minimum_nonconsecutive_events_for_30s")
            if minimum != 3:
                issues.append(Issue("EATING_PLAN_MISSING", "旧版吃食计划必须声明 minimum_nonconsecutive_events_for_30s=3"))
            ordered_indexes = [item[0] for item in sorted(event_positions)]
            if len(set(ordered_indexes)) != len(ordered_indexes) or any(
                right - left <= 1 for left, right in zip(ordered_indexes, ordered_indexes[1:])
            ):
                issues.append(Issue("EATING_EVENTS_NOT_DISTRIBUTED", "旧版吃食事件使用了相同或相邻生成镜 ID"))
            if len(narrative_sections) < min(3, len(events)):
                issues.append(Issue("EATING_EVENTS_NOT_DISTRIBUTED", "旧版吃食事件没有分散到至少三个不同叙事段"))

    return issues


def _docx_body_text(document: Any) -> str:
    """Return editable Word paragraph/table text, including nested tables."""
    chunks: list[str] = [paragraph.text for paragraph in document.paragraphs]

    def visit_table(table: Any) -> None:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(paragraph.text for paragraph in cell.paragraphs)
                for nested in cell.tables:
                    visit_table(nested)

    for table in document.tables:
        visit_table(table)
    return "\n".join(chunks)


def validate_final_docx(path: Path) -> list[Issue]:
    """Reject renamed/fake DOCX files and image-only Word deliveries."""
    issues: list[Issue] = []
    if not zipfile.is_zipfile(path):
        return [Issue("FINAL_DOCX_INVALID", f"{path.name} 不是真实 OPC/ZIP DOCX")]

    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            required_parts = {"[Content_Types].xml", "word/document.xml", "word/_rels/document.xml.rels"}
            missing = sorted(required_parts - names)
            if missing:
                return [Issue("FINAL_DOCX_INVALID", f"{path.name} 缺 OPC 核心部件：{'、'.join(missing)}")]

            try:
                content_types = ElementTree.fromstring(archive.read("[Content_Types].xml"))
                document_xml = ElementTree.fromstring(archive.read("word/document.xml"))
                relationships_xml = ElementTree.fromstring(archive.read("word/_rels/document.xml.rels"))
            except ElementTree.ParseError as exc:
                return [Issue("FINAL_DOCX_INVALID", f"{path.name} 的 OPC XML 不可解析：{exc}")]

            word_main_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
            main_declared = any(
                node.attrib.get("PartName") == "/word/document.xml"
                and node.attrib.get("ContentType") == word_main_type
                for node in content_types
            )
            if not main_declared:
                issues.append(Issue("FINAL_DOCX_INVALID", f"{path.name} 未正确声明 Word document.xml 内容类型"))

            image_relationships: dict[str, str] = {}
            for relationship in relationships_xml:
                relation_type = relationship.attrib.get("Type", "")
                relation_id = relationship.attrib.get("Id", "")
                target = relationship.attrib.get("Target", "")
                if relation_type.endswith("/image") and relation_id and target:
                    member = target.lstrip("/") if target.startswith("/") else posixpath.normpath(posixpath.join("word", target))
                    image_relationships[relation_id] = member
            embedded_ids = {
                node.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                for node in document_xml.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
            }
            used_images = {
                relation_id: member
                for relation_id, member in image_relationships.items()
                if relation_id in embedded_ids and member in names
            }
            if not used_images:
                issues.append(
                    Issue(
                        "FINAL_DOCX_IMAGE_MISSING",
                        f"{path.name} 的 document.xml 没有引用任何真实图片关系",
                    )
                )
    except (OSError, zipfile.BadZipFile, KeyError) as exc:
        return [Issue("FINAL_DOCX_INVALID", f"{path.name} 无法作为 DOCX 打开：{exc}")]

    try:
        from docx import Document

        document = Document(str(path))
    except ImportError:
        issues.append(Issue("FINAL_DOCX_INVALID", "当前运行时缺 python-docx，无法证明 Word 可打开"))
        return issues
    except Exception as exc:  # python-docx raises several package/XML exceptions.
        issues.append(Issue("FINAL_DOCX_INVALID", f"python-docx 无法打开 {path.name}：{exc}"))
        return issues

    body_text = _docx_body_text(document)
    editable_label_groups = (
        ("准确秒数", "原片时间"),
        ("分镜描述",),
        ("口播稿",),
        ("即梦可复制 Prompt", "完整Prompt", "完整 Prompt"),
    )
    missing_labels = ["/".join(group) for group in editable_label_groups if not any(label in body_text for label in group)]
    if missing_labels:
        issues.append(
            Issue(
                "FINAL_DOCX_EDITABLE_TEXT_MISSING",
                f"{path.name} 缺原生可编辑正文标签：{'、'.join(missing_labels)}；不得把文字做成图片",
            )
        )
    if len(re.sub(r"\s+", "", body_text)) < 40:
        issues.append(Issue("FINAL_DOCX_EDITABLE_TEXT_MISSING", f"{path.name} 的可编辑正文过少，疑似只放图片"))
    return issues


def lint_delivery_directory(path: Path) -> list[Issue]:
    if not path.is_dir():
        return [Issue("FINAL_DOCX_MISSING", f"用户输出目录不存在：{path}")]
    files = [item for item in path.rglob("*") if item.is_file() and not item.name.startswith(".")]
    docx_files = [item for item in files if item.suffix.lower() == ".docx"]
    issues: list[Issue] = []
    if not docx_files:
        issues.append(Issue("FINAL_DOCX_MISSING", f"用户输出目录没有 DOCX：{path}"))
    elif len(docx_files) > 1:
        issues.append(Issue("FINAL_DOCX_COUNT_MISMATCH", f"用户输出目录有{len(docx_files)}个 DOCX，必须只保留唯一终稿"))
    else:
        issues.extend(validate_final_docx(docx_files[0]))
    leaked = [item.name for item in files if item.suffix.lower() != ".docx"]
    if leaked:
        issues.append(Issue("USER_DELIVERY_ARTIFACT_LEAK", f"用户输出目录混入非 DOCX：{'、'.join(sorted(leaked)[:8])}"))
    return issues


def lint(
    path: Path,
    role_lock_path: Path | None = None,
    min_prompt_chars: int | None = None,
    max_prompt_chars: int | None = None,
    story_plan_path: Path | None = None,
    delivery_dir: Path | None = None,
    *,
    stage: str | None = None,
    text_handoff_path: Path | None = None,
    enforce_prompt_length: bool | None = None,
) -> list[Issue]:
    text = path.read_text(encoding="utf-8-sig")
    shots = split_shots(text)
    if not shots:
        return [Issue("TXT_EXPORT_MISSING", "未找到 S001｜… 形式的分镜标题")]

    issues: list[Issue] = lint_stage_inputs(stage, story_plan_path, text_handoff_path, delivery_dir)
    story_plan = load_story_plan(story_plan_path)
    length_issues, effective_min, effective_max = resolve_prompt_length_contract(
        story_plan,
        enforce_prompt_length,
        min_prompt_chars,
        max_prompt_chars,
    )
    issues.extend(length_issues)
    prompts: dict[str, str] = {}
    all_lines: list[tuple[str, str, str]] = []
    for title, block in shots:
        shot_issues, prompt, lines = lint_shot(
            title,
            block,
            min_prompt_chars=effective_min,
            max_prompt_chars=effective_max,
        )
        issues.extend(shot_issues)
        sid = shot_id(title)
        prompts[sid] = prompt
        all_lines.extend((sid, label, text) for label, text in lines)

    role_lock = load_role_lock(role_lock_path)
    if role_lock is not None:
        issues.extend(lint_role_lock(shots, prompts, all_lines, role_lock))
    if story_plan is not None:
        issues.extend(lint_story_plan(shots, prompts, story_plan))
    text_handoff = load_text_handoff(text_handoff_path)
    if text_handoff is not None:
        issues.extend(Issue(issue.code, issue.message) for issue in validate_text_handoff(text_handoff))
        if story_plan is not None:
            issues.extend(lint_text_handoff_against_story_plan(text_handoff, story_plan))
    if delivery_dir is not None and stage in (None, "full_delivery_postexport"):
        issues.extend(lint_delivery_directory(delivery_dir))

    seen: set[tuple[str, str]] = set()
    unique: list[Issue] = []
    for issue in issues:
        key = (issue.code, issue.message)
        if key not in seen:
            seen.add(key)
            unique.append(issue)
    return unique


def main() -> int:
    parser = argparse.ArgumentParser(description="检查逐分镜 Prompt 的结构和可编码的角色锁冲突")
    parser.add_argument("txt", type=Path)
    parser.add_argument(
        "--stage",
        required=True,
        choices=("text_branch", "full_delivery_precompile", "full_delivery_postexport"),
        help="显式声明文分支、导出前或DOCX导出后阶段",
    )
    parser.add_argument("--role-lock", type=Path, default=None)
    parser.add_argument("--story-plan", type=Path, default=None)
    parser.add_argument("--text-handoff", type=Path, default=None, help="text_branch 必需的 text-handoff-v2.0 JSON")
    parser.add_argument("--delivery-dir", type=Path, default=None, help="full_delivery 用户侧输出目录；必须只有唯一 DOCX")
    parser.add_argument(
        "--enforce-prompt-length",
        action="store_const",
        const=True,
        default=None,
        help="显式启用长度契约；若 story_plan 已声明契约，必须与其一致",
    )
    parser.add_argument(
        "--min-prompt-chars",
        type=int,
        default=None,
        help="已启用长度契约时的 Prompt 非空白字符下限；默认3000",
    )
    parser.add_argument(
        "--max-prompt-chars",
        type=int,
        default=None,
        help="已启用长度契约时的 Prompt 非空白字符上限；默认4000",
    )
    args = parser.parse_args()
    if not args.txt.is_file():
        print(f"文件不存在：{args.txt}", file=sys.stderr)
        return 2
    if args.role_lock is not None and not args.role_lock.is_file():
        print(f"角色锁不存在：{args.role_lock}", file=sys.stderr)
        return 2
    if args.story_plan is not None and not args.story_plan.is_file():
        print(f"故事计划不存在：{args.story_plan}", file=sys.stderr)
        return 2
    if args.text_handoff is not None and not args.text_handoff.is_file():
        print(f"text handoff 不存在：{args.text_handoff}", file=sys.stderr)
        return 2
    try:
        issues = lint(
            args.txt,
            args.role_lock,
            min_prompt_chars=args.min_prompt_chars,
            max_prompt_chars=args.max_prompt_chars,
            story_plan_path=args.story_plan,
            delivery_dir=args.delivery_dir,
            stage=args.stage,
            text_handoff_path=args.text_handoff,
            enforce_prompt_length=args.enforce_prompt_length,
        )
    except (OSError, ValueError, json.JSONDecodeError) as exc:
        print(f"无法检查：{exc}", file=sys.stderr)
        return 2
    if issues:
        print(f"发现结构性阻断：{len(issues)}项")
        for issue in issues:
            print(f"- [{issue.code}] {issue.message}")
        return 1
    print("未发现结构性阻断。此结果只说明可自动检查的结构与角色锁一致，不代表内容审核完成。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
